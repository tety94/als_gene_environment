"""
build_annotated_tables.py

Builds Word (.docx) tables from the `get_annotated_results` stored
procedure output, split by `neuro_plausibility_score`:

    - rows with neuro_plausibility_score > 0 are listed FIRST
    - rows with neuro_plausibility_score == 0 are listed SECOND

Two documents are produced:

    supplementary_tables.docx
        Full annotation: exposure, gene_name, variant + every column
        from gene_neuro_annotation (except the bookkeeping
        `last_updated` timestamp, which isn't meaningful in a paper).

    main_text_tables.docx
        Compact subset for the main text: exposure, gene_name, variant,
        gene_symbol, neuro_plausibility_score.

Rows with a NULL neuro_plausibility_score fall into neither table; they
are counted and logged, not silently dropped.

Usage
-----
    python build_annotated_tables.py
    python build_annotated_tables.py --outdir ./tables
    python build_annotated_tables.py --input-csv results.csv   # bypass the DB call

Requires: python-docx, pandas
    pip install python-docx pandas
"""

from __future__ import annotations

import argparse
import logging
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

log = logging.getLogger("build_annotated_tables")
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# --------------------------------------------------------------------------
# Exposure name mapping (Italian land-use category -> English)
# --------------------------------------------------------------------------

EXPOSURE_LABELS = {
    "seminativi": "Arable land",
    "vigneti": "Vineyards",
    "risaie": "Rice fields",
}


def translate_exposure(df: pd.DataFrame) -> pd.DataFrame:
    """Map the `exposure` column from Italian land-use terms to English.
    Values not found in EXPOSURE_LABELS are left unchanged, with a warning
    so unmapped categories don't silently slip into the paper."""
    df = df.copy()
    unmapped = sorted(set(df["exposure"].dropna()) - set(EXPOSURE_LABELS))
    if unmapped:
        log.warning("Exposure values with no English mapping (left as-is): %s", unmapped)
    df["exposure"] = df["exposure"].map(lambda v: EXPOSURE_LABELS.get(v, v))
    return df


# --------------------------------------------------------------------------
# Column sets
# --------------------------------------------------------------------------

FULL_COLUMNS = [
    "exposure",
    "gene_name",
    "variant",
    "gene_id",
    "gene_symbol",
    "gene_type",
    "expressed_brain",
    "brain_tissues",
    "expressed_neurons",
    "expressed_glia",
    "cell_types",
    "go_neuro_processes",
    "go_toxic_response",
    "ctd_chemicals",
    "ctd_neuro_diseases",
    "ctd_neuro_disease_direct",
    "ctd_neuro_disease_pesticide_mediated",
    "als_panelapp_confidence",
    "als_opentargets_score",
    "neuro_plausibility_score",
]

SUBSET_COLUMNS = [
    "exposure",
    "gene_name",
    "variant",
    "gene_symbol",
    "neuro_plausibility_score",
]

# tinyint(1) columns that should render as Yes/No rather than 1/0
BOOLEAN_COLUMNS = {
    "expressed_brain",
    "expressed_neurons",
    "expressed_glia",
    "ctd_neuro_disease_direct",
    "ctd_neuro_disease_pesticide_mediated",
}

HEADER_LABELS = {
    "exposure": "Exposure",
    "gene_name": "Gene name",
    "variant": "Variant",
    "gene_id": "Gene ID",
    "gene_symbol": "Gene symbol",
    "gene_type": "Gene type",
    "expressed_brain": "Expressed in brain",
    "brain_tissues": "Brain tissues",
    "expressed_neurons": "Expressed in neurons",
    "expressed_glia": "Expressed in glia",
    "cell_types": "Cell types",
    "go_neuro_processes": "GO neuro processes",
    "go_toxic_response": "GO toxic response",
    "ctd_chemicals": "CTD chemicals",
    "ctd_neuro_diseases": "CTD neuro diseases",
    "ctd_neuro_disease_direct": "CTD neuro disease (direct)",
    "ctd_neuro_disease_pesticide_mediated": "CTD neuro disease (pesticide-mediated)",
    "als_panelapp_confidence": "ALS PanelApp confidence",
    "als_opentargets_score": "ALS Open Targets score",
    "neuro_plausibility_score": "Neuro plausibility score",
}


# --------------------------------------------------------------------------
# Formatting helpers
# --------------------------------------------------------------------------

def format_cell(col: str, value) -> str:
    """Render a raw DB value as clean table text for a given column."""
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""

    if col in BOOLEAN_COLUMNS:
        if value in (1, True, "1"):
            return "Yes"
        if value in (0, False, "0"):
            return "No"
        return ""

    if isinstance(value, float):
        return str(int(value)) if value.is_integer() else f"{value:.3f}"

    return str(value)


def set_repeat_header(row) -> None:
    """Mark a table row so it repeats on every page it spans."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def shade_cell(cell, fill: str = "D9D9D9") -> None:
    """Apply flat grey shading to a table cell (never use w:val=SOLID)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


# --------------------------------------------------------------------------
# Table building
# --------------------------------------------------------------------------

def set_landscape(doc: Document, margin: float = 0.5) -> None:
    """Switch the document's section to landscape and use narrower margins
    (in inches) so wide tables have more room."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)


def add_table(doc: Document, title: str, df: pd.DataFrame, columns: list[str]) -> None:
    doc.add_heading(title, level=2)

    if df.empty:
        doc.add_paragraph("No rows in this category.")
        return

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.autofit = True

    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = HEADER_LABELS.get(col, col)
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
        shade_cell(header_cells[i])
    set_repeat_header(table.rows[0])

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            cells[i].text = format_cell(col, row.get(col))

    doc.add_paragraph()  # spacing after the table


def build_document(
    df: pd.DataFrame, columns: list[str], out_path: Path, doc_title: str, landscape: bool = False
) -> None:
    pos = df[df["neuro_plausibility_score"] > 0].reset_index(drop=True)
    zero = df[df["neuro_plausibility_score"] == 0].reset_index(drop=True)
    null_score = df[df["neuro_plausibility_score"].isna()]

    if not null_score.empty:
        log.warning(
            "%d rows have a NULL neuro_plausibility_score and were excluded "
            "from both tables in %s",
            len(null_score), out_path.name,
        )

    doc = Document()
    if landscape:
        set_landscape(doc)
    doc.add_heading(doc_title, level=1)

    add_table(doc, "Genes with neuro plausibility score > 0", pos, columns)
    add_table(doc, "Genes with neuro plausibility score = 0", zero, columns)

    doc.save(out_path)
    log.info(
        "Wrote %s (%d rows score>0, %d rows score=0, %d excluded/NULL)",
        out_path, len(pos), len(zero), len(null_score),
    )


# --------------------------------------------------------------------------
# Entry point
# --------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument(
        "--input-csv", type=Path, default=None,
        help="Skip the DB/stored-procedure call and load results from a CSV "
             "dump instead (must have the same columns as get_annotated_results()).",
    )
    parser.add_argument(
        "--outdir", type=Path, default=Path("."),
        help="Directory to write the .docx files into (default: current directory).",
    )
    args = parser.parse_args()
    args.outdir.mkdir(parents=True, exist_ok=True)

    if args.input_csv is not None:
        df = pd.read_csv(args.input_csv)
    else:
        from gene_environment.db.repository import get_annotated_results
        df = get_annotated_results()

    if df.empty:
        log.warning("No rows returned — nothing to write.")
        return

    if "neuro_plausibility_score" not in df.columns:
        raise RuntimeError("Expected column 'neuro_plausibility_score' not found in the results.")

    df = translate_exposure(df)

    build_document(
        df, FULL_COLUMNS,
        args.outdir / "supplementary_tables.docx",
        "Supplementary Tables \u2014 Full Gene Neuro-Annotation",
        landscape=True,
    )
    build_document(
        df, SUBSET_COLUMNS,
        args.outdir / "main_text_tables.docx",
        "Main Text Tables \u2014 Gene Neuro-Annotation Summary",
    )


if __name__ == "__main__":
    main()
# -*- coding: utf-8 -*-
"""
Generate Table 2 Word files and figures by calling the stored routine
get_significant_results_table_2 via the project's MySQL connection helpers.

Usage (from project root):
    python3 -m gene_environment.report.generate_table2 [--generation 2] [--alpha 0.05]

Requirements:
    pip install pandas python-docx matplotlib seaborn mysql-connector-python regex scipy statsmodels

Behavior:
- Calls the stored routine `get_significant_results_table_2()` and uses the first
  resultset returned.
- Drops gna.* columns before producing outputs.
- Translates the `exposure` column from the source dataset's Italian land-use
  terms to English (see `gene_environment.report.exposure_labels`) before
  anything is written to a table, figure, or filename.
- Produces:
    output/table2/Table2_top10.docx
    output/table2/Table2_full_supplementary.docx   (now includes all figures embedded)
    output/table2/figures/variants_per_chromosome.png
    output/table2/figures/genes_vs_variants_scatter.png
    output/table2/figures/empirical_p_g1_histogram.png
    output/table2/figures/observed_vs_expected_by_chromosome.png
    output/table2/figures/observed_vs_expected_by_chromosome_per_exposure.png
    output/table2/figures/by_exposure/observed_vs_expected_chrom_<exposure>.png  (one per exposure)
    output/table2/figures/by_exposure/pvalues_<exposure>.csv                    (one per exposure)
    output/table2/table2_chromosome_enrichment_stats.csv
    output/table2/table2_chromosome_enrichment_by_exposure_stats.csv
- Numeric formatting: p-values 3 significant digits, coefficients 2 decimals.

Statistical model for the per-chromosome enrichment test (both the pooled
"all exposures" version and the per-exposure version):
    For a given chromosome (and, in the per-exposure analysis, a given
    exposure), let
        n = number of variants TESTED on that chromosome
            (COUNT(*) FROM variant_results WHERE exposure=... AND
             generation=... GROUP BY chromosome -- i.e. exactly the query
             used by fetch_tested_variant_counts_by_chromosome)
        k = number of those variants found SIGNIFICANT
            (empirical_p_g1 < alpha)
    The null model is PROPORTIONAL ALLOCATION: within the table being
    analyzed (all chromosomes of one exposure, or all chromosomes pooled
    across exposures), the overall observed significance rate is
        p_rate = sum(k) / sum(n)   [i.e. total_significant / total_tested]
    and each chromosome's expected count of significant hits, under a null
    of "no chromosome-level concentration", is proportional to how much
    testing was done there:
        expected = n * p_rate
    The per-chromosome p-value is a one-sided ("greater") exact binomial
    test:
        binom_p = P(X >= k | X ~ Binomial(n, p_rate))
    i.e. we are testing whether a chromosome concentrates more significant
    hits than its share of testing volume would predict, not the raw
    significance threshold alpha. Within each table (per exposure, or
    pooled), binom_p is BH-adjusted (binom_p_adj) across chromosomes.
    NOTE: an earlier version of this analysis used alpha itself as the null
    probability (expected = n * alpha). That produced p-values of 1.0 across
    every chromosome, because get_significant_results_table_2 already
    returns a far more stringent subset than a naive
    empirical_p_g1 < alpha filter would (the true significance rate in this
    dataset is far below alpha, e.g. ~0.04% vs alpha = 5%), so alpha is not
    a meaningful "by chance" baseline for these numbers. The proportional
    model above is calibrated to the data instead and is the current,
    correct version.

Note on exposure translation and DB queries: the tested-variant-count
queries (`fetch_tested_variant_counts_by_chromosome`) filter
`WHERE exposure = %s` against the *raw* (Italian) values stored in the
DB, so the exposure column is only translated to English AFTER those
queries have run, right before anything is written out for the report
(tables, figure titles, filenames).
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import Dict, List, Optional

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from scipy.stats import binomtest
from statsmodels.stats.multitest import multipletests

from gene_environment.report.db_utils import (
    call_stored_routine_to_df,
    chrom_sort_key,
    extract_chromosome,
    normalize_chrom_label,
    slugify,
)
from gene_environment.report.exposure_labels import translate_exposure, translate_exposure_value
from gene_environment.report.word_utils import (
    add_figure_to_doc,
    repeat_header_row,
    set_cell_bg,
    set_col_width,
    set_table_borders,
)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OUT_DIR = Path("output/table2")
FIG_DIR = OUT_DIR / "figures"

ASTORE_NAME = "get_significant_results_table_2"

# Columns to include in Word tables (order).
# NOTE: "empirical_p_2" (no "g") is kept exactly as in the original script.
# Every other column follows the "_g1"/"_g2" convention, so this looks like it
# could be a typo for "empirical_p_g2" -- please confirm against the real
# astore output before renaming it, otherwise the column will silently come
# back empty in the report.
TABLE_COLUMNS = [
    "exposure",
    "variant",
    "empirical_p_g1",
    "obs_coef_g1",
    "muted_g1",
    "not_muted_g1",
    "empirical_p_2",
    "obs_coef_g2",
    "muted_g2",
    "not_muted_g2",
]

# Nicer header labels for the Word table (falls back to the raw column name)
COLUMN_LABELS = {
    "exposure": "Exposure",
    "variant": "Variant",
    "empirical_p_g1": "Emp. p (G1)",
    "obs_coef_g1": "Coef. (G1)",
    "muted_g1": "Muted (G1)",
    "not_muted_g1": "Not muted (G1)",
    "empirical_p_2": "Emp. p (G2)",
    "obs_coef_g2": "Coef. (G2)",
    "muted_g2": "Muted (G2)",
    "not_muted_g2": "Not muted (G2)",
}

# Column widths in inches
COL_WIDTHS_IN = {
    "exposure": 1.5,
    "variant": 2.2,
    "empirical_p_g1": 0.85,
    "obs_coef_g1": 0.75,
    "muted_g1": 0.75,
    "not_muted_g1": 0.85,
    "empirical_p_2": 0.85,
    "obs_coef_g2": 0.75,
    "muted_g2": 0.75,
    "not_muted_g2": 0.85,
}

NUMERIC_PREFIXES = ("empirical_p", "obs_coef", "muted", "not_muted")
SIG_ALPHA_DEFAULT = 0.05

# Table color scheme
HEADER_FILL = "44546A"        # dark blue-grey
HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_FILL = "EEF1F6"         # very light blue-grey
SIG_FONT_COLOR = RGBColor(0xC0, 0x00, 0x00)  # highlight p < alpha


# ---------------------------------------------------------------------------
# docx table helpers
# ---------------------------------------------------------------------------

def add_table_to_doc(
    doc: Document,
    df: pd.DataFrame,
    title: Optional[str] = None,
    caption: Optional[str] = None,
    max_rows: Optional[int] = None,
    alpha: float = SIG_ALPHA_DEFAULT,
) -> None:
    """Add a formatted table to a python-docx Document.

    - Dark header band with white bold text and friendly column labels
    - Fixed column widths (in inches, correctly converted to twips)
    - Numeric columns right-aligned, text columns left-aligned
    - Zebra row shading + thin consistent borders
    - empirical_p_* values below `alpha` are bolded/highlighted
    - Header row repeats on each printed page
    """
    if title:
        h = doc.add_heading(title, level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    df_to_use = df[TABLE_COLUMNS]
    if max_rows is not None:
        df_to_use = df_to_use.head(max_rows)

    table = doc.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(TABLE_COLUMNS):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(COLUMN_LABELS.get(col, col))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = HEADER_FONT_COLOR
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], HEADER_FILL)
        set_col_width(hdr_cells[i], COL_WIDTHS_IN.get(col, 1.0))

    # Data rows
    for ridx, (_, row) in enumerate(df_to_use.iterrows()):
        cells = table.add_row().cells
        shade = ZEBRA_FILL if ridx % 2 == 1 else "FFFFFF"
        for i, col in enumerate(TABLE_COLUMNS):
            raw_val = row.get(col, pd.NA)
            text = format_value_for_word(col, raw_val)
            cell = cells[i]
            set_col_width(cell, COL_WIDTHS_IN.get(col, 1.0))
            para = cell.paragraphs[0]
            para.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT
                if col.startswith(NUMERIC_PREFIXES)
                else WD_PARAGRAPH_ALIGNMENT.LEFT
            )
            run = para.add_run(text)
            run.font.size = Pt(9)

            if col.startswith("empirical_p") and _is_significant(raw_val, alpha):
                run.bold = True
                run.font.color.rgb = SIG_FONT_COLOR

            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, shade)

    repeat_header_row(table)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def _is_significant(val, alpha: float) -> bool:
    try:
        return pd.notna(val) and float(val) < alpha
    except (TypeError, ValueError):
        return False


def add_pvalue_table_to_doc(
    doc: Document,
    df: pd.DataFrame,
    title: Optional[str] = None,
    caption: Optional[str] = None,
    alpha: float = SIG_ALPHA_DEFAULT,
) -> None:
    """Add a compact chromosome-level enrichment table (chromosome, n_tested,
    n_significant_observed, expected, binom_p, binom_p_adj) to a docx.
    binom_p_adj values below `alpha` are highlighted, same convention as
    add_table_to_doc.
    """
    cols = ["chromosome", "n_tested", "n_significant_observed", "expected", "binom_p", "binom_p_adj"]
    labels = {
        "chromosome": "Chromosome",
        "n_tested": "N tested",
        "n_significant_observed": "N significant (obs.)",
        "expected": "Expected (n\u00b7p_rate)",
        "binom_p": "Binomial p",
        "binom_p_adj": "Binomial p (BH-adj.)",
    }

    if title:
        h = doc.add_heading(title, level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    if df.empty:
        doc.add_paragraph("No tested-variant data available for this exposure.")
        return

    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(labels[col])
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = HEADER_FONT_COLOR
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], HEADER_FILL)

    for ridx, (_, row) in enumerate(df[cols].iterrows()):
        cells = table.add_row().cells
        shade = ZEBRA_FILL if ridx % 2 == 1 else "FFFFFF"
        for i, col in enumerate(cols):
            val = row[col]
            if col in ("binom_p", "binom_p_adj", "expected"):
                text = "{:.3g}".format(float(val)) if pd.notna(val) else ""
            else:
                text = str(val) if pd.notna(val) else ""
            para = cells[i].paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if col != "chromosome" else WD_PARAGRAPH_ALIGNMENT.LEFT
            run = para.add_run(text)
            run.font.size = Pt(9)
            if col == "binom_p_adj" and _is_significant(val, alpha):
                run.bold = True
                run.font.color.rgb = SIG_FONT_COLOR
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cells[i], shade)

    repeat_header_row(table)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


# ---------------------------------------------------------------------------
# Data access
# ---------------------------------------------------------------------------

def drop_gna_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Drop gna.* columns if present."""
    to_drop = []
    for col in df.columns:
        if col in ("neuro_plausibility_score", "expressed_neurons"):
            to_drop.append(col)
        if isinstance(col, str) and col.startswith("gna."):
            to_drop.append(col)
    return df.drop(columns=to_drop, errors="ignore") if to_drop else df


def fetch_tested_variant_counts_by_chromosome(get_connection, cursor_scope, exposure: str, generation: int) -> Dict[str, int]:
    """
    Tested-variant counts per chromosome for a single exposure, mirroring:
        SELECT chromosome, COUNT(*)
        FROM variant_results
        WHERE exposure = <exposure> AND generation = <generation>
        GROUP BY chromosome
    `exposure` must be the RAW (source-language) value here, since it is
    matched against the DB, not the translated display label.
    """
    sql = (
        "SELECT chromosome, COUNT(*) "
        "FROM variant_results "
        "WHERE exposure = %s AND generation = %s "
        "GROUP BY chromosome"
    )
    counts: Dict[str, int] = {}
    with get_connection() as conn:
        with cursor_scope(conn) as cur:
            cur.execute(sql, (exposure, generation))
            for chrom, cnt in cur.fetchall():
                if chrom is None:
                    continue
                label = normalize_chrom_label(chrom)
                counts[label] = counts.get(label, 0) + int(cnt)
    return counts


def fetch_tested_variant_counts_for_exposures(get_connection, cursor_scope, exposures: List[str], generation: int) -> pd.DataFrame:
    """
    Run `fetch_tested_variant_counts_by_chromosome` once per exposure and
    stack the results into a DataFrame [exposure, chromosome, n_tested].
    Returns an empty DataFrame (with the right columns) if there are no
    exposures or no rows come back for any of them.
    """
    rows: List[dict] = []
    for exposure in exposures:
        counts = fetch_tested_variant_counts_by_chromosome(get_connection, cursor_scope, exposure, generation)
        for chrom, n_tested in counts.items():
            rows.append({"exposure": exposure, "chromosome": chrom, "n_tested": n_tested})

    if not rows:
        return pd.DataFrame(columns=["exposure", "chromosome", "n_tested"])
    return pd.DataFrame(rows)


# ---------------------------------------------------------------------------
# Formatting / parsing helpers
# ---------------------------------------------------------------------------

def format_value_for_word(col: str, val) -> str:
    """empirical_p_*: 3 significant digits; obs_coef_*: 2 decimals; else str()."""
    if pd.isna(val):
        return ""
    try:
        if col.startswith("empirical_p"):
            return "{:.3g}".format(float(val))
        if col.startswith("obs_coef"):
            return "{:.2f}".format(float(val))
    except (TypeError, ValueError):
        return str(val)
    return str(val)


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------

def make_figures(df: pd.DataFrame, fig_dir: Path) -> None:
    """Chromosome-level descriptive figures (unchanged grouping)."""
    df = df.copy()
    df["chromosome"] = df["variant"].apply(lambda v: extract_chromosome(v) if pd.notna(v) else "NA")

    variants_per_chrom = df.groupby("chromosome")["variant"].nunique().rename("n_variants").reset_index()
    genes_per_chrom = df.groupby("chromosome")["exposure"].nunique().rename("n_genes").reset_index()
    merged = pd.merge(variants_per_chrom, genes_per_chrom, on="chromosome", how="outer").fillna(0)

    merged = merged.sort_values(by="chromosome", key=lambda s: s.map(chrom_sort_key))

    sns.set_theme(style="whitegrid")

    plt.figure(figsize=(10, 6))
    ax = sns.barplot(data=merged, x="chromosome", y="n_variants", color="#4472C4")
    ax.set_xlabel("Chromosome")
    ax.set_ylabel("Number of unique variants")
    ax.set_title("Variants per chromosome")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(fig_dir / "variants_per_chromosome.png", dpi=300)
    plt.close()

    plt.figure(figsize=(8, 6))
    ax = sns.scatterplot(data=merged, x="n_genes", y="n_variants", s=100, color="#4472C4")
    for _, r in merged.iterrows():
        ax.text(r["n_genes"], r["n_variants"], str(r["chromosome"]), fontsize=9,
                 horizontalalignment="left", verticalalignment="bottom")
    ax.set_xlabel("Number of unique genes (exposures) per chromosome")
    ax.set_ylabel("Number of unique variants per chromosome")
    ax.set_title("Genes vs variants per chromosome")
    plt.tight_layout()
    plt.savefig(fig_dir / "genes_vs_variants_scatter.png", dpi=300)
    plt.close()

    plt.figure(figsize=(8, 5))
    pvals = pd.to_numeric(df.get("empirical_p_g1", pd.Series(dtype=float)), errors="coerce").dropna()
    hist_path = fig_dir / "empirical_p_g1_histogram.png"
    if not pvals.empty:
        ax = sns.histplot(pvals, bins=50, kde=False, color="#548235")
        ax.set_xlabel("empirical_p_g1")
        ax.set_ylabel("Count")
        ax.set_title("Histogram of empirical_p_g1")
        plt.tight_layout()
    else:
        plt.text(0.5, 0.5, "No empirical_p_g1 data available", ha="center", va="center")
        plt.axis("off")
    plt.savefig(hist_path, dpi=300)
    plt.close()


def _empty_placeholder_figure(path: Path, message: str) -> None:
    plt.figure(figsize=(8, 5))
    plt.text(0.5, 0.5, message, ha="center", va="center", fontsize=12)
    plt.axis("off")
    plt.savefig(path, dpi=300)
    plt.close()


def _add_binomial_enrichment_stats(merged: pd.DataFrame) -> pd.DataFrame:
    """Add expected/binomial-p columns to a per-chromosome table with
    n_tested / n_significant_observed columns.

    Null model: proportional allocation. Given the table's own totals
    (summed across all its chromosome rows -- e.g. all chromosomes within
    one exposure, or all chromosomes pooled across exposures), the overall
    significance rate is
        p_rate = sum(n_significant_observed) / sum(n_tested)
    Under H0 (no chromosome-level concentration of hits), each chromosome's
    expected count of significant hits is proportional to how much of it was
    tested:
        expected  = n_tested * p_rate
        binom_p   = P(X >= n_significant_observed | X ~ Binomial(n_tested, p_rate))
    (one-sided "greater" exact binomial test -- testing whether this
    chromosome concentrates more hits than its share of testing volume would
    predict). binom_p_adj is the BH (fdr_bh) correction of binom_p across the
    rows of this table.

    NOTE: this intentionally does NOT use the raw significance threshold
    alpha as the null rate -- in this dataset the actual significance rate is
    far below alpha (e.g. get_significant_results_table_2 already returns a
    much more stringent subset than a naive empirical_p_g1 < 0.05 filter
    would), so alpha itself is not a meaningful "by chance" baseline here.
    """
    merged = merged.copy()
    total_tested = int(merged["n_tested"].sum())
    total_sig = int(merged["n_significant_observed"].sum())

    if total_tested == 0 or total_sig == 0:
        merged["expected"] = 0.0
        merged["binom_p"] = 1.0
        merged["binom_p_adj"] = 1.0
        return merged

    p_rate = total_sig / total_tested
    merged["expected"] = merged["n_tested"] * p_rate

    binom_pvals = []
    for _, row in merged.iterrows():
        n, k = int(row["n_tested"]), int(row["n_significant_observed"])
        if n <= 0:
            binom_pvals.append(1.0)
            continue
        try:
            binom_pvals.append(binomtest(k, n, p_rate, alternative="greater").pvalue)
        except ValueError:
            binom_pvals.append(1.0)
    merged["binom_p"] = binom_pvals

    try:
        _, p_adj, _, _ = multipletests(merged["binom_p"].fillna(1.0).values, method="fdr_bh")
        merged["binom_p_adj"] = p_adj
    except ValueError:
        merged["binom_p_adj"] = merged["binom_p"]
    return merged


def _draw_chrom_bars(ax, merged: pd.DataFrame, title: str) -> None:
    """Draw one observed-vs-expected-per-chromosome panel on `ax`."""
    x_labels = merged["chromosome"].astype(str).tolist()
    xi = range(len(x_labels))
    bar_w = 0.4

    ax.bar([i - bar_w / 2 for i in xi], merged["n_significant_observed"], width=bar_w,
           label="Observed", color="#4472C4")
    ax.bar([i + bar_w / 2 for i in xi], merged["expected"], width=bar_w,
           label="Expected (n\u00b7p_rate)", color="#ED7D31", alpha=0.85)
    ax.set_xticks(list(xi))
    ax.set_xticklabels(x_labels, rotation=45, ha="right", fontsize=8)
    ax.set_ylabel("Count", fontsize=9)
    ax.set_title(title, fontsize=10)
    ax.legend(fontsize=7)
    ax.set_ylim(bottom=0)


def compute_chromosome_enrichment_global(
    df: pd.DataFrame,
    tested_df: pd.DataFrame,
    alpha: float = SIG_ALPHA_DEFAULT,
    out_dir: Path = OUT_DIR,
):
    """
    Observed vs. expected significant variants *per chromosome*, pooling all
    exposures together. `tested_df` (exposure, chromosome, n_tested) is summed
    across exposure to get the total number of variants ever tested on each
    chromosome. Expected count of significant hits per chromosome is
    n_tested * (total_significant / total_tested) -- the proportional null
    model described in the module docstring -- tested with a one-sided
    binomial test per chromosome, BH-adjusted across chromosomes.

    Chromosomes with zero significant hits still appear (zero-height bar).
    """
    fig_path = out_dir / "figures" / "observed_vs_expected_by_chromosome.png"
    stats_path = out_dir / "table2_chromosome_enrichment_stats.csv"

    if tested_df.empty:
        print("[warn] no tested-variant counts -- skipping chromosome enrichment stats.", file=sys.stderr)
        _empty_placeholder_figure(fig_path, "No tested-variant counts available")
        empty = pd.DataFrame(columns=["chromosome", "n_tested", "n_significant_observed",
                                       "expected", "binom_p", "binom_p_adj", "sig_variants"])
        empty.to_csv(stats_path, index=False)
        return empty, {"total_tested": 0, "total_significant": 0, "overall_rate": None,
                        "per_chromosome_csv": str(stats_path), "figure": str(fig_path)}

    tested_by_chrom = tested_df.groupby("chromosome", as_index=False)["n_tested"].sum()

    df = df.copy()
    df["chromosome"] = df["variant"].apply(lambda v: extract_chromosome(v) if pd.notna(v) else "NA")
    df["empirical_p_g1_num"] = pd.to_numeric(df.get("empirical_p_g1", pd.Series(dtype=float)), errors="coerce")
    df["is_sig_raw"] = df["empirical_p_g1_num"] < alpha

    sig_df = df.loc[df["is_sig_raw"]]
    obs_counts = sig_df.groupby("chromosome")["variant"].nunique().rename("n_significant_observed").reset_index()
    sig_lists = sig_df.groupby("chromosome")["variant"].apply(list).rename("sig_variants").reset_index()

    merged = tested_by_chrom.merge(obs_counts, on="chromosome", how="left")
    merged = merged.merge(sig_lists, on="chromosome", how="left")
    merged["n_significant_observed"] = merged["n_significant_observed"].fillna(0).astype(int)
    merged["sig_variants"] = merged["sig_variants"].apply(lambda v: v if isinstance(v, list) else [])
    merged["n_tested"] = merged["n_tested"].astype(int)

    total_tested = int(merged["n_tested"].sum())
    total_sig = int(merged["n_significant_observed"].sum())

    if total_tested == 0:
        print("[warn] total tested variants is zero -- skipping enrichment stats.", file=sys.stderr)
        _empty_placeholder_figure(fig_path, "No tested variants for this generation")
        merged.to_csv(stats_path, index=False)
        return merged, {"total_tested": 0, "total_significant": total_sig, "overall_rate": None,
                         "per_chromosome_csv": str(stats_path), "figure": str(fig_path)}

    merged = _add_binomial_enrichment_stats(merged)

    # NOTE: there is no meaningful "overall" p-value here beyond the
    # per-chromosome binomial tests above. Under the proportional null model,
    # expected is defined so that sum(expected) == sum(observed) by
    # construction (expected = n_tested * (total_sig/total_tested)), so a
    # pooled test on the totals would be tautological. The real answer to
    # "is there chromosome-level concentration of hits" lives in the
    # per-chromosome binom_p / binom_p_adj columns.
    merged = merged.sort_values(by="chromosome", key=lambda s: s.map(chrom_sort_key))
    merged.to_csv(stats_path, index=False)

    plt.figure(figsize=(max(8, 0.5 * len(merged)), 6))
    ax = plt.gca()
    _draw_chrom_bars(ax, merged, title="Observed vs expected significant variants per chromosome (all exposures)")
    plt.tight_layout()
    plt.savefig(fig_path, dpi=300)
    plt.close()

    summary = {
        "total_tested": total_tested,
        "total_significant": total_sig,
        "overall_rate": (total_sig / total_tested) if total_tested else None,
        "per_chromosome_csv": str(stats_path),
        "figure": str(fig_path),
    }
    return merged, summary


def compute_chromosome_enrichment_by_exposure(
    df: pd.DataFrame,
    tested_df: pd.DataFrame,
    alpha: float = SIG_ALPHA_DEFAULT,
    out_dir: Path = OUT_DIR,
):
    """
    Same idea as `compute_chromosome_enrichment_global`, but computed
    separately for each exposure (its own tested-per-chromosome baseline
    from `tested_df`, its own observed significant variants), and rendered
    as one grid figure with one panel per exposure plus one standalone
    figure + CSV per exposure. Expected count and p-value per chromosome use
    the proportional null model described in the module docstring
    (expected = n_tested * (exposure's own total_significant/total_tested);
    one-sided binomial test), BH-adjusted across chromosomes within each
    exposure.

    Exposures with no tested-variant rows for this generation are skipped
    (nothing to compare against); exposures with tested variants but zero
    significant hits still get a panel, with every bar at zero height.

    Chart titles and filenames use the English exposure label
    (translate_exposure_value); grouping/matching against `df` and
    `tested_df` still uses the raw exposure value passed in from `main`.
    """
    fig_path = out_dir / "figures" / "observed_vs_expected_by_chromosome_per_exposure.png"
    stats_path = out_dir / "table2_chromosome_enrichment_by_exposure_stats.csv"

    if tested_df.empty:
        print("[warn] no tested-variant counts -- skipping per-exposure chromosome enrichment.", file=sys.stderr)
        _empty_placeholder_figure(fig_path, "No tested-variant counts available")
        empty = pd.DataFrame(columns=["exposure", "chromosome", "n_tested", "n_significant_observed",
                                       "expected", "binom_p", "binom_p_adj", "sig_variants"])
        empty.to_csv(stats_path, index=False)
        return empty, {"n_exposures_plotted": 0, "per_exposure_csv": str(stats_path), "figure": str(fig_path)}

    df = df.copy()
    df["chromosome"] = df["variant"].apply(lambda v: extract_chromosome(v) if pd.notna(v) else "NA")
    df["empirical_p_g1_num"] = pd.to_numeric(df.get("empirical_p_g1", pd.Series(dtype=float)), errors="coerce")
    df["is_sig_raw"] = df["empirical_p_g1_num"] < alpha

    exposures = sorted(set(tested_df["exposure"].unique()) | set(df["exposure"].dropna().unique()))

    per_exposure_tables = {}
    all_rows = []
    for exposure in exposures:
        t_sub = tested_df.loc[tested_df["exposure"] == exposure, ["chromosome", "n_tested"]]
        if t_sub.empty:
            # No tested-variant baseline for this exposure/generation -> can't
            # compute an "expected" count, so skip rather than guess.
            continue

        df_sub = df.loc[df["exposure"] == exposure]
        sig_sub = df_sub.loc[df_sub["is_sig_raw"]]
        obs = sig_sub.groupby("chromosome")["variant"].nunique().rename("n_significant_observed").reset_index()
        sig_lists = sig_sub.groupby("chromosome")["variant"].apply(list).rename("sig_variants").reset_index()

        merged = t_sub.merge(obs, on="chromosome", how="left").merge(sig_lists, on="chromosome", how="left")
        merged["n_significant_observed"] = merged["n_significant_observed"].fillna(0).astype(int)
        merged["sig_variants"] = merged["sig_variants"].apply(lambda v: v if isinstance(v, list) else [])
        merged["n_tested"] = merged["n_tested"].astype(int)

        merged = _add_binomial_enrichment_stats(merged)
        merged = merged.sort_values(by="chromosome", key=lambda s: s.map(chrom_sort_key))
        merged["exposure"] = exposure

        per_exposure_tables[exposure] = merged
        all_rows.append(merged)

    if not all_rows:
        print("[warn] no exposure had tested-variant rows -- skipping per-exposure chromosome enrichment.", file=sys.stderr)
        _empty_placeholder_figure(fig_path, "No exposures with tested-variant data")
        empty = pd.DataFrame(columns=["exposure", "chromosome", "n_tested", "n_significant_observed",
                                       "expected", "binom_p", "binom_p_adj", "sig_variants"])
        empty.to_csv(stats_path, index=False)
        return empty, {"n_exposures_plotted": 0, "per_exposure_csv": str(stats_path), "figure": str(fig_path)}

    combined = pd.concat(all_rows, ignore_index=True)
    combined.to_csv(stats_path, index=False)

    # One standalone figure AND one standalone p-value CSV per exposure, in
    # addition to the combined grid/CSV below. This is the per-exposure,
    # per-chromosome p-value CSV requested: chromosome, n_tested,
    # n_significant_observed, expected (n_tested * exposure's own significance
    # rate), binom_p, binom_p_adj.
    by_exposure_dir = out_dir / "figures" / "by_exposure"
    by_exposure_dir.mkdir(parents=True, exist_ok=True)
    individual_paths = []
    individual_csv_paths = []
    for exposure, merged in per_exposure_tables.items():
        label = translate_exposure_value(exposure)
        slug = slugify(label)

        indiv_path = by_exposure_dir / f"observed_vs_expected_chrom_{slug}.png"
        plt.figure(figsize=(max(6, 0.6 * len(merged)), 4.5))
        _draw_chrom_bars(plt.gca(), merged, title=str(label))
        plt.tight_layout()
        plt.savefig(indiv_path, dpi=200)
        plt.close()
        individual_paths.append(indiv_path)

        csv_cols = ["chromosome", "n_tested", "n_significant_observed", "expected", "binom_p", "binom_p_adj"]
        indiv_csv_path = by_exposure_dir / f"pvalues_{slug}.csv"
        merged[csv_cols].to_csv(indiv_csv_path, index=False)
        individual_csv_paths.append(indiv_csv_path)

    n = len(per_exposure_tables)
    ncols = min(3, n)
    nrows = -(-n // ncols)  # ceil
    fig, axes = plt.subplots(nrows, ncols, figsize=(6.2 * ncols, 4.6 * nrows), squeeze=False)
    for idx, (exposure, merged) in enumerate(per_exposure_tables.items()):
        r, c = divmod(idx, ncols)
        _draw_chrom_bars(axes[r][c], merged, title=str(translate_exposure_value(exposure)))
    for idx in range(n, nrows * ncols):
        r, c = divmod(idx, ncols)
        axes[r][c].axis("off")
    fig.suptitle("Observed vs expected significant variants per chromosome, by exposure "
                 "(expected = n_tested \u00d7 exposure's own significance rate)", fontsize=12)
    plt.tight_layout(rect=(0, 0, 1, 0.97))
    plt.savefig(fig_path, dpi=200)
    plt.close(fig)

    summary = {
        "n_exposures_plotted": n,
        "per_exposure_csv": str(stats_path),
        "figure": str(fig_path),
        "individual_figures": [str(p) for p in individual_paths],
        "individual_pvalue_csvs": [str(p) for p in individual_csv_paths],
        "per_exposure_tables": per_exposure_tables,
    }
    return combined, summary


# ---------------------------------------------------------------------------
# Main (callable, reusable from the report runner and the CLI)
# ---------------------------------------------------------------------------

def run_table2(generation: int = 2, alpha: float = SIG_ALPHA_DEFAULT, out_dir: Path = OUT_DIR) -> None:
    from gene_environment.db.connection import cursor_scope, get_connection

    out_dir = Path(out_dir)
    fig_dir = out_dir / "figures"
    out_dir.mkdir(parents=True, exist_ok=True)
    fig_dir.mkdir(parents=True, exist_ok=True)

    print(f"Calling stored routine: {ASTORE_NAME}() ...")
    df = call_stored_routine_to_df(ASTORE_NAME, get_connection, cursor_scope, columns=TABLE_COLUMNS)
    df = drop_gna_columns(df)
    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    make_figures(df, fig_dir)

    # Raw (source-language) exposure values, needed to query the DB below.
    raw_exposures = sorted(df["exposure"].dropna().unique().tolist())
    tested_df = fetch_tested_variant_counts_for_exposures(get_connection, cursor_scope, raw_exposures, generation)

    chrom_stats_df, chrom_summary = compute_chromosome_enrichment_global(
        df, tested_df, alpha=alpha, out_dir=out_dir
    )
    print("Chromosome enrichment summary (all exposures):", chrom_summary)

    chrom_by_exposure_df, chrom_by_exposure_summary = compute_chromosome_enrichment_by_exposure(
        df, tested_df, alpha=alpha, out_dir=out_dir
    )
    print("Chromosome enrichment summary (by exposure):",
          {k: v for k, v in chrom_by_exposure_summary.items() if k != "per_exposure_tables"})

    # Every DB-dependent computation above needed the raw exposure values;
    # from here on only display (tables/captions) is left, so translate now.
    df = translate_exposure(df)

    # --- Table 2 (top 10) ---
    doc_top10 = Document()
    doc_top10.add_heading("Table 2. Significant variants (top 10)", level=1)
    doc_top10.add_paragraph(
        "Table shows the top 10 significant variants from get_significant_results_table_2. "
        "p-values highlighted in red are below the significance threshold "
        f"(alpha = {alpha}). gna.* columns omitted."
    )
    add_table_to_doc(doc_top10, df, max_rows=10, alpha=alpha)
    top10_path = out_dir / "Table2_top10.docx"
    doc_top10.save(top10_path)

    # --- Supplementary Word (full results + all figures) ---
    doc_full = Document()
    doc_full.add_heading("Supplementary Table: full results", level=1)
    doc_full.add_paragraph(
        "Full results from get_significant_results_table_2. All rows included. "
        "gna.* columns omitted from the table but available in the database."
    )
    add_table_to_doc(doc_full, df, alpha=alpha)

    doc_full.add_heading("Figures", level=1)
    add_figure_to_doc(doc_full, fig_dir / "variants_per_chromosome.png",
                       "Figure 1. Number of unique variants per chromosome.")
    add_figure_to_doc(doc_full, fig_dir / "genes_vs_variants_scatter.png",
                       "Figure 2. Genes vs variants per chromosome.")
    add_figure_to_doc(doc_full, fig_dir / "empirical_p_g1_histogram.png",
                       "Figure 3. Distribution of empirical_p_g1.")
    add_figure_to_doc(doc_full, fig_dir / "observed_vs_expected_by_chromosome.png",
                       "Figure 4. Observed vs expected significant variants per chromosome, all exposures pooled "
                       f"(generation {generation}). Expected = n_tested \u00d7 (total significant / total "
                       "tested, pooled across all exposures); p-values from a one-sided binomial test, "
                       "BH-adjusted across chromosomes.")
    add_figure_to_doc(doc_full, fig_dir / "observed_vs_expected_by_chromosome_per_exposure.png",
                       "Figure 5. Observed vs expected significant variants per chromosome, one panel per exposure "
                       f"(generation {generation}). Same proportional null model as Figure 4, but computed "
                       "independently for each exposure (using that exposure's own significance rate).",
                       width_in=6.5)

    # --- Per-exposure enrichment tables (chromosome-by-chromosome p-values) ---
    per_exposure_tables = chrom_by_exposure_summary.get("per_exposure_tables", {})
    if per_exposure_tables:
        doc_full.add_heading("Per-exposure chromosome enrichment", level=1)
        doc_full.add_paragraph(
            f"For each exposure and chromosome: n_tested variants (COUNT(*) FROM variant_results "
            f"WHERE exposure=... AND generation={generation} GROUP BY chromosome), the number "
            f"found significant (empirical_p_g1 < {alpha}), the expected count under a "
            "proportional-allocation null (n_tested \u00d7 the exposure's own overall significance "
            "rate, i.e. that exposure's total significant / total tested), and a one-sided binomial "
            f"p-value testing chromosome-level concentration of hits (BH-adjusted within each "
            f"exposure). Adjusted p-values below {alpha} are highlighted."
        )
        for exposure, merged in per_exposure_tables.items():
            label = translate_exposure_value(exposure)
            add_pvalue_table_to_doc(
                doc_full, merged,
                title=f"Exposure: {label}",
                caption=f"Chromosome-level enrichment for {label} (generation {generation}).",
                alpha=alpha,
            )
            slug = slugify(label)
            add_figure_to_doc(
                doc_full, fig_dir / "by_exposure" / f"observed_vs_expected_chrom_{slug}.png",
                f"Observed vs expected significant variants per chromosome -- {label}.",
                width_in=5.5,
            )

    full_path = out_dir / "Table2_full_supplementary.docx"
    doc_full.save(full_path)

    print("Done.")
    print(f"Top 10 Word: {top10_path}")
    print(f"Full supplementary Word: {full_path}")
    print(f"Figures saved in: {fig_dir}")
    print(f"Chromosome enrichment CSV (pooled): {chrom_summary['per_chromosome_csv']}")
    print(f"Chromosome enrichment CSV (by exposure, combined): {chrom_by_exposure_summary['per_exposure_csv']}")
    print(f"Per-exposure figures + p-value CSVs ({chrom_by_exposure_summary['n_exposures_plotted']}): "
          f"{fig_dir / 'by_exposure'}")


def parse_args():
    p = argparse.ArgumentParser(description="Generate Table 2 report (Word + figures).")
    p.add_argument("--generation", type=int, default=2,
                    help="Generation number used to pull tested-variant counts per exposure.")
    p.add_argument("--alpha", type=float, default=SIG_ALPHA_DEFAULT,
                    help="Significance threshold for empirical_p_g1, and the null probability "
                         "used in the per-chromosome binomial enrichment test.")
    return p.parse_args()


def main():
    args = parse_args()
    run_table2(generation=args.generation, alpha=args.alpha)


if __name__ == "__main__":
    main()

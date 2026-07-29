"""
word_utils.py

Shared python-docx helpers used by every report module (generate_table1,
generate_table2, generate_table2b, build_annotated_tables). Consolidated
here so cell shading, column widths, borders, landscape orientation and
figure embedding are done identically -- and fixed in one place -- across
all reports instead of being copy-pasted per script.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def set_cell_bg(cell, color_hex: str = "D9D9D9") -> None:
    """Apply flat shading to a table cell (never use w:val=SOLID, which
    Word renders as a solid fill of `color_hex` OVER a pattern -- 'clear'
    is the flat-fill value)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_col_width(cell, width_inches: float) -> None:
    """Set a table cell's width. Word column widths (w:type='dxa') are
    expressed in twips (1/1440 inch) -- NOT EMU (1/914400 inch)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def set_table_borders(table, color_hex: str = "BFBFBF", size: int = 4) -> None:
    """Apply thin, consistent borders to the whole table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tblPr.append(borders)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Set a table cell's text with optional bold, replacing any existing
    content. Lighter-weight than the shaded/bordered tables built by
    add_table_to_doc-style helpers -- for reports that just need a plain
    "Light Grid Accent 1"-style table with bold highlighting."""
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def repeat_header_row(table) -> None:
    """Mark a table's first row so it repeats on every page it spans."""
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def set_landscape(doc: Document, left: float = 0.5, right: float = 0.5,
                   top: float = 0.5, bottom: float = 0.5) -> None:
    """Switch the document's first section to landscape and apply the
    given margins (in inches) -- useful for wide tables (many columns)."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)


def add_figure_to_doc(doc: Document, fig_path: Path, caption: str, width_in: float = 6.0) -> None:
    """Embed a figure (if it exists) into the document with an italic
    caption centered below it. Silently skipped if the figure is missing,
    so a report can still be built if one plot failed upstream."""
    if not fig_path.exists():
        return
    doc.add_picture(str(fig_path), width=Inches(width_in))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)

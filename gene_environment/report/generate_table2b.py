# -*- coding: utf-8 -*-
"""
Generate Table 2b Word files and figures by calling the stored routine
get_significant_results_table_2b via the project's MySQL connection helpers.

Usage (from project root):
    python3 -m gene_environment.report.generate_table2b

Requirements:
    pip install pandas python-docx matplotlib seaborn mysql-connector-python

Behavior:
- Calls the stored routine `get_significant_results_table_2b()` and uses the
  first resultset returned. Columns expected (per the astore signature):
      variant, exposure, gene_id, gene_symbol, gene_type,
      expressed_brain, expressed_neurons, expressed_glia,
      ctd_chemicals, als_opentargets_score, neuro_plausibility_score
- This dataset has NO p-value / coefficient columns (unlike Table 2), so
  this script is purely descriptive: no significance highlighting, no
  statistical enrichment test. It reports gene annotation content instead.
- Translates the `exposure` column from the source dataset's Italian
  land-use terms to English (see `gene_environment.report.exposure_labels`)
  right after fetching, since (unlike Table 2) nothing downstream needs the
  raw value to query the DB again.
- Produces POOLED (all-exposures) outputs exactly as before:
    output/table2b/Table2b_top10.docx
    output/table2b/Table2b_full_supplementary.docx
    output/table2b/table2b_raw_results.csv
    output/table2b/figures/variants_per_chromosome.png
    output/table2b/figures/genes_vs_variants_scatter.png
    output/table2b/figures/gene_type_distribution.png
    output/table2b/figures/expressed_brain_neurons_glia.png
    output/table2b/figures/als_opentargets_score_histogram.png
    output/table2b/figures/neuro_plausibility_score_histogram.png
    output/table2b/figures/ctd_chemicals_top20.png
    output/table2b/table2b_gene_type_counts.csv
    output/table2b/table2b_chemicals_frequency.csv
- NEW: Produces PER-EXPOSURE outputs, mirroring the by_exposure pattern used
  in generate_table2.py:
    output/table2b/figures/by_exposure/<slug>/variants_per_chromosome.png
    output/table2b/figures/by_exposure/<slug>/gene_type_distribution.png
    output/table2b/figures/by_exposure/<slug>/expressed_brain_neurons_glia.png
    output/table2b/figures/by_exposure/<slug>/als_opentargets_score_histogram.png
    output/table2b/figures/by_exposure/<slug>/neuro_plausibility_score_histogram.png
    output/table2b/figures/by_exposure/<slug>/ctd_chemicals_top20.png
    output/table2b/figures/by_exposure/<slug>/gene_type_counts.csv
    output/table2b/figures/by_exposure/<slug>/chemicals_frequency.csv
  plus one combined grid figure per metric (one panel per exposure), same
  idea as `observed_vs_expected_by_chromosome_per_exposure.png` in Table 2:
    output/table2b/figures/gene_type_distribution_by_exposure.png
    output/table2b/figures/expressed_brain_neurons_glia_by_exposure.png
  <slug> uses the same `slugify(translate_exposure_value(...))` convention
  as Table 2, and is computed from the already-translated (English)
  exposure label, since (unlike Table 2) this script has no further DB
  queries downstream that need the raw Italian value.
  The full supplementary Word doc gets a new "Per-exposure figures" section
  with one subsection (heading + all 6 figures) per exposure.

Notes on ambiguous column types (documented here since the astore only
gives column names, not types):
- expressed_brain / expressed_neurons / expressed_glia: could be a 0/1 flag
  or a continuous score. This script inspects the actual observed values at
  runtime -- if every non-null value across the three columns is in {0, 1},
  it treats them as binary flags and plots proportion-of-genes-expressed;
  otherwise it treats them as continuous scores and plots histograms. A
  message is printed to stderr saying which path was taken. This detection
  is done ONCE on the pooled dataset, and the same mode is reused for every
  per-exposure panel so that all panels are visually comparable (a mode
  that flips exposure-to-exposure would make the by-exposure grid useless).
- ctd_chemicals: assumed to be a delimited list of chemical names per row
  (comma / semicolon / pipe separated). Frequency is counted after
  splitting on any of those delimiters and stripping whitespace.
"""

from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from gene_environment.report.db_utils import chrom_sort_key, call_stored_routine_to_df, extract_chromosome, slugify
from gene_environment.report.exposure_labels import translate_exposure
from gene_environment.report.word_utils import (
    add_figure_to_doc,
    repeat_header_row,
    set_cell_bg,
    set_col_width,
    set_landscape,
    set_table_borders,
)

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OUT_DIR = Path("output/table2b")
FIG_DIR = OUT_DIR / "figures"
BY_EXPOSURE_DIR = FIG_DIR / "by_exposure"

ASTORE_NAME = "get_significant_results_table_2b"

TABLE_COLUMNS = [
    "exposure",
    "variant",
    "gene_id",
    "gene_symbol",
    "gene_type",
    "expressed_brain",
    "expressed_neurons",
    "expressed_glia",
    "ctd_chemicals",
    "als_opentargets_score",
    "neuro_plausibility_score",
]

COLUMN_LABELS = {
    "exposure": "Exposure",
    "variant": "Variant",
    "gene_id": "Gene ID",
    "gene_symbol": "Gene symbol",
    "gene_type": "Gene type",
    "expressed_brain": "Expr. brain",
    "expressed_neurons": "Expr. neurons",
    "expressed_glia": "Expr. glia",
    "ctd_chemicals": "CTD chemicals",
    "als_opentargets_score": "ALS OpenTargets score",
    "neuro_plausibility_score": "Neuro plausibility score",
}

# Column widths in inches. Landscape US Letter with 0.5" margins gives ~10"
# of usable width -- these sum to 10.0.
COL_WIDTHS_IN = {
    "exposure": 1.1,
    "variant": 1.3,
    "gene_id": 0.7,
    "gene_symbol": 0.8,
    "gene_type": 0.8,
    "expressed_brain": 0.7,
    "expressed_neurons": 0.7,
    "expressed_glia": 0.7,
    "ctd_chemicals": 1.5,
    "als_opentargets_score": 0.8,
    "neuro_plausibility_score": 0.9,
}

CTD_CHEMICALS_TRUNCATE_CHARS = 80
CTD_CHEMICALS_TOP_N = 20
CHEM_SPLIT_RE = re.compile(r"[,;|]")

# Table color scheme (same convention as Table 2)
HEADER_FILL = "44546A"
HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_FILL = "EEF1F6"

_TRUE_TOKENS = {"1", "1.0", "true", "t", "yes", "y"}
_FALSE_TOKENS = {"0", "0.0", "false", "f", "no", "n"}


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

def _format_bool_like(val) -> str:
    """Best-effort formatting of a value that might be a 0/1 flag, a
    True/False, or a Y/N string, as 'Yes' / 'No'. Falls back to the raw
    string if it doesn't look boolean (e.g. a continuous score)."""
    s = str(val).strip().lower()
    if s in _TRUE_TOKENS:
        return "Yes"
    if s in _FALSE_TOKENS:
        return "No"
    try:
        return "{:.3g}".format(float(val))
    except (TypeError, ValueError):
        return str(val)


def format_value_for_word(col: str, val) -> str:
    if pd.isna(val):
        return ""
    if col == "ctd_chemicals":
        text = str(val)
        if len(text) > CTD_CHEMICALS_TRUNCATE_CHARS:
            return text[:CTD_CHEMICALS_TRUNCATE_CHARS].rstrip() + "\u2026"
        return text
    if col in ("expressed_brain", "expressed_neurons", "expressed_glia"):
        return _format_bool_like(val)
    if col in ("als_opentargets_score", "neuro_plausibility_score"):
        try:
            return "{:.3g}".format(float(val))
        except (TypeError, ValueError):
            return str(val)
    return str(val)


def _split_chemicals(raw) -> List[str]:
    if pd.isna(raw):
        return []
    parts = CHEM_SPLIT_RE.split(str(raw))
    return [p.strip() for p in parts if p.strip()]


# ---------------------------------------------------------------------------
# docx table helpers
# ---------------------------------------------------------------------------

def add_table_to_doc(
    doc: Document,
    df: pd.DataFrame,
    title: Optional[str] = None,
    caption: Optional[str] = None,
    max_rows: Optional[int] = None,
) -> None:
    if title:
        h = doc.add_heading(title, level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    df_to_use = df[TABLE_COLUMNS]
    if max_rows is not None:
        df_to_use = df_to_use.head(max_rows)

    table = doc.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(TABLE_COLUMNS):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(COLUMN_LABELS.get(col, col))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = HEADER_FONT_COLOR
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], HEADER_FILL)
        set_col_width(hdr_cells[i], COL_WIDTHS_IN.get(col, 1.0))

    numeric_cols = ("als_opentargets_score", "neuro_plausibility_score")
    for ridx, (_, row) in enumerate(df_to_use.iterrows()):
        cells = table.add_row().cells
        shade = ZEBRA_FILL if ridx % 2 == 1 else "FFFFFF"
        for i, col in enumerate(TABLE_COLUMNS):
            text = format_value_for_word(col, row.get(col, pd.NA))
            cell = cells[i]
            set_col_width(cell, COL_WIDTHS_IN.get(col, 1.0))
            para = cell.paragraphs[0]
            para.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT if col in numeric_cols else WD_PARAGRAPH_ALIGNMENT.LEFT
            )
            run = para.add_run(text)
            run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, shade)

    repeat_header_row(table)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


# ---------------------------------------------------------------------------
# Figures -- all of these now take an explicit `fig_dir` / `out_dir` so they
# can be reused unchanged for both the pooled (all-exposures) run and each
# per-exposure run: pooled calls pass FIG_DIR/OUT_DIR, per-exposure calls
# pass BY_EXPOSURE_DIR/<slug>/ for both (figures and their companion CSVs
# live together in that per-exposure folder).
# ---------------------------------------------------------------------------

def make_chromosome_figures(df: pd.DataFrame, fig_dir: Path, title_suffix: str = "") -> None:
    """Purely descriptive chromosome-level figures (no significance concept
    here -- every row in this dataset is already the astore's output)."""
    df = df.copy()
    df["chromosome"] = df["variant"].apply(lambda v: extract_chromosome(v) if pd.notna(v) else "NA")

    variants_per_chrom = df.groupby("chromosome")["variant"].nunique().rename("n_variants").reset_index()
    genes_per_chrom = df.groupby("chromosome")["gene_symbol"].nunique().rename("n_genes").reset_index()
    merged = pd.merge(variants_per_chrom, genes_per_chrom, on="chromosome", how="outer").fillna(0)
    merged = merged.sort_values(by="chromosome", key=lambda s: s.map(chrom_sort_key))

    sns.set_theme(style="whitegrid")

    plt.figure(figsize=(10, 6))
    ax = sns.barplot(data=merged, x="chromosome", y="n_variants", color="#4472C4")
    ax.set_xlabel("Chromosome")
    ax.set_ylabel("Number of unique variants")
    ax.set_title(f"Variants per chromosome (Table 2b){title_suffix}")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(fig_dir / "variants_per_chromosome.png", dpi=300)
    plt.close()

    plt.figure(figsize=(8, 6))
    ax = sns.scatterplot(data=merged, x="n_genes", y="n_variants", s=100, color="#4472C4")
    for _, r in merged.iterrows():
        ax.text(r["n_genes"], r["n_variants"], str(r["chromosome"]), fontsize=9,
                 horizontalalignment="left", verticalalignment="bottom")
    ax.set_xlabel("Number of unique genes per chromosome")
    ax.set_ylabel("Number of unique variants per chromosome")
    ax.set_title(f"Genes vs variants per chromosome (Table 2b){title_suffix}")
    plt.tight_layout()
    plt.savefig(fig_dir / "genes_vs_variants_scatter.png", dpi=300)
    plt.close()


def make_gene_type_figure(df: pd.DataFrame, out_dir: Path, fig_dir: Path, title_suffix: str = "") -> pd.DataFrame:
    """Bar chart of unique genes per gene_type. Returns the counts table,
    also saved to CSV."""
    genes = df.drop_duplicates(subset=["gene_symbol"])
    counts = genes["gene_type"].fillna("Unknown").value_counts().rename_axis("gene_type").reset_index(name="n_genes")
    counts = counts.sort_values("n_genes", ascending=False)
    counts.to_csv(out_dir / "gene_type_counts.csv" if out_dir != OUT_DIR else out_dir / "table2b_gene_type_counts.csv",
                  index=False)

    plt.figure(figsize=(max(6, 0.5 * len(counts)), 5))
    ax = sns.barplot(data=counts, x="gene_type", y="n_genes", color="#4472C4")
    ax.set_xlabel("Gene type")
    ax.set_ylabel("Number of unique genes")
    ax.set_title(f"Gene type distribution{title_suffix}")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(fig_dir / "gene_type_distribution.png", dpi=300)
    plt.close()
    return counts


def make_expression_figure(df: pd.DataFrame, fig_dir: Path, mode: Optional[str] = None,
                            title_suffix: str = "") -> str:
    """expressed_brain / expressed_neurons / expressed_glia: if `mode` is
    given ('binary' or 'continuous') it is used as-is -- this lets
    per-exposure calls reuse the mode detected once on the pooled dataset,
    so every panel in the by-exposure grid is visually comparable. If
    `mode` is None (pooled call), auto-detects whether the observed values
    look like binary flags (all in {0,1}) or continuous scores. Returns
    the mode used ('binary' or 'continuous') for logging / reuse."""
    cols = ["expressed_brain", "expressed_neurons", "expressed_glia"]
    genes = df.drop_duplicates(subset=["gene_symbol"])[["gene_symbol"] + cols].copy()

    numeric = {}
    for c in cols:
        numeric[c] = pd.to_numeric(genes[c], errors="coerce")
    numeric_df = pd.DataFrame(numeric)

    if mode is None:
        observed_values = set(numeric_df.stack().dropna().unique().tolist())
        is_binary = observed_values.issubset({0.0, 1.0}) and len(observed_values) > 0
        mode = "binary" if is_binary else "continuous"
        print(f"[info] expressed_brain/neurons/glia detected as {mode} "
              "(this mode will be reused for all per-exposure panels).", file=sys.stderr)

    fig_path = fig_dir / "expressed_brain_neurons_glia.png"

    if mode == "binary":
        props = numeric_df.mean(numeric_only=True).reset_index()
        props.columns = ["compartment", "proportion_expressed"]
        props["compartment"] = props["compartment"].str.replace("expressed_", "", regex=False)

        plt.figure(figsize=(6, 5))
        ax = sns.barplot(data=props, x="compartment", y="proportion_expressed", color="#548235")
        ax.set_ylim(0, 1)
        ax.set_xlabel("Compartment")
        ax.set_ylabel("Proportion of unique genes expressed")
        ax.set_title(f"Gene expression by compartment (binary flags){title_suffix}")
        plt.tight_layout()
        plt.savefig(fig_path, dpi=300)
        plt.close()
    else:
        fig, axes = plt.subplots(1, 3, figsize=(15, 4.5))
        for ax, c in zip(axes, cols):
            vals = numeric_df[c].dropna()
            if vals.empty:
                ax.text(0.5, 0.5, "No data", ha="center", va="center", transform=ax.transAxes)
                ax.axis("off")
                continue
            sns.histplot(vals, bins=30, ax=ax, color="#548235")
            ax.set_title(c.replace("expressed_", "").capitalize())
            ax.set_xlabel("Score")
        fig.suptitle(title_suffix.strip(" ()") or None)
        plt.tight_layout()
        plt.savefig(fig_path, dpi=300)
        plt.close(fig)

    return mode


def make_score_histogram(df: pd.DataFrame, col: str, fig_dir: Path, fig_name: str, title: str) -> None:
    """Histogram of a numeric score column, deduplicated by gene_symbol so a
    gene tested against multiple exposures/variants isn't overweighted."""
    genes = df.drop_duplicates(subset=["gene_symbol"])
    vals = pd.to_numeric(genes.get(col, pd.Series(dtype=float)), errors="coerce").dropna()

    fig_path = fig_dir / fig_name
    plt.figure(figsize=(8, 5))
    if not vals.empty:
        ax = sns.histplot(vals, bins=40, color="#ED7D31")
        ax.set_xlabel(col)
        ax.set_ylabel("Number of unique genes")
        ax.set_title(title)
        plt.tight_layout()
    else:
        plt.text(0.5, 0.5, f"No {col} data available", ha="center", va="center")
        plt.axis("off")
    plt.savefig(fig_path, dpi=300)
    plt.close()


def make_chemicals_figure(df: pd.DataFrame, out_dir: Path, fig_dir: Path, title_suffix: str = "") -> pd.DataFrame:
    """Top-N most frequent CTD chemicals across all rows. Frequency counts
    rows (variant x exposure x gene), not unique genes, since the same
    chemical linked to different genes/variants is still relevant signal.
    Returns the frequency table, also saved to CSV."""
    counter: Counter = Counter()
    for raw in df.get("ctd_chemicals", pd.Series(dtype=object)):
        counter.update(_split_chemicals(raw))

    fig_path = fig_dir / "ctd_chemicals_top20.png"
    freq_path = (out_dir / "chemicals_frequency.csv" if out_dir != OUT_DIR
                 else out_dir / "table2b_chemicals_frequency.csv")

    if not counter:
        freq = pd.DataFrame(columns=["chemical", "n_occurrences"])
        freq.to_csv(freq_path, index=False)
        plt.figure(figsize=(8, 5))
        plt.text(0.5, 0.5, "No ctd_chemicals data available", ha="center", va="center")
        plt.axis("off")
        plt.savefig(fig_path, dpi=300)
        plt.close()
        return freq

    freq = pd.DataFrame(counter.most_common(), columns=["chemical", "n_occurrences"])
    freq.to_csv(freq_path, index=False)

    top = freq.head(CTD_CHEMICALS_TOP_N).iloc[::-1]  # reverse for horizontal bar (largest on top)
    plt.figure(figsize=(8, max(5, 0.3 * len(top))))
    ax = sns.barplot(data=top, y="chemical", x="n_occurrences", color="#4472C4", orient="h")
    ax.set_xlabel("Occurrences")
    ax.set_ylabel("Chemical")
    ax.set_title(f"Top {min(CTD_CHEMICALS_TOP_N, len(freq))} CTD chemicals{title_suffix}")
    plt.tight_layout()
    plt.savefig(fig_path, dpi=300)
    plt.close()
    return freq


# ---------------------------------------------------------------------------
# Per-exposure orchestration (new)
# ---------------------------------------------------------------------------

def run_per_exposure_figures(
    df: pd.DataFrame,
    exposures: List[str],
    expression_mode: str,
    by_exposure_dir: Path,
    fig_dir: Path,
) -> Dict[str, Path]:
    """For each (already-translated/English) exposure label, filter df and
    regenerate the full set of descriptive figures + companion CSVs into
    their own subfolder `by_exposure/<slug>/`, then build one combined grid
    figure per metric (gene-type distribution, expression) with one panel
    per exposure -- mirroring `observed_vs_expected_by_chromosome_per_exposure`
    in generate_table2.py. Returns {exposure: slug_dir} for use when
    building the Word doc.

    Exposures with zero rows (shouldn't normally happen since `exposures`
    is derived from df itself) are skipped defensively.
    """
    slug_dirs: Dict[str, Path] = {}
    gene_type_panels: List[Tuple[str, pd.DataFrame]] = []
    expression_panels: List[str] = []  # just track exposures with data, for the grid

    for exposure in exposures:
        df_sub = df.loc[df["exposure"] == exposure]
        if df_sub.empty:
            continue

        slug = slugify(exposure)
        slug_dir = by_exposure_dir / slug
        slug_dir.mkdir(parents=True, exist_ok=True)
        slug_dirs[exposure] = slug_dir

        title_suffix = f" -- {exposure}"

        make_chromosome_figures(df_sub, slug_dir, title_suffix=title_suffix)
        counts = make_gene_type_figure(df_sub, slug_dir, slug_dir, title_suffix=title_suffix)
        make_expression_figure(df_sub, slug_dir, mode=expression_mode, title_suffix=title_suffix)
        make_score_histogram(df_sub, "als_opentargets_score", slug_dir,
                              "als_opentargets_score_histogram.png",
                              f"Distribution of ALS OpenTargets score{title_suffix}")
        make_score_histogram(df_sub, "neuro_plausibility_score", slug_dir,
                              "neuro_plausibility_score_histogram.png",
                              f"Distribution of neuro plausibility score{title_suffix}")
        make_chemicals_figure(df_sub, slug_dir, slug_dir, title_suffix=title_suffix)

        counts["exposure"] = exposure
        gene_type_panels.append((exposure, counts))
        expression_panels.append(exposure)

    # --- combined grid: gene type distribution by exposure ---
    if gene_type_panels:
        n = len(gene_type_panels)
        ncols = min(3, n)
        nrows = -(-n // ncols)
        fig, axes = plt.subplots(nrows, ncols, figsize=(6.2 * ncols, 4.6 * nrows), squeeze=False)
        for idx, (exposure, counts) in enumerate(gene_type_panels):
            r, c = divmod(idx, ncols)
            ax = axes[r][c]
            sns.barplot(data=counts, x="gene_type", y="n_genes", color="#4472C4", ax=ax)
            ax.set_title(exposure, fontsize=10)
            ax.set_xlabel("")
            ax.set_ylabel("N genes", fontsize=9)
            ax.tick_params(axis="x", rotation=45, labelsize=8)
        for idx in range(n, nrows * ncols):
            r, c = divmod(idx, ncols)
            axes[r][c].axis("off")
        fig.suptitle("Gene type distribution by exposure", fontsize=12)
        plt.tight_layout(rect=(0, 0, 1, 0.97))
        plt.savefig(fig_dir / "gene_type_distribution_by_exposure.png", dpi=200)
        plt.close(fig)

    # --- combined grid: expression by exposure (reuses the per-exposure PNGs) ---
    if expression_panels:
        n = len(expression_panels)
        ncols = min(3, n)
        nrows = -(-n // ncols)
        fig, axes = plt.subplots(nrows, ncols, figsize=(6.0 * ncols, 4.2 * nrows), squeeze=False)
        for idx, exposure in enumerate(expression_panels):
            r, c = divmod(idx, ncols)
            ax = axes[r][c]
            img = plt.imread(slug_dirs[exposure] / "expressed_brain_neurons_glia.png")
            ax.imshow(img)
            ax.set_title(exposure, fontsize=10)
            ax.axis("off")
        for idx in range(n, nrows * ncols):
            r, c = divmod(idx, ncols)
            axes[r][c].axis("off")
        fig.suptitle(f"Expression by compartment, by exposure (mode: {expression_mode})", fontsize=12)
        plt.tight_layout(rect=(0, 0, 1, 0.96))
        plt.savefig(fig_dir / "expressed_brain_neurons_glia_by_exposure.png", dpi=150)
        plt.close(fig)

    return slug_dirs


# ---------------------------------------------------------------------------
# Main (callable, reusable from the report runner and the CLI)
# ---------------------------------------------------------------------------

def run_table2b(out_dir: Path = OUT_DIR) -> None:
    from gene_environment.db.connection import cursor_scope, get_connection

    out_dir = Path(out_dir)
    fig_dir = out_dir / "figures"
    by_exposure_dir = fig_dir / "by_exposure"
    out_dir.mkdir(parents=True, exist_ok=True)
    fig_dir.mkdir(parents=True, exist_ok=True)
    by_exposure_dir.mkdir(parents=True, exist_ok=True)

    print(f"Calling stored routine: {ASTORE_NAME}() ...")
    df = call_stored_routine_to_df(ASTORE_NAME, get_connection, cursor_scope, columns=TABLE_COLUMNS)
    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    # No further DB round-trip depends on the raw exposure value here (this
    # script, unlike Table 2, has no per-exposure DB queries), so translate
    # immediately -- everything downstream, including the by-exposure slugs,
    # uses the English label.
    df = translate_exposure(df)

    df.to_csv(out_dir / "table2b_raw_results.csv", index=False)

    if df.empty:
        print("[warn] no rows returned by the stored routine -- nothing further to do.", file=sys.stderr)
        return

    # --- pooled (all-exposures) figures, unchanged from before ---
    make_chromosome_figures(df, fig_dir)
    gene_type_counts = make_gene_type_figure(df, out_dir, fig_dir)
    expression_mode = make_expression_figure(df, fig_dir)
    make_score_histogram(df, "als_opentargets_score", fig_dir, "als_opentargets_score_histogram.png",
                          "Distribution of ALS OpenTargets score (unique genes)")
    make_score_histogram(df, "neuro_plausibility_score", fig_dir, "neuro_plausibility_score_histogram.png",
                          "Distribution of neuro plausibility score (unique genes)")
    chemicals_freq = make_chemicals_figure(df, out_dir, fig_dir)

    # --- NEW: per-exposure figures ---
    exposures = sorted(df["exposure"].dropna().unique().tolist())
    slug_dirs = run_per_exposure_figures(df, exposures, expression_mode, by_exposure_dir, fig_dir)
    print(f"Per-exposure figures generated for {len(slug_dirs)} exposures in: {by_exposure_dir}")

    n_unique_genes = df["gene_symbol"].nunique()
    n_unique_variants = df["variant"].nunique()
    n_unique_exposures = df["exposure"].nunique()

    # --- Table 2b (top 10) ---
    doc_top10 = Document()
    set_landscape(doc_top10, top=0.6, bottom=0.6)
    doc_top10.add_heading("Table 2b. Gene annotations for significant variants (top 10)", level=1)
    doc_top10.add_paragraph(
        "Table shows the top 10 rows from get_significant_results_table_2b. This dataset has no "
        "p-value / coefficient columns, so no significance highlighting is applied here -- it "
        "reports gene annotation content (gene type, tissue expression, CTD chemicals, "
        "OpenTargets / neuro-plausibility scores) for the variant-gene-exposure combinations "
        "returned by the routine."
    )
    add_table_to_doc(doc_top10, df, max_rows=10)
    top10_path = out_dir / "Table2b_top10.docx"
    doc_top10.save(top10_path)

    # --- Supplementary Word (full results + all figures) ---
    doc_full = Document()
    set_landscape(doc_full, top=0.6, bottom=0.6)
    doc_full.add_heading("Supplementary Table 2b: gene annotations, full results", level=1)
    doc_full.add_paragraph(
        f"Full results from get_significant_results_table_2b. {len(df)} rows covering "
        f"{n_unique_variants} unique variants, {n_unique_genes} unique genes, and "
        f"{n_unique_exposures} unique exposures. CTD chemical lists are truncated to "
        f"{CTD_CHEMICALS_TRUNCATE_CHARS} characters in this table; the full text is in "
        "table2b_raw_results.csv."
    )
    add_table_to_doc(doc_full, df)

    doc_full.add_heading("Figures (all exposures pooled)", level=1)
    add_figure_to_doc(doc_full, fig_dir / "variants_per_chromosome.png",
                       "Figure 1. Number of unique variants per chromosome.")
    add_figure_to_doc(doc_full, fig_dir / "genes_vs_variants_scatter.png",
                       "Figure 2. Genes vs variants per chromosome.")
    add_figure_to_doc(doc_full, fig_dir / "gene_type_distribution.png",
                       "Figure 3. Distribution of gene types among unique genes.")
    expr_caption = (
        "Figure 4. Proportion of unique genes expressed in brain / neurons / glia."
        if expression_mode == "binary" else
        "Figure 4. Distribution of brain / neuron / glia expression scores among unique genes "
        "(detected as continuous, not binary flags -- see script notes)."
    )
    add_figure_to_doc(doc_full, fig_dir / "expressed_brain_neurons_glia.png", expr_caption, width_in=6.5)
    add_figure_to_doc(doc_full, fig_dir / "als_opentargets_score_histogram.png",
                       "Figure 5. Distribution of ALS OpenTargets score among unique genes.")
    add_figure_to_doc(doc_full, fig_dir / "neuro_plausibility_score_histogram.png",
                       "Figure 6. Distribution of neuro plausibility score among unique genes.")
    add_figure_to_doc(doc_full, fig_dir / "ctd_chemicals_top20.png",
                       f"Figure 7. Top {min(CTD_CHEMICALS_TOP_N, len(chemicals_freq))} most frequent "
                       "CTD chemicals across all rows (row-level occurrence count).", width_in=6.5)

    # --- NEW: combined by-exposure grids, then one subsection per exposure ---
    if slug_dirs:
        doc_full.add_heading("Figures by exposure", level=1)
        add_figure_to_doc(doc_full, fig_dir / "gene_type_distribution_by_exposure.png",
                           "Figure 8. Gene type distribution, one panel per exposure.", width_in=6.5)
        add_figure_to_doc(doc_full, fig_dir / "expressed_brain_neurons_glia_by_exposure.png",
                           f"Figure 9. Expression by compartment, one panel per exposure (mode: {expression_mode}).",
                           width_in=6.5)

        doc_full.add_heading("Per-exposure figures", level=1)
        doc_full.add_paragraph(
            "The following subsections repeat the full descriptive figure set (Figures 1-7 above) "
            "separately for each exposure, using the same gene-expression detection mode "
            f"('{expression_mode}') as the pooled analysis for comparability."
        )
        for exposure, slug_dir in slug_dirs.items():
            doc_full.add_heading(f"Exposure: {exposure}", level=2)
            add_figure_to_doc(doc_full, slug_dir / "variants_per_chromosome.png",
                               f"Variants per chromosome -- {exposure}.")
            add_figure_to_doc(doc_full, slug_dir / "genes_vs_variants_scatter.png",
                               f"Genes vs variants per chromosome -- {exposure}.")
            add_figure_to_doc(doc_full, slug_dir / "gene_type_distribution.png",
                               f"Gene type distribution -- {exposure}.")
            add_figure_to_doc(doc_full, slug_dir / "expressed_brain_neurons_glia.png",
                               f"Expression by compartment -- {exposure}.", width_in=6.5)
            add_figure_to_doc(doc_full, slug_dir / "als_opentargets_score_histogram.png",
                               f"ALS OpenTargets score distribution -- {exposure}.")
            add_figure_to_doc(doc_full, slug_dir / "neuro_plausibility_score_histogram.png",
                               f"Neuro plausibility score distribution -- {exposure}.")
            add_figure_to_doc(doc_full, slug_dir / "ctd_chemicals_top20.png",
                               f"Top CTD chemicals -- {exposure}.", width_in=6.5)

    full_path = out_dir / "Table2b_full_supplementary.docx"
    doc_full.save(full_path)

    print("Done.")
    print(f"Top 10 Word: {top10_path}")
    print(f"Full supplementary Word: {full_path}")
    print(f"Raw results CSV: {out_dir / 'table2b_raw_results.csv'}")
    print(f"Gene type counts CSV: {out_dir / 'table2b_gene_type_counts.csv'}")
    print(f"CTD chemicals frequency CSV: {out_dir / 'table2b_chemicals_frequency.csv'}")
    print(f"Figures saved in: {fig_dir}")
    print(f"Per-exposure figures + CSVs saved in: {by_exposure_dir}")
    print(f"Expression columns detected as: {expression_mode}")


def main():
    run_table2b()


if __name__ == "__main__":
    main()
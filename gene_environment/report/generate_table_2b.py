# -*- coding: utf-8 -*-
"""
Generate Table 2b Word files and figures by calling the stored routine
get_significant_results_table_2b via the project's MySQL connection helpers.

Usage (from project root):
    python3 -m gene_environment.report.generate_table2b

Requirements:
    pip install pandas python-docx matplotlib seaborn mysql-connector-python

Behavior:
- Calls the stored routine `get_significant_results_table_2b()` and uses the
  first resultset returned. Columns expected (per the astore signature):
      variant, exposure, gene_id, gene_symbol, gene_type,
      expressed_brain, expressed_neurons, expressed_glia,
      ctd_chemicals, als_opentargets_score, neuro_plausibility_score
- This dataset has NO p-value / coefficient columns (unlike Table 2's
  get_significant_results_table_2), so this script is purely descriptive:
  no significance highlighting, no chromosome enrichment test. It reports
  gene annotation content instead.
- Produces:
    output/table2b/Table2b_top10.docx
    output/table2b/Table2b_full_supplementary.docx
    output/table2b/table2b_raw_results.csv
    output/table2b/figures/variants_per_chromosome.png
    output/table2b/figures/genes_vs_variants_scatter.png
    output/table2b/figures/gene_type_distribution.png
    output/table2b/figures/expressed_brain_neurons_glia.png
    output/table2b/figures/als_opentargets_score_histogram.png
    output/table2b/figures/neuro_plausibility_score_histogram.png
    output/table2b/figures/ctd_chemicals_top20.png
    output/table2b/table2b_gene_type_counts.csv
    output/table2b/table2b_chemicals_frequency.csv

Notes on ambiguous column types (documented here since the astore only
gives column names, not types):
- expressed_brain / expressed_neurons / expressed_glia: could be a 0/1 flag
  or a continuous score. This script inspects the actual observed values at
  runtime -- if every non-null value across the three columns is in {0, 1},
  it treats them as binary flags and plots proportion-of-genes-expressed;
  otherwise it treats them as continuous scores and plots histograms. A
  message is printed to stderr saying which path was taken.
- ctd_chemicals: assumed to be a delimited list of chemical names per row
  (comma / semicolon / pipe separated). Frequency is counted after
  splitting on any of those delimiters and stripping whitespace.
"""

from __future__ import annotations

import re
import sys
from collections import Counter
from pathlib import Path
from typing import List, Optional

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

OUT_DIR = Path("output/table2b")
FIG_DIR = OUT_DIR / "figures"
OUT_DIR.mkdir(parents=True, exist_ok=True)
FIG_DIR.mkdir(parents=True, exist_ok=True)

ASTORE_NAME = "get_significant_results_table_2b"

TABLE_COLUMNS = [
    "exposure",
    "variant",
    "gene_id",
    "gene_symbol",
    "gene_type",
    "expressed_brain",
    "expressed_neurons",
    "expressed_glia",
    "ctd_chemicals",
    "als_opentargets_score",
    "neuro_plausibility_score",
]

COLUMN_LABELS = {
    "exposure": "Exposure",
    "variant": "Variant",
    "gene_id": "Gene ID",
    "gene_symbol": "Gene symbol",
    "gene_type": "Gene type",
    "expressed_brain": "Expr. brain",
    "expressed_neurons": "Expr. neurons",
    "expressed_glia": "Expr. glia",
    "ctd_chemicals": "CTD chemicals",
    "als_opentargets_score": "ALS OpenTargets score",
    "neuro_plausibility_score": "Neuro plausibility score",
}

# Column widths in inches. Landscape US Letter with 0.5" margins gives ~10"
# of usable width -- these sum to 10.0.
COL_WIDTHS_IN = {
    "exposure": 1.1,
    "variant": 1.3,
    "gene_id": 0.7,
    "gene_symbol": 0.8,
    "gene_type": 0.8,
    "expressed_brain": 0.7,
    "expressed_neurons": 0.7,
    "expressed_glia": 0.7,
    "ctd_chemicals": 1.5,
    "als_opentargets_score": 0.8,
    "neuro_plausibility_score": 0.9,
}

CTD_CHEMICALS_TRUNCATE_CHARS = 80
CTD_CHEMICALS_TOP_N = 20
CHEM_SPLIT_RE = re.compile(r"[,;|]")

# Table color scheme (same convention as Table 2)
HEADER_FILL = "44546A"
HEADER_FONT_COLOR = RGBColor(0xFF, 0xFF, 0xFF)
ZEBRA_FILL = "EEF1F6"


# ---------------------------------------------------------------------------
# docx table helpers (same conventions as generate_table2.py)
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, color_hex: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_col_width(cell, width_inches: float) -> None:
    """Twips (1/1440 inch), not EMU -- see note in generate_table2.py."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_borders(table, color_hex: str = "BFBFBF", size: int = 4) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        borders.append(el)
    tblPr.append(borders)


def make_landscape(doc: Document) -> None:
    """Switch the current (first) section to landscape with narrow margins,
    needed because this table has 11 columns."""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)


def format_value_for_word(col: str, val) -> str:
    if pd.isna(val):
        return ""
    if col == "ctd_chemicals":
        text = str(val)
        if len(text) > CTD_CHEMICALS_TRUNCATE_CHARS:
            return text[:CTD_CHEMICALS_TRUNCATE_CHARS].rstrip() + "\u2026"
        return text
    if col in ("expressed_brain", "expressed_neurons", "expressed_glia"):
        return _format_bool_like(val)
    if col in ("als_opentargets_score", "neuro_plausibility_score"):
        try:
            return "{:.3g}".format(float(val))
        except (TypeError, ValueError):
            return str(val)
    return str(val)


_TRUE_TOKENS = {"1", "1.0", "true", "t", "yes", "y"}
_FALSE_TOKENS = {"0", "0.0", "false", "f", "no", "n"}


def _format_bool_like(val) -> str:
    """Best-effort formatting of a value that might be a 0/1 flag, a
    True/False, or a Y/N string, as 'Yes' / 'No'. Falls back to the raw
    string if it doesn't look boolean (e.g. a continuous score)."""
    s = str(val).strip().lower()
    if s in _TRUE_TOKENS:
        return "Yes"
    if s in _FALSE_TOKENS:
        return "No"
    try:
        return "{:.3g}".format(float(val))
    except (TypeError, ValueError):
        return str(val)


def add_table_to_doc(
    doc: Document,
    df: pd.DataFrame,
    title: Optional[str] = None,
    caption: Optional[str] = None,
    max_rows: Optional[int] = None,
) -> None:
    if title:
        h = doc.add_heading(title, level=2)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    df_to_use = df[TABLE_COLUMNS]
    if max_rows is not None:
        df_to_use = df_to_use.head(max_rows)

    table = doc.add_table(rows=1, cols=len(TABLE_COLUMNS))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(TABLE_COLUMNS):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(COLUMN_LABELS.get(col, col))
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = HEADER_FONT_COLOR
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_bg(hdr_cells[i], HEADER_FILL)
        _set_col_width(hdr_cells[i], COL_WIDTHS_IN.get(col, 1.0))

    numeric_cols = ("als_opentargets_score", "neuro_plausibility_score")
    for ridx, (_, row) in enumerate(df_to_use.iterrows()):
        cells = table.add_row().cells
        shade = ZEBRA_FILL if ridx % 2 == 1 else "FFFFFF"
        for i, col in enumerate(TABLE_COLUMNS):
            text = format_value_for_word(col, row.get(col, pd.NA))
            cell = cells[i]
            _set_col_width(cell, COL_WIDTHS_IN.get(col, 1.0))
            para = cell.paragraphs[0]
            para.alignment = (
                WD_PARAGRAPH_ALIGNMENT.RIGHT if col in numeric_cols else WD_PARAGRAPH_ALIGNMENT.LEFT
            )
            run = para.add_run(text)
            run.font.size = Pt(8)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_bg(cell, shade)

    first_tr = table.rows[0]._tr
    trPr = first_tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)

    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)


def add_figure_to_doc(doc: Document, fig_path: Path, caption: str, width_in: float = 6.0) -> None:
    if not fig_path.exists():
        return
    doc.add_picture(str(fig_path), width=Inches(width_in))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)


# ---------------------------------------------------------------------------
# Data access
# ---------------------------------------------------------------------------

def call_stored_routine_to_df(astore_name: str, get_connection, cursor_scope) -> pd.DataFrame:
    """Call `astore_name()` and return a DataFrame from the first resultset.
    Identical logic to generate_table2.py's version (handles drivers that
    need nextset() iteration)."""
    rows: List[dict] = []

    with get_connection() as conn:
        with cursor_scope(conn, dictionary=True) as cur:
            try:
                cur.execute(f"CALL {astore_name}()")
            except Exception as e:
                raise RuntimeError(f"Failed to CALL {astore_name}(). DB error: {e}") from e

            try:
                fetched = cur.fetchall()
                if fetched:
                    rows = fetched
            except Exception as e:
                print(f"[warn] initial fetchall() failed, will try nextset(): {e}", file=sys.stderr)
                rows = []

            try:
                while (not rows) and cur.nextset():
                    try:
                        fetched = cur.fetchall()
                        if fetched:
                            rows = fetched
                            break
                    except Exception as e:
                        print(f"[warn] fetchall() on a later resultset failed: {e}", file=sys.stderr)
                        continue
            except Exception as e:
                print(f"[warn] nextset() not supported by this driver: {e}", file=sys.stderr)

    if not rows:
        return pd.DataFrame(columns=TABLE_COLUMNS)

    df = pd.DataFrame(rows)
    df.columns = [c.strip() if isinstance(c, str) else c for c in df.columns]
    return df


# ---------------------------------------------------------------------------
# Formatting / parsing helpers
# ---------------------------------------------------------------------------

def extract_chromosome(variant: str) -> str:
    """Same convention as generate_table2.py."""
    if pd.isna(variant):
        return "NA"
    if not isinstance(variant, str):
        variant = str(variant)
    m = re.match(r"^(?:chr)?([^:]+):", variant, flags=re.IGNORECASE)
    chrom = m.group(1) if m else re.split(r"[:_\-]", variant)[0]
    chrom = chrom.lower().lstrip("chr")
    if chrom.isalpha():
        chrom = chrom.upper()
    return chrom


def _chrom_sort_key(ch):
    try:
        return (0, int(ch))
    except ValueError:
        return (1, ch)


def _split_chemicals(raw) -> List[str]:
    if pd.isna(raw):
        return []
    parts = CHEM_SPLIT_RE.split(str(raw))
    return [p.strip() for p in parts if p.strip()]


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------

def make_chromosome_figures(df: pd.DataFrame) -> None:
    """Purely descriptive chromosome-level figures (no significance concept
    here -- every row in this dataset is already the astore's output)."""
    df = df.copy()
    df["chromosome"] = df["variant"].apply(lambda v: extract_chromosome(v) if pd.notna(v) else "NA")

    variants_per_chrom = df.groupby("chromosome")["variant"].nunique().rename("n_variants").reset_index()
    genes_per_chrom = df.groupby("chromosome")["gene_symbol"].nunique().rename("n_genes").reset_index()
    merged = pd.merge(variants_per_chrom, genes_per_chrom, on="chromosome", how="outer").fillna(0)
    merged = merged.sort_values(by="chromosome", key=lambda s: s.map(_chrom_sort_key))

    sns.set_theme(style="whitegrid")

    plt.figure(figsize=(10, 6))
    ax = sns.barplot(data=merged, x="chromosome", y="n_variants", color="#4472C4")
    ax.set_xlabel("Chromosome")
    ax.set_ylabel("Number of unique variants")
    ax.set_title("Variants per chromosome (Table 2b)")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig(FIG_DIR / "variants_per_chromosome.png", dpi=300)
    plt.close()

    plt.figure(figsize=(8, 6))
    ax = sns.scatterplot(data=merged, x="n_genes", y="n_variants", s=100, color="#4472C4")
    for _, r in merged.iterrows():
        ax.text(r["n_genes"], r["n_variants"], str(r["chromosome"]), fontsize=9,
                 horizontalalignment="left", verticalalignment="bottom")
    ax.set_xlabel("Number of unique genes per chromosome")
    ax.set_ylabel("Number of unique variants per chromosome")
    ax.set_title("Genes vs variants per chromosome (Table 2b)")
    plt.tight_layout()
    plt.savefig(FIG_DIR / "genes_vs_variants_scatter.png", dpi=300)
    plt.close()


def make_gene_type_figure(df: pd.DataFrame) -> pd.DataFrame:
    """Bar chart of unique genes per gene_type. Returns the counts table,
    also saved to CSV."""
    genes = df.drop_duplicates(subset=["gene_symbol"])
    counts = genes["gene_type"].fillna("Unknown").value_counts().rename_axis("gene_type").reset_index(name="n_genes")
    counts = counts.sort_values("n_genes", ascending=False)
    counts.to_csv(OUT_DIR / "table2b_gene_type_counts.csv", index=False)

    plt.figure(figsize=(max(6, 0.5 * len(counts)), 5))
    ax = sns.barplot(data=counts, x="gene_type", y="n_genes", color="#4472C4")
    ax.set_xlabel("Gene type")
    ax.set_ylabel("Number of unique genes")
    ax.set_title("Gene type distribution")
    plt.xticks(rotation=45, ha="right")
    plt.tight_layout()
    plt.savefig(FIG_DIR / "gene_type_distribution.png", dpi=300)
    plt.close()
    return counts


def make_expression_figure(df: pd.DataFrame) -> str:
    """expressed_brain / expressed_neurons / expressed_glia: auto-detects
    whether these look like binary flags (all observed values in {0,1}) or
    continuous scores, and plots accordingly. Returns which mode was used
    ('binary' or 'continuous') for logging."""
    cols = ["expressed_brain", "expressed_neurons", "expressed_glia"]
    genes = df.drop_duplicates(subset=["gene_symbol"])[["gene_symbol"] + cols].copy()

    numeric = {}
    for c in cols:
        numeric[c] = pd.to_numeric(genes[c], errors="coerce")
    numeric_df = pd.DataFrame(numeric)

    observed_values = set(numeric_df.stack().dropna().unique().tolist())
    is_binary = observed_values.issubset({0.0, 1.0}) and len(observed_values) > 0

    fig_path = FIG_DIR / "expressed_brain_neurons_glia.png"

    if is_binary:
        print("[info] expressed_brain/neurons/glia detected as binary (0/1) flags -- "
              "plotting proportion of genes expressed.", file=sys.stderr)
        props = numeric_df.mean(numeric_only=True).reset_index()
        props.columns = ["compartment", "proportion_expressed"]
        props["compartment"] = props["compartment"].str.replace("expressed_", "", regex=False)

        plt.figure(figsize=(6, 5))
        ax = sns.barplot(data=props, x="compartment", y="proportion_expressed", color="#548235")
        ax.set_ylim(0, 1)
        ax.set_xlabel("Compartment")
        ax.set_ylabel("Proportion of unique genes expressed")
        ax.set_title("Gene expression by compartment (binary flags)")
        plt.tight_layout()
        plt.savefig(fig_path, dpi=300)
        plt.close()
        return "binary"
    else:
        print("[info] expressed_brain/neurons/glia detected as continuous scores (not all "
              "values in {0,1}) -- plotting histograms instead of proportions.", file=sys.stderr)
        fig, axes = plt.subplots(1, 3, figsize=(15, 4.5))
        for ax, c in zip(axes, cols):
            vals = numeric_df[c].dropna()
            if vals.empty:
                ax.text(0.5, 0.5, "No data", ha="center", va="center", transform=ax.transAxes)
                ax.axis("off")
                continue
            sns.histplot(vals, bins=30, ax=ax, color="#548235")
            ax.set_title(c.replace("expressed_", "").capitalize())
            ax.set_xlabel("Score")
        plt.tight_layout()
        plt.savefig(fig_path, dpi=300)
        plt.close(fig)
        return "continuous"


def make_score_histogram(df: pd.DataFrame, col: str, fig_name: str, title: str) -> None:
    """Histogram of a numeric score column, deduplicated by gene_symbol so a
    gene tested against multiple exposures/variants isn't overweighted."""
    genes = df.drop_duplicates(subset=["gene_symbol"])
    vals = pd.to_numeric(genes.get(col, pd.Series(dtype=float)), errors="coerce").dropna()

    fig_path = FIG_DIR / fig_name
    plt.figure(figsize=(8, 5))
    if not vals.empty:
        ax = sns.histplot(vals, bins=40, color="#ED7D31")
        ax.set_xlabel(col)
        ax.set_ylabel("Number of unique genes")
        ax.set_title(title)
        plt.tight_layout()
    else:
        plt.text(0.5, 0.5, f"No {col} data available", ha="center", va="center")
        plt.axis("off")
    plt.savefig(fig_path, dpi=300)
    plt.close()


def make_chemicals_figure(df: pd.DataFrame) -> pd.DataFrame:
    """Top-N most frequent CTD chemicals across all rows. Frequency counts
    rows (variant x exposure x gene), not unique genes, since the same
    chemical linked to different genes/variants is still relevant signal.
    Returns the frequency table, also saved to CSV."""
    counter: Counter = Counter()
    for raw in df.get("ctd_chemicals", pd.Series(dtype=object)):
        counter.update(_split_chemicals(raw))

    fig_path = FIG_DIR / "ctd_chemicals_top20.png"
    freq_path = OUT_DIR / "table2b_chemicals_frequency.csv"

    if not counter:
        freq = pd.DataFrame(columns=["chemical", "n_occurrences"])
        freq.to_csv(freq_path, index=False)
        plt.figure(figsize=(8, 5))
        plt.text(0.5, 0.5, "No ctd_chemicals data available", ha="center", va="center")
        plt.axis("off")
        plt.savefig(fig_path, dpi=300)
        plt.close()
        return freq

    freq = pd.DataFrame(counter.most_common(), columns=["chemical", "n_occurrences"])
    freq.to_csv(freq_path, index=False)

    top = freq.head(CTD_CHEMICALS_TOP_N).iloc[::-1]  # reverse for horizontal bar (largest on top)
    plt.figure(figsize=(8, max(5, 0.3 * len(top))))
    ax = sns.barplot(data=top, y="chemical", x="n_occurrences", color="#4472C4", orient="h")
    ax.set_xlabel("Occurrences")
    ax.set_ylabel("Chemical")
    ax.set_title(f"Top {min(CTD_CHEMICALS_TOP_N, len(freq))} CTD chemicals")
    plt.tight_layout()
    plt.savefig(fig_path, dpi=300)
    plt.close()
    return freq


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    from gene_environment.db.connection import cursor_scope, get_connection

    print(f"Calling stored routine: {ASTORE_NAME}() ...")
    df = call_stored_routine_to_df(ASTORE_NAME, get_connection, cursor_scope)
    for c in TABLE_COLUMNS:
        if c not in df.columns:
            df[c] = pd.NA

    df.to_csv(OUT_DIR / "table2b_raw_results.csv", index=False)

    if df.empty:
        print("[warn] no rows returned by the stored routine -- nothing further to do.", file=sys.stderr)
        return

    make_chromosome_figures(df)
    gene_type_counts = make_gene_type_figure(df)
    expression_mode = make_expression_figure(df)
    make_score_histogram(df, "als_opentargets_score", "als_opentargets_score_histogram.png",
                          "Distribution of ALS OpenTargets score (unique genes)")
    make_score_histogram(df, "neuro_plausibility_score", "neuro_plausibility_score_histogram.png",
                          "Distribution of neuro plausibility score (unique genes)")
    chemicals_freq = make_chemicals_figure(df)

    n_unique_genes = df["gene_symbol"].nunique()
    n_unique_variants = df["variant"].nunique()
    n_unique_exposures = df["exposure"].nunique()

    # --- Table 2b (top 10) ---
    doc_top10 = Document()
    make_landscape(doc_top10)
    doc_top10.add_heading("Table 2b. Gene annotations for significant variants (top 10)", level=1)
    doc_top10.add_paragraph(
        "Table shows the top 10 rows from get_significant_results_table_2b. This dataset has no "
        "p-value / coefficient columns, so no significance highlighting is applied here -- it "
        "reports gene annotation content (gene type, tissue expression, CTD chemicals, "
        "OpenTargets / neuro-plausibility scores) for the variant-gene-exposure combinations "
        "returned by the routine."
    )
    add_table_to_doc(doc_top10, df, max_rows=10)
    top10_path = OUT_DIR / "Table2b_top10.docx"
    doc_top10.save(top10_path)

    # --- Supplementary Word (full results + all figures) ---
    doc_full = Document()
    make_landscape(doc_full)
    doc_full.add_heading("Supplementary Table 2b: gene annotations, full results", level=1)
    doc_full.add_paragraph(
        f"Full results from get_significant_results_table_2b. {len(df)} rows covering "
        f"{n_unique_variants} unique variants, {n_unique_genes} unique genes, and "
        f"{n_unique_exposures} unique exposures. CTD chemical lists are truncated to "
        f"{CTD_CHEMICALS_TRUNCATE_CHARS} characters in this table; the full text is in "
        "table2b_raw_results.csv."
    )
    add_table_to_doc(doc_full, df)

    doc_full.add_heading("Figures", level=1)
    add_figure_to_doc(doc_full, FIG_DIR / "variants_per_chromosome.png",
                       "Figure 1. Number of unique variants per chromosome.")
    add_figure_to_doc(doc_full, FIG_DIR / "genes_vs_variants_scatter.png",
                       "Figure 2. Genes vs variants per chromosome.")
    add_figure_to_doc(doc_full, FIG_DIR / "gene_type_distribution.png",
                       "Figure 3. Distribution of gene types among unique genes.")
    expr_caption = (
        "Figure 4. Proportion of unique genes expressed in brain / neurons / glia."
        if expression_mode == "binary" else
        "Figure 4. Distribution of brain / neuron / glia expression scores among unique genes "
        "(detected as continuous, not binary flags -- see script notes)."
    )
    add_figure_to_doc(doc_full, FIG_DIR / "expressed_brain_neurons_glia.png", expr_caption, width_in=6.5)
    add_figure_to_doc(doc_full, FIG_DIR / "als_opentargets_score_histogram.png",
                       "Figure 5. Distribution of ALS OpenTargets score among unique genes.")
    add_figure_to_doc(doc_full, FIG_DIR / "neuro_plausibility_score_histogram.png",
                       "Figure 6. Distribution of neuro plausibility score among unique genes.")
    add_figure_to_doc(doc_full, FIG_DIR / "ctd_chemicals_top20.png",
                       f"Figure 7. Top {min(CTD_CHEMICALS_TOP_N, len(chemicals_freq))} most frequent "
                       "CTD chemicals across all rows (row-level occurrence count).", width_in=6.5)

    full_path = OUT_DIR / "Table2b_full_supplementary.docx"
    doc_full.save(full_path)

    print("Done.")
    print(f"Top 10 Word: {top10_path}")
    print(f"Full supplementary Word: {full_path}")
    print(f"Raw results CSV: {OUT_DIR / 'table2b_raw_results.csv'}")
    print(f"Gene type counts CSV: {OUT_DIR / 'table2b_gene_type_counts.csv'}")
    print(f"CTD chemicals frequency CSV: {OUT_DIR / 'table2b_chemicals_frequency.csv'}")
    print(f"Figures saved in: {FIG_DIR}")
    print(f"Expression columns detected as: {expression_mode}")


if __name__ == "__main__":
    main()
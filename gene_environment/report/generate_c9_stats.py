"""
generation_stats.py

For each environmental component E (e.g. seminativi_1500, risaie_1500,
vigneti_1500):
  1. Restrict the cohort to "close" samples only (E > 0).
  2. Use ONLY the variants that get_annotated_results() associated with
     that specific exposure (variant_exposure_map.json, built by
     build_c9_check.py).
  3. Within that restricted cohort, for each variant column (0/1), split
     mutated vs non-mutated and test against: sex, onset_site,
     diagnostic_delay, education_years, survival, survival_null, mutaz_bin
     (mutaz_bin / C9ORF72 is treated like the other variables here, but
     also gets its own advanced dedicated report).
  4. Everything is still split by generation (1 / 2), with a combined
     section that is the union of the two generations' significant
     results (no re-running of stats on pooled data).

Output layout per component E:
  stats/by_exposure/E/gen{1,2}_variant_stats.csv
  stats/by_exposure/E/plots/...
  stats/by_exposure/E/reports/gen{1,2}_report.docx
  stats/by_exposure/E/reports/combined_report.docx
  stats/by_exposure/E/reports/c9_report.docx (+ c9 CSVs/plots)

All output (logs, docx content, column headers) is in English.
"""

import json
import logging
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from scipy import stats

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ----------------------------------------------------------------------
# CONFIG
# ----------------------------------------------------------------------
OUT_DIR = Path("/mnt/cresla_prod/stefano_ge/c9_check")
MERGED_CSV = OUT_DIR / "c9_check_merged.csv"
VARIANT_EXPOSURE_MAP_FILE = OUT_DIR / "variant_exposure_map.json"

STATS_DIR = OUT_DIR / "stats"
BY_EXPOSURE_DIR = STATS_DIR / "by_exposure"

ID_COL = "id"
GENERATION_COL = "generation"

SEX_COL = "sex"
ONSET_SITE_COL = "onset_site"
DIAGNOSTIC_DELAY_COL = "diagnostic_delay"
EDUCATION_YEARS_COL = "education_years"
SURVIVAL_COL = "survival"
MUTAZ_RAW_COL = "mutaz"

ALPHA = 0.05  # significance threshold applied to the Bonferroni-corrected p-value

EXPOSURE_LABELS = {
    "seminativi_1500": "Arable land 1500mt",
    "vigneti_1500": "Vineyards 1500mt",
    "risaie_1500": "Rice fields 1500mt",
    "seminativi_1000": "Arable land 1000mt",
    "vigneti_1000": "Vineyards 1000mt",
    "risaie_1000": "Rice fields 1000mt",
}


def exposure_label(exposure: str) -> str:
    return EXPOSURE_LABELS.get(exposure, exposure)

CATEGORICAL_VARS = [SEX_COL, ONSET_SITE_COL, "mutaz_bin", "survival_null"]
CONTINUOUS_VARS = [DIAGNOSTIC_DELAY_COL, EDUCATION_YEARS_COL, SURVIVAL_COL]


# ----------------------------------------------------------------------
# Data loading
# ----------------------------------------------------------------------
def load_data():
    df = pd.read_csv(MERGED_CSV)
    with open(VARIANT_EXPOSURE_MAP_FILE) as f:
        variant_exposure_map = json.load(f)  # {column_name: [exposure, ...]}

    df["mutaz_bin"] = df[MUTAZ_RAW_COL].astype(str).str.contains("C9ORF72", na=False).astype(int)
    df["survival_null"] = df[SURVIVAL_COL].isna()
    df[EDUCATION_YEARS_COL] = pd.to_numeric(df[EDUCATION_YEARS_COL], errors="coerce").astype("Int64")

    log.info("Data loaded: %d rows, %d variant columns in exposure map", len(df), len(variant_exposure_map))
    return df, variant_exposure_map


def get_components(variant_exposure_map: dict, df_columns) -> dict:
    """Returns {exposure: [variant_columns...]} restricted to exposures
    that also exist as a column in the merged dataframe (needed to define
    'close' vs 'far'). Logs a warning for exposures without a matching
    column."""
    exposure_to_variants = {}
    for col, exposures in variant_exposure_map.items():
        for exp in exposures:
            exposure_to_variants.setdefault(exp, []).append(col)

    valid = {}
    for exp, cols in exposure_to_variants.items():
        if exp in df_columns:
            valid[exp] = sorted(set(cols))
        else:
            log.warning("Exposure '%s' has %d variants but no matching column in the merged CSV, skipping.", exp, len(cols))

    log.info("Environmental components found: %s", list(valid.keys()))
    return valid


# ----------------------------------------------------------------------
# Statistical tests
# ----------------------------------------------------------------------
def test_categorical(df: pd.DataFrame, group_col: str, var_col: str) -> dict:
    sub = df[[group_col, var_col]].dropna()
    table = pd.crosstab(sub[group_col], sub[var_col])
    if table.shape[0] < 2 or table.shape[1] < 2:
        return {"test": "n/a", "pvalue": None, "n0": None, "n1": None}

    if table.shape == (2, 2):
        _, p = stats.fisher_exact(table.values)
        test_name = "fisher_exact"
    else:
        _, p, _, _ = stats.chi2_contingency(table.values)
        test_name = "chi2"

    n0 = int((sub[group_col] == 0).sum())
    n1 = int((sub[group_col] == 1).sum())
    return {"test": test_name, "pvalue": p, "n0": n0, "n1": n1}


def test_continuous(df: pd.DataFrame, group_col: str, var_col: str) -> dict:
    sub = df[[group_col, var_col]].dropna()
    g0 = sub.loc[sub[group_col] == 0, var_col]
    g1 = sub.loc[sub[group_col] == 1, var_col]
    if len(g0) < 2 or len(g1) < 2:
        return {"test": "n/a", "pvalue": None, "n0": len(g0), "n1": len(g1)}

    _, p = stats.mannwhitneyu(g0, g1, alternative="two-sided")
    return {
        "test": "mannwhitney",
        "pvalue": p,
        "n0": len(g0),
        "n1": len(g1),
        "median0": float(g0.median()),
        "median1": float(g1.median()),
    }


def run_tests_for_generation(df_gen: pd.DataFrame, variant_cols: list) -> pd.DataFrame:
    """df_gen must already be restricted to the 'close' cohort (exposure > 0)
    and to a single generation. variant_cols must already be restricted to
    the variants relevant for the current exposure."""
    rows = []
    for variant in variant_cols:
        if variant not in df_gen.columns:
            continue
        group = df_gen[variant]
        if group.dropna().nunique() < 2:
            continue  # monomorphic variant in this subset, skip

        for var_col in CATEGORICAL_VARS:
            res = test_categorical(df_gen.assign(_grp=group), "_grp", var_col)
            rows.append({"variant": variant, "variable": var_col, "type": "categorical", **res})

        for var_col in CONTINUOUS_VARS:
            res = test_continuous(df_gen.assign(_grp=group), "_grp", var_col)
            rows.append({"variant": variant, "variable": var_col, "type": "continuous", **res})

    results = pd.DataFrame(rows)

    n_tests = int(results["pvalue"].notna().sum()) if not results.empty else 0
    if n_tests > 0:
        results["pvalue_bonf"] = (results["pvalue"] * n_tests).clip(upper=1.0)
    else:
        results["pvalue_bonf"] = pd.Series(dtype=float)
    results.attrs["n_tests"] = n_tests
    log.info("Bonferroni correction applied over %d tests run", n_tests)

    return results


def bonferroni_threshold(results: pd.DataFrame) -> float:
    n_tests = results.attrs.get("n_tests") or int(results["pvalue"].notna().sum())
    if n_tests == 0:
        return 0.0
    return ALPHA / n_tests


# ----------------------------------------------------------------------
# Plots
# ----------------------------------------------------------------------
def plot_categorical(df_gen: pd.DataFrame, variant: str, var_col: str, label: str, out_path: Path):
    sub = df_gen[[variant, var_col]].dropna()
    prop = pd.crosstab(sub[variant], sub[var_col], normalize="index")
    fig, ax = plt.subplots(figsize=(6, 4))
    prop.plot(kind="bar", stacked=True, ax=ax)
    ax.set_title(f"{label} | {variant} vs {var_col}")
    ax.set_xlabel(f"{variant} (0/1)")
    ax.set_ylabel("proportion")
    ax.legend(title=var_col, bbox_to_anchor=(1.05, 1), loc="upper left")
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def plot_continuous(df_gen: pd.DataFrame, variant: str, var_col: str, label: str, out_path: Path):
    sub = df_gen[[variant, var_col]].dropna()
    fig, ax = plt.subplots(figsize=(5, 4))
    sns.boxplot(data=sub, x=variant, y=var_col, ax=ax)
    sns.stripplot(data=sub, x=variant, y=var_col, ax=ax, color="black", alpha=0.4, size=3)
    ax.set_title(f"{label} | {variant} vs {var_col}")
    fig.tight_layout()
    fig.savefig(out_path, dpi=150)
    plt.close(fig)


def generate_plots(df_gen: pd.DataFrame, results: pd.DataFrame, label: str, plots_dir: Path) -> pd.DataFrame:
    plots_dir.mkdir(parents=True, exist_ok=True)
    sig = results[results["pvalue_bonf"].notna() & (results["pvalue_bonf"] < ALPHA)].copy()
    sig["plot_path"] = None

    for idx, row in sig.iterrows():
        variant, var_col, vtype = row["variant"], row["variable"], row["type"]
        fname = f"{label}_{variant}_{var_col}.png".replace("/", "_")
        out_path = plots_dir / fname
        try:
            if vtype == "categorical":
                plot_categorical(df_gen, variant, var_col, label, out_path)
            else:
                plot_continuous(df_gen, variant, var_col, label, out_path)
            sig.at[idx, "plot_path"] = str(out_path)
        except Exception as e:
            log.warning("Plot failed for %s/%s: %s", variant, var_col, e)

    log.info("%s: %d plots generated (Bonferroni p < %.2f)", label, sig["plot_path"].notna().sum(), ALPHA)
    return sig


# ----------------------------------------------------------------------
# Word report helpers
# ----------------------------------------------------------------------
def _set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold


def _add_significant_results_table(doc, sig: pd.DataFrame):
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 1"
    headers = ["Variant", "Variable", "Test", "p-value", "N group 0", "N group 1"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)

    for _, row in sig.sort_values("pvalue_bonf").iterrows():
        cells = table.add_row().cells
        is_sig = pd.notna(row["pvalue_bonf"]) and row["pvalue_bonf"] < ALPHA
        _set_cell_text(cells[0], str(row["variant"]))
        _set_cell_text(cells[1], str(row["variable"]))
        _set_cell_text(cells[2], str(row["test"]))
        _set_cell_text(cells[3], f"{row['pvalue']:.4g}", bold=is_sig)
        _set_cell_text(cells[4], str(row.get("n0", "")))
        _set_cell_text(cells[5], str(row.get("n1", "")))


def _add_generation_section(doc, results: pd.DataFrame, sig: pd.DataFrame, generation: int, heading_level: int = 2):
    from docx.shared import Inches

    threshold = bonferroni_threshold(results)
    doc.add_heading(f"Generation {generation}", level=heading_level)
    doc.add_paragraph(
        f"Total tests run: {len(results)}. "
        f"Bonferroni-corrected significance threshold: raw p-value < {threshold:.4g} "
        f"(alpha={ALPHA} / {results.attrs.get('n_tests', len(results))} tests). "
        f"Significant results: {len(sig)}. Significant p-values are shown in bold."
    )

    if sig.empty:
        doc.add_paragraph("No significant results in this generation.")
        return

    _add_significant_results_table(doc, sig)

    doc.add_heading("Plots", level=heading_level + 1)
    for _, row in sig.sort_values("pvalue_bonf").iterrows():
        if not row["plot_path"]:
            continue
        doc.add_paragraph(f"{row['variant']} vs {row['variable']} (p = {row['pvalue']:.4g})")
        doc.add_picture(row["plot_path"], width=Inches(5))


def build_word_report(results: pd.DataFrame, sig: pd.DataFrame, generation: int, exposure: str, cohort_label: str, out_path: Path):
    from docx import Document

    doc = Document()
    doc.add_heading(f"Variant analysis - Exposure: {exposure_label(exposure)} ({cohort_label}) - Generation {generation}", level=1)
    _add_generation_section(doc, results, sig, generation, heading_level=2)

    doc.save(out_path)
    log.info("Report saved: %s", out_path)


def build_combined_report(results1: pd.DataFrame, results2: pd.DataFrame,
                           sig1: pd.DataFrame, sig2: pd.DataFrame, exposure: str, cohort_label: str, out_path: Path):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(f"Variant analysis - Exposure: {exposure_label(exposure)} ({cohort_label}) - Generation 1 vs Generation 2", level=1)
    doc.add_paragraph(
        "This report lists Generation 1 and Generation 2 significant results "
        "separately, followed by a Combined section that is the union of "
        "both generations' significant results (simple concatenation, "
        "no statistics re-run on pooled data)."
    )

    _add_generation_section(doc, results1, sig1, 1, heading_level=2)
    _add_generation_section(doc, results2, sig2, 2, heading_level=2)

    doc.add_heading("Combined (Generation 1 + Generation 2)", level=2)
    sig1_tagged = sig1.copy()
    sig1_tagged["generation"] = 1
    sig2_tagged = sig2.copy()
    sig2_tagged["generation"] = 2
    union = pd.concat([sig1_tagged, sig2_tagged], ignore_index=True)
    doc.add_paragraph(
        f"Union of significant (variant, variable) results across both generations: {len(union)} rows "
        f"({len(sig1)} from Generation 1, {len(sig2)} from Generation 2)."
    )

    if not union.empty:
        table = doc.add_table(rows=1, cols=7)
        table.style = "Light Grid Accent 1"
        headers = ["Generation", "Variant", "Variable", "Test", "p-value", "N group 0", "N group 1"]
        for i, h in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)
        for _, row in union.sort_values(["generation", "pvalue_bonf"]).iterrows():
            cells = table.add_row().cells
            _set_cell_text(cells[0], str(row["generation"]))
            _set_cell_text(cells[1], str(row["variant"]))
            _set_cell_text(cells[2], str(row["variable"]))
            _set_cell_text(cells[3], str(row["test"]))
            _set_cell_text(cells[4], f"{row['pvalue']:.4g}", bold=True)
            _set_cell_text(cells[5], str(row.get("n0", "")))
            _set_cell_text(cells[6], str(row.get("n1", "")))

        doc.add_heading("Plots", level=2)
        for _, row in union.sort_values(["generation", "pvalue_bonf"]).iterrows():
            if not row["plot_path"]:
                continue
            doc.add_paragraph(f"Gen{row['generation']} | {row['variant']} vs {row['variable']} (p = {row['pvalue']:.4g})")
            doc.add_picture(row["plot_path"], width=Inches(5))

    doc.save(out_path)
    log.info("Combined report saved: %s", out_path)


# ----------------------------------------------------------------------
# C9 (mutaz_bin) advanced report - ALL variants (of this exposure), regardless
# of significance
# ----------------------------------------------------------------------
def _add_c9_contingency_counts(df_gen: pd.DataFrame, c9: pd.DataFrame) -> pd.DataFrame:
    counts = []
    for _, row in c9.iterrows():
        variant = row["variant"]
        sub = df_gen[[variant, "mutaz_bin"]].dropna()
        ct = pd.crosstab(sub[variant], sub["mutaz_bin"])

        def get(v, m):
            return int(ct.loc[v, m]) if v in ct.index and m in ct.columns else 0

        counts.append({
            "variant0_c9pos": get(0, 1),
            "variant0_c9neg": get(0, 0),
            "variant1_c9pos": get(1, 1),
            "variant1_c9neg": get(1, 0),
        })
    counts_df = pd.DataFrame(counts, index=c9.index)
    return pd.concat([c9, counts_df], axis=1)


def _c9_results_for_generation(df_gen: pd.DataFrame, results: pd.DataFrame, generation: int, plots_dir: Path) -> pd.DataFrame:
    c9 = results[results["variable"] == "mutaz_bin"].copy()
    c9 = c9.sort_values("pvalue")
    c9 = _add_c9_contingency_counts(df_gen, c9)

    plots_dir.mkdir(parents=True, exist_ok=True)
    c9["plot_path"] = None
    for idx, row in c9.iterrows():
        variant = row["variant"]
        fname = f"gen{generation}_{variant}_mutaz_bin.png".replace("/", "_")
        out_img = plots_dir / fname
        try:
            plot_categorical(df_gen, variant, "mutaz_bin", f"gen{generation}", out_img)
            c9.at[idx, "plot_path"] = str(out_img)
        except Exception as e:
            log.warning("C9 plot failed for %s: %s", variant, e)

    return c9


def _add_c9_generation_section(doc, results: pd.DataFrame, c9: pd.DataFrame, generation: int, heading_level: int = 2, include_plots: bool = True):
    from docx.shared import Inches

    threshold = bonferroni_threshold(results)
    doc.add_heading(f"Generation {generation}", level=heading_level)
    doc.add_paragraph(
        f"All variants (for this exposure) tested against mutaz_bin (C9ORF72), regardless of significance. "
        f"Total variants: {len(c9)}. "
        f"Bonferroni-corrected significance threshold: raw p-value < {threshold:.4g} "
        f"(alpha={ALPHA} / {results.attrs.get('n_tests', len(results))} tests). "
        f"Significant p-values are shown in bold."
    )

    if c9.empty:
        doc.add_paragraph("No variants available for this generation/exposure.")
        return

    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    headers = ["Variant", "Test", "p-value", "Variant=0 C9+", "Variant=0 C9-", "Variant=1 C9+", "Variant=1 C9-"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for _, row in c9.iterrows():
        cells = table.add_row().cells
        is_sig = pd.notna(row["pvalue_bonf"]) and row["pvalue_bonf"] < ALPHA
        _set_cell_text(cells[0], str(row["variant"]))
        _set_cell_text(cells[1], str(row["test"]))
        pv_txt = f"{row['pvalue']:.4g}" if pd.notna(row["pvalue"]) else "n/a"
        _set_cell_text(cells[2], pv_txt, bold=is_sig)
        _set_cell_text(cells[3], str(row.get("variant0_c9pos", "")))
        _set_cell_text(cells[4], str(row.get("variant0_c9neg", "")))
        _set_cell_text(cells[5], str(row.get("variant1_c9pos", "")))
        _set_cell_text(cells[6], str(row.get("variant1_c9neg", "")))

    if not include_plots:
        return

    doc.add_heading("Plots", level=heading_level + 1)
    for _, row in c9.iterrows():
        if not row["plot_path"]:
            continue
        doc.add_paragraph(f"{row['variant']} (p = {row['pvalue']:.4g})" if pd.notna(row["pvalue"]) else str(row["variant"]))
        doc.add_picture(row["plot_path"], width=Inches(4.5))


def _add_c9_combined_table(doc, union: pd.DataFrame, results_by_gen: dict):
    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = ["Generation", "Variant", "Test", "p-value", "Variant=0 C9+", "Variant=0 C9-", "Variant=1 C9+", "Variant=1 C9-"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for _, row in union.iterrows():
        threshold = bonferroni_threshold(results_by_gen[row["generation"]][0])
        is_sig = pd.notna(row["pvalue"]) and row["pvalue"] < threshold
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(row["generation"]))
        _set_cell_text(cells[1], str(row["variant"]))
        _set_cell_text(cells[2], str(row["test"]))
        pv_txt = f"{row['pvalue']:.4g}" if pd.notna(row["pvalue"]) else "n/a"
        _set_cell_text(cells[3], pv_txt, bold=is_sig)
        _set_cell_text(cells[4], str(row.get("variant0_c9pos", "")))
        _set_cell_text(cells[5], str(row.get("variant0_c9neg", "")))
        _set_cell_text(cells[6], str(row.get("variant1_c9pos", "")))
        _set_cell_text(cells[7], str(row.get("variant1_c9neg", "")))


def build_c9_report(df: pd.DataFrame, results_by_gen: dict, exposure: str, c9_dir: Path, out_path: Path):
    from docx import Document

    c9_dir.mkdir(parents=True, exist_ok=True)
    c9_plots_dir = c9_dir / "plots"

    doc = Document()
    doc.add_heading(f"C9ORF72 (mutaz_bin) advanced report - exposure: {exposure}", level=1)
    doc.add_paragraph(
        "This report includes ALL variants (for this exposure, 'close' cohort only) "
        "tested against mutaz_bin, regardless of statistical significance. "
        "Generation 1 and Generation 2 are shown separately, followed by a "
        "Combined section that is the union of both generations' rows "
        "(no statistics re-run on pooled data)."
    )

    c9_by_gen = {}
    for generation in (1, 2):
        if generation not in results_by_gen:
            continue
        results, df_gen = results_by_gen[generation]
        c9 = _c9_results_for_generation(df_gen, results, generation, c9_plots_dir)
        c9_by_gen[generation] = c9

        csv_path = c9_dir / f"c9_mutaz_bin_gen{generation}.csv"
        c9_out = c9.copy()
        c9_out.insert(0, "exposure", exposure)
        c9_out.to_csv(csv_path, index=False)
        log.info("C9 CSV saved for generation %d (%s): %s (%d variants)", generation, exposure, csv_path, len(c9))

        _add_c9_generation_section(doc, results, c9, generation, heading_level=2)

    union = None
    if 1 in c9_by_gen and 2 in c9_by_gen:
        doc.add_heading("Combined (Generation 1 + Generation 2)", level=2)
        doc.add_paragraph(
            "Union of all variant rows from both generations (concatenation, not a joint re-analysis)."
        )
        c9_1 = c9_by_gen[1].copy()
        c9_1["generation"] = 1
        c9_2 = c9_by_gen[2].copy()
        c9_2["generation"] = 2
        union = pd.concat([c9_1, c9_2], ignore_index=True).sort_values(["generation", "pvalue"])

        combined_csv = c9_dir / "c9_mutaz_bin_combined.csv"
        union_out = union.copy()
        union_out.insert(0, "exposure", exposure)
        union_out.to_csv(combined_csv, index=False)
        log.info("C9 combined CSV saved: %s", combined_csv)

        _add_c9_combined_table(doc, union, results_by_gen)

    doc.save(out_path)
    log.info("C9 report saved: %s", out_path)

    return c9_by_gen, union


# ----------------------------------------------------------------------
# Per-exposure pipeline
# ----------------------------------------------------------------------
def run_cohort_pipeline(df_cohort: pd.DataFrame, exposure: str, cohort_label: str, variant_cols: list, cohort_dir: Path):
    log.info("=== Exposure: %s | cohort: %s (%d variants) ===", exposure, cohort_label, len(variant_cols))
    log.info("%s cohort for %s: %d rows", cohort_label, exposure, len(df_cohort))

    plots_dir = cohort_dir / "plots"
    reports_dir = cohort_dir / "reports"
    c9_dir = reports_dir / "c9_report"
    reports_dir.mkdir(parents=True, exist_ok=True)

    sig_by_gen = {}
    results_for_c9 = {}  # generation -> (results, df_gen)
    results_by_gen = {}
    for generation in (1, 2):
        df_gen = df_cohort[df_cohort[GENERATION_COL] == generation].copy()
        log.info("[%s/%s] Generation %d: %d samples", exposure, cohort_label, generation, len(df_gen))
        if df_gen.empty:
            log.warning("[%s/%s] Generation %d is empty, skipping.", exposure, cohort_label, generation)
            continue

        results = run_tests_for_generation(df_gen, variant_cols)
        if results.empty:
            log.warning("[%s/%s] Generation %d produced no test results, skipping.", exposure, cohort_label, generation)
            continue

        results_by_gen[generation] = results
        results_for_c9[generation] = (results, df_gen)

        results_csv = cohort_dir / f"gen{generation}_variant_stats.csv"
        results_out = results.copy()
        results_out.insert(0, "cohort", cohort_label)
        results_out.insert(0, "exposure", exposure_label(exposure))
        results_out.to_csv(results_csv, index=False)
        log.info("[%s/%s] Results CSV saved for generation %d: %s", exposure, cohort_label, generation, results_csv)

        sig = generate_plots(df_gen, results, f"{exposure}_{cohort_label}_gen{generation}", plots_dir)
        sig_by_gen[generation] = sig

        build_word_report(results, sig, generation, exposure, cohort_label, reports_dir / f"gen{generation}_report.docx")

    if 1 in sig_by_gen and 2 in sig_by_gen:
        build_combined_report(
            results_by_gen[1], results_by_gen[2],
            sig_by_gen[1], sig_by_gen[2],
            exposure, cohort_label,
            reports_dir / "combined_report.docx",
        )

    c9_data = None
    if results_for_c9:
        c9_by_gen, union = build_c9_report(df_cohort, results_for_c9, f"{exposure_label(exposure)} ({cohort_label})", c9_dir, reports_dir / "c9_report.docx")
        c9_data = {
            "exposure": exposure,
            "cohort": cohort_label,
            "results_for_c9": results_for_c9,
            "c9_by_gen": c9_by_gen,
            "union": union,
        }

    return results_by_gen, sig_by_gen, c9_data


def build_close_vs_far_summary(exposure: str, close_res: dict, far_res: dict, out_path: Path):
    """Side-by-side summary: for every (variant, variable) tested in either
    cohort, show p-value and Bonferroni significance in close vs far, so it's
    easy to spot subgroup-specific (environment-dependent) effects."""
    from docx import Document

    doc = Document()
    doc.add_heading(f"Close vs Far summary - exposure: {exposure_label(exposure)}", level=1)
    doc.add_paragraph(
        "For each generation, compares the 'close' cohort (exposure > 0) against "
        "the 'far' cohort (exposure = 0) on the same variant/variable pairs. "
        "A pattern significant in 'close' but not in 'far' suggests a subgroup "
        "effect specific to that environmental exposure. Bold = significant "
        "after Bonferroni correction within that cohort/generation."
    )

    for generation in (1, 2):
        if generation not in close_res and generation not in far_res:
            continue
        doc.add_heading(f"Generation {generation}", level=2)

        r_close = close_res.get(generation)
        r_far = far_res.get(generation)
        thr_close = bonferroni_threshold(r_close) if r_close is not None else None
        thr_far = bonferroni_threshold(r_far) if r_far is not None else None

        keys = set()
        if r_close is not None:
            keys |= set(zip(r_close["variant"], r_close["variable"]))
        if r_far is not None:
            keys |= set(zip(r_far["variant"], r_far["variable"]))

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        headers = ["Variant", "Variable", "p-value (close)", "p-value (far)", "N close / N far"]
        for i, h in enumerate(headers):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)

        for variant, variable in sorted(keys):
            row_close = r_close[(r_close["variant"] == variant) & (r_close["variable"] == variable)] if r_close is not None else pd.DataFrame()
            row_far = r_far[(r_far["variant"] == variant) & (r_far["variable"] == variable)] if r_far is not None else pd.DataFrame()

            p_close = row_close["pvalue"].iloc[0] if not row_close.empty else None
            p_far = row_far["pvalue"].iloc[0] if not row_far.empty else None
            sig_close = p_close is not None and pd.notna(p_close) and thr_close and p_close < thr_close
            sig_far = p_far is not None and pd.notna(p_far) and thr_far and p_far < thr_far

            n_close = f"{row_close['n0'].iloc[0]}/{row_close['n1'].iloc[0]}" if not row_close.empty and pd.notna(row_close['n0'].iloc[0]) else "n/a"
            n_far = f"{row_far['n0'].iloc[0]}/{row_far['n1'].iloc[0]}" if not row_far.empty and pd.notna(row_far['n0'].iloc[0]) else "n/a"

            cells = table.add_row().cells
            _set_cell_text(cells[0], str(variant))
            _set_cell_text(cells[1], str(variable))
            _set_cell_text(cells[2], f"{p_close:.4g}" if p_close is not None and pd.notna(p_close) else "n/a", bold=sig_close)
            _set_cell_text(cells[3], f"{p_far:.4g}" if p_far is not None and pd.notna(p_far) else "n/a", bold=sig_far)
            _set_cell_text(cells[4], f"{n_close} / {n_far}")

    doc.save(out_path)
    log.info("Close vs Far summary saved: %s", out_path)


def run_exposure_pipeline(df: pd.DataFrame, exposure: str, variant_cols: list, exposure_dir: Path):
    df_close = df[df[exposure] > 0]
    df_far = df[df[exposure] == 0]

    close_results_by_gen, _, close_c9 = run_cohort_pipeline(df_close, exposure, "close", variant_cols, exposure_dir / "close")
    far_results_by_gen, _, far_c9 = run_cohort_pipeline(df_far, exposure, "far", variant_cols, exposure_dir / "far")

    build_close_vs_far_summary(exposure, close_results_by_gen, far_results_by_gen, exposure_dir / "close_vs_far_summary.docx")

    return [c for c in (close_c9, far_c9) if c is not None]


def _build_master_summary_table(doc, all_c9_data: list):
    """Initial table: every (exposure, cohort, generation, variant) row with
    pvalue < 0.05, across ALL exposures, with an extra 'Exposure' column."""
    summary_rows = []
    for entry in all_c9_data:
        label = f"{exposure_label(entry['exposure'])} ({entry['cohort']})"
        c9_by_gen = entry["c9_by_gen"]
        for generation, c9 in c9_by_gen.items():
            sig = c9[c9["pvalue"].notna() & (c9["pvalue"] < 0.05)]
            for _, row in sig.iterrows():
                summary_rows.append({
                    "exposure": label,
                    "generation": generation,
                    "variant": row["variant"],
                    "test": row["test"],
                    "pvalue": row["pvalue"],
                    "pvalue_bonf": row.get("pvalue_bonf"),
                    "variant0_c9pos": row.get("variant0_c9pos", ""),
                    "variant0_c9neg": row.get("variant0_c9neg", ""),
                    "variant1_c9pos": row.get("variant1_c9pos", ""),
                    "variant1_c9neg": row.get("variant1_c9neg", ""),
                })

    doc.add_heading("Summary: all exposures, p-value < 0.05", level=2)
    if not summary_rows:
        doc.add_paragraph("No (exposure, cohort, generation, variant) row reached p-value < 0.05.")
        return

    summary = pd.DataFrame(summary_rows).sort_values("pvalue")
    doc.add_paragraph(
        f"{len(summary)} rows across all exposures/cohorts/generations with raw p-value < 0.05. "
        f"Bold = also significant after Bonferroni correction (within its own exposure/cohort/generation)."
    )

    table = doc.add_table(rows=1, cols=8)
    table.style = "Light Grid Accent 1"
    headers = ["Exposure", "Generation", "Variant", "Test", "p-value", "Variant=0 C9+", "Variant=0 C9-", "Variant=1 C9+/C9- (1)"]
    # NB: last header shortened for space; actual data still has separate C9+/C9- for variant=1 below
    headers = ["Exposure", "Generation", "Variant", "Test", "p-value", "Var=0 C9+", "Var=0 C9-", "Var=1 C9+"]
    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True)
    for _, row in summary.iterrows():
        is_sig = pd.notna(row["pvalue_bonf"]) and row["pvalue_bonf"] < ALPHA
        cells = table.add_row().cells
        _set_cell_text(cells[0], str(row["exposure"]))
        _set_cell_text(cells[1], str(row["generation"]))
        _set_cell_text(cells[2], str(row["variant"]))
        _set_cell_text(cells[3], str(row["test"]))
        _set_cell_text(cells[4], f"{row['pvalue']:.4g}", bold=is_sig)
        _set_cell_text(cells[5], str(row["variant0_c9pos"]))
        _set_cell_text(cells[6], str(row["variant0_c9neg"]))
        _set_cell_text(cells[7], str(row["variant1_c9pos"]))


def build_master_c9_report(all_c9_data: list, out_path: Path):
    """Single docx with ALL exposures' C9 tables (per generation + combined),
    one after another, so everything can be reviewed in one file. No images
    (those are already in each exposure's own c9_report.docx)."""
    from docx import Document

    doc = Document()
    doc.add_heading("C9ORF72 (mutaz_bin) - all exposures", level=1)
    doc.add_paragraph(
        "This document collects, for every environmental exposure and cohort "
        "(close/far), the same C9 tables already available in each exposure's "
        "individual c9_report.docx, gathered here in a single file for convenience. "
        "Plots are omitted here; see the per-exposure c9_report.docx for images."
    )

    _build_master_summary_table(doc, all_c9_data)

    for entry in all_c9_data:
        label = f"{exposure_label(entry['exposure'])} ({entry['cohort']})"
        doc.add_heading(f"Exposure: {label}", level=2)

        results_for_c9 = entry["results_for_c9"]
        c9_by_gen = entry["c9_by_gen"]
        union = entry["union"]

        for generation in (1, 2):
            if generation not in c9_by_gen:
                continue
            results, _ = results_for_c9[generation]
            _add_c9_generation_section(doc, results, c9_by_gen[generation], generation, heading_level=3, include_plots=False)

        if union is not None:
            doc.add_heading("Combined (Generation 1 + Generation 2)", level=3)
            doc.add_paragraph(
                "Union of all variant rows from both generations (concatenation, not a joint re-analysis)."
            )
            _add_c9_combined_table(doc, union, results_for_c9)

    doc.save(out_path)
    log.info("Master C9 report saved: %s", out_path)


# ----------------------------------------------------------------------
# Main
# ----------------------------------------------------------------------
def main():
    STATS_DIR.mkdir(parents=True, exist_ok=True)
    BY_EXPOSURE_DIR.mkdir(parents=True, exist_ok=True)

    df, variant_exposure_map = load_data()
    components = get_components(variant_exposure_map, set(df.columns))

    if not components:
        raise RuntimeError(
            "No environmental component matched a column in the merged CSV. "
            "Check variant_exposure_map.json exposure names against c9_check_merged.csv columns."
        )

    all_c9_data = []
    for exposure, variant_cols in components.items():
        exposure_dir = BY_EXPOSURE_DIR / exposure
        c9_data = run_exposure_pipeline(df, exposure, variant_cols, exposure_dir)
        all_c9_data.extend(c9_data)

    if all_c9_data:
        build_master_c9_report(all_c9_data, BY_EXPOSURE_DIR / "all_c9_report.docx")


if __name__ == "__main__":
    main()
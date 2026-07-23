#!/usr/bin/env python3
"""
generate_table1.py

Generates Table 1 for the paper (descriptive statistics for two cohorts) from:
  - a CSV containing clinical/environmental patient data (one row per id)
  - a CSV mapping id -> generation/cohort, produced by build_cohort_mapping.py
    by reading VCF headers (avoids loading gen.parquet, too heavy/unstable)

Output (in OUTPUT_DIR):
  - table1_stats.csv        -> raw statistics table, reusable
  - Table1.docx             -> Word table ready for the paper
  - figures/*.png           -> boxplots/barplots comparing the two cohorts

Usage:
    python build_cohort_mapping.py   # generates id -> generation mapping
    python generate_table1.py        # generates Table 1

Modify only the CONFIG section below to adapt paths/column names.
"""

import os
import sys
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

warnings.filterwarnings("ignore")

# ============================================================
# CONFIG — edit here
# ============================================================

CSV_PATH = "/srv/python-projects/gene_environment_v2/data/componenti_ambientali_full.csv"

COHORT_MAPPING_CSV = "output/table1/id_generation_mapping.csv"

OUTPUT_DIR = Path("output/table1")

ID_COL_CSV = "id"
ID_COL_MAPPING = "id"

COHORT_COL = "generazione"

# If more than 2 distinct values exist, specify EXACTLY which two to compare.
# Example: COHORT_VALUES = ["gen1", "gen2"]
COHORT_VALUES = None  # None = auto-detect (requires exactly 2 distinct values)

# Human-readable labels for the two cohorts
COHORT_LABELS = None  # e.g. {"gen1": "Cohort 1 (PARALS)", "gen2": "Cohort 2 (Sweden)"}

CATEGORICAL_VARS = ["sex", "onset_site"]
NUMERIC_VARS = [
    "diagnostic_delay",
    "onset_age",
    "survival",
    "seminativi_1500",
    "vigneti_1500",
    "risaie_1500",
    "seminativi_1000",
    "vigneti_1000",
    "risaie_1000",
]

VAR_LABELS = {
    "sex": "Sex",
    "onset_site": "Onset site",
    "diagnostic_delay": "Diagnostic delay (months)",
    "onset_age": "Age at onset (years)",
    "survival": "Survival (years)",
    "seminativi_1500": "Arable land within 1500 m (%)",
    "vigneti_1500": "Vineyards within 1500 m (%)",
    "risaie_1500": "Rice fields within 1500 m (%)",
    "seminativi_1000": "Arable land within 1000 m (%)",
    "vigneti_1000": "Vineyards within 1000 m (%)",
    "risaie_1000": "Rice fields within 1000 m (%)",
}

ALPHA = 0.05

# ============================================================
# FUNCTIONS
# ============================================================

def load_data():
    print(f"Loading CSV: {CSV_PATH}")
    df = pd.read_csv(CSV_PATH)
    if ID_COL_CSV not in df.columns:
        sys.exit(f"ERROR: id column '{ID_COL_CSV}' not found in CSV. Columns: {list(df.columns)}")

    print(f"Loading cohort mapping from CSV: {COHORT_MAPPING_CSV}")
    if not Path(COHORT_MAPPING_CSV).exists():
        sys.exit(
            f"ERROR: {COHORT_MAPPING_CSV} not found.\n"
            f"Generate mapping first with: python build_cohort_mapping.py"
        )
    gen = pd.read_csv(COHORT_MAPPING_CSV)

    if ID_COL_MAPPING not in gen.columns:
        sys.exit(f"ERROR: id column '{ID_COL_MAPPING}' not found in mapping CSV. Columns: {list(gen.columns)}")
    if COHORT_COL not in gen.columns:
        sys.exit(f"ERROR: cohort column '{COHORT_COL}' not found in mapping CSV. Columns: {list(gen.columns)}")

    gen = gen[[ID_COL_MAPPING, COHORT_COL]].drop_duplicates()

    merged = df.merge(
        gen, left_on=ID_COL_CSV, right_on=ID_COL_MAPPING, how="inner"
    )
    n_lost = len(df) - len(merged)
    if n_lost > 0:
        print(f"WARNING: {n_lost} patients in CSV not found in cohort mapping (excluded).")

    return merged


def resolve_cohorts(merged):
    values = sorted(merged[COHORT_COL].dropna().unique().tolist())

    if COHORT_VALUES is not None:
        chosen = COHORT_VALUES
        missing = [v for v in chosen if v not in values]
        if missing:
            sys.exit(f"ERROR: COHORT_VALUES {missing} not present in '{COHORT_COL}'. Found: {values}")
    else:
        if len(values) != 2:
            sys.exit(
                f"ERROR: found {len(values)} distinct values in '{COHORT_COL}': {values}.\n"
                f"Set COHORT_VALUES = [value1, value2] in CONFIG."
            )
        chosen = values  # sorted ensures stable ordering

    labels = COHORT_LABELS or {v: str(v) for v in chosen}
    for v in chosen:
        labels.setdefault(v, str(v))

    sub = merged[merged[COHORT_COL].isin(chosen)].copy()
    print(f"Selected cohorts: {chosen} -> N = {sub[COHORT_COL].value_counts().to_dict()}")
    return sub, chosen, labels


def is_normal(series, alpha=0.05):
    series = series.dropna()
    if len(series) < 8:
        return True
    stat, p = stats.shapiro(series)
    return p > alpha


def fmt_p(p):
    if pd.isna(p):
        return "-"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def summarize_numeric(sub, var, cohort_col, groups):
    rows = []
    g1 = sub.loc[sub[cohort_col] == groups[0], var].dropna()
    g2 = sub.loc[sub[cohort_col] == groups[1], var].dropna()

    normal = is_normal(g1) and is_normal(g2)

    if normal:
        desc1 = f"{g1.mean():.2f} ± {g1.std():.2f}"
        desc2 = f"{g2.mean():.2f} ± {g2.std():.2f}"
        stat, p = stats.ttest_ind(g1, g2, equal_var=False, nan_policy="omit")
        test_used = "Welch t-test"
    else:
        desc1 = f"{g1.median():.2f} [{g1.quantile(.25):.2f}-{g1.quantile(.75):.2f}]"
        desc2 = f"{g2.median():.2f} [{g2.quantile(.25):.2f}-{g2.quantile(.75):.2f}]"
        if len(g1) > 0 and len(g2) > 0:
            stat, p = stats.mannwhitneyu(g1, g2, alternative="two-sided")
        else:
            p = np.nan
        test_used = "Mann–Whitney U"

    rows.append(
        {
            "variable": VAR_LABELS.get(var, var),
            "type": "numeric",
            "group1_n": len(g1),
            "group1_stat": desc1,
            "group2_n": len(g2),
            "group2_stat": desc2,
            "test": test_used,
            "p_value": p,
            "p_value_fmt": fmt_p(p),
        }
    )
    return rows


def summarize_categorical(sub, var, cohort_col, groups):
    rows = []
    ct = pd.crosstab(sub[var], sub[cohort_col])
    ct = ct[[c for c in groups if c in ct.columns]]

    if ct.shape[0] == 2 and (ct.values < 5).any():
        _, p = stats.fisher_exact(ct.values) if ct.shape == (2, 2) else (None, np.nan)
        test_used = "Fisher exact"
    else:
        chi2, p, _, _ = stats.chi2_contingency(ct)
        test_used = "Chi-square"

    first = True
    for level in ct.index:
        n1 = ct.loc[level, groups[0]] if groups[0] in ct.columns else 0
        n2 = ct.loc[level, groups[1]] if groups[1] in ct.columns else 0
        tot1 = ct[groups[0]].sum() if groups[0] in ct.columns else 0
        tot2 = ct[groups[1]].sum() if groups[1] in ct.columns else 0
        pct1 = 100 * n1 / tot1 if tot1 else 0
        pct2 = 100 * n2 / tot2 if tot2 else 0
        rows.append(
            {
                "variable": f"{VAR_LABELS.get(var, var)} - {level}" if not first else VAR_LABELS.get(var, var),
                "type": "categorical",
                "group1_n": tot1,
                "group1_stat": f"{n1} ({pct1:.1f}%)",
                "group2_n": tot2,
                "group2_stat": f"{n2} ({pct2:.1f}%)",
                "test": test_used if first else "",
                "p_value": p if first else np.nan,
                "p_value_fmt": fmt_p(p) if first else "",
            }
        )
        first = False
    return rows


def build_stats_table(sub, cohort_col, groups):
    all_rows = []
    for var in CATEGORICAL_VARS:
        all_rows.extend(summarize_categorical(sub, var, cohort_col, groups))
    for var in NUMERIC_VARS:
        all_rows.extend(summarize_numeric(sub, var, cohort_col, groups))
    return pd.DataFrame(all_rows)


# ------------------------------------------------------------
# FIGURES
# ------------------------------------------------------------

def make_figures(sub, cohort_col, groups, labels, fig_dir):
    fig_dir.mkdir(parents=True, exist_ok=True)
    sns.set_style("whitegrid")
    palette = {groups[0]: "#4C72B0", groups[1]: "#DD8452"}

    plot_df = sub.copy()
    plot_df["Cohort"] = plot_df[cohort_col].map(labels)

    for var in NUMERIC_VARS:
        if var not in plot_df.columns:
            continue
        fig, ax = plt.subplots(figsize=(5, 4))
        sns.boxplot(
            data=plot_df,
            x="Cohort",
            y=var,
            ax=ax,
            palette=[palette[g] for g in groups],
        )
        sns.stripplot(
            data=plot_df, x="Cohort", y=var, ax=ax, color="black", alpha=0.3, size=3, jitter=True
        )
        ax.set_title(VAR_LABELS.get(var, var))
        ax.set_xlabel("")
        ax.set_ylabel(VAR_LABELS.get(var, var))
        fig.tight_layout()
        fig.savefig(fig_dir / f"boxplot_{var}.png", dpi=200)
        plt.close(fig)

    for var in CATEGORICAL_VARS:
        if var not in plot_df.columns:
            continue
        fig, ax = plt.subplots(figsize=(5, 4))
        ct = pd.crosstab(plot_df["Cohort"], plot_df[var], normalize="index") * 100
        ct.plot(kind="bar", stacked=True, ax=ax, colormap="tab10")
        ax.set_ylabel("%")
        ax.set_xlabel("")
        ax.set_title(VAR_LABELS.get(var, var))
        ax.legend(title=var, bbox_to_anchor=(1.02, 1), loc="upper left", fontsize=8)
        fig.tight_layout()
        fig.savefig(fig_dir / f"barplot_{var}.png", dpi=200)
        plt.close(fig)

    print(f"Figures saved in: {fig_dir}")


# ------------------------------------------------------------
# WORD TABLE
# ------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def make_docx_table(stats_df, groups, labels, n_total, output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    title = doc.add_paragraph()
    run = title.add_run("Table 1. Clinical and environmental characteristics of the two cohorts")
    run.bold = True
    run.font.size = Pt(12)

    n1 = n_total.get(groups[0], 0)
    n2 = n_total.get(groups[1], 0)

    col_headers = [
        "Variable",
        f"{labels[groups[0]]} (n={n1})",
        f"{labels[groups[1]]} (n={n2})",
        "Test",
        "p",
    ]

    table = doc.add_table(rows=1, cols=len(col_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Cm(6.0), Cm(4.0), Cm(4.0), Cm(3.0), Cm(2.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w

    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(col_headers):
        hdr_cells[i].text = htext
        hdr_cells[i].width = widths[i]
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "D9D9D9")

    for _, row in stats_df.iterrows():
        cells = table.add_row().cells
        values = [
            row["variable"],
            row["group1_stat"],
            row["group2_stat"],
            row["test"],
            row["p_value_fmt"],
        ]
        for i, v in enumerate(values):
            cells[i].text = "" if pd.isna(v) else str(v)
            cells[i].width = widths[i]
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
                    if i == 4 and row["p_value"] is not np.nan and not pd.isna(row["p_value"]) and row["p_value"] < ALPHA:
                        r.bold = True

    note = doc.add_paragraph()
    note_run = note.add_run(
        "Numeric variables: mean ± SD (Welch t-test) if normally distributed, "
        "otherwise median [IQR] (Mann–Whitney U). Categorical variables: n (%) "
        "(Chi-square or Fisher exact test if expected counts <5)."
    )
    note_run.italic = True
    note_run.font.size = Pt(8)

    doc.save(output_path)
    print(f"Word table saved in: {output_path}")


# ============================================================
# MAIN
# ============================================================

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    fig_dir = OUTPUT_DIR / "figures"

    merged = load_data()
    sub, groups, labels = resolve_cohorts(merged)

    n_total = sub[COHORT_COL].value_counts().to_dict()

    stats_df = build_stats_table(sub, COHORT_COL, groups)

    csv_out = OUTPUT_DIR / "table1_stats.csv"
    stats_df.to_csv(csv_out, index=False)
    print(f"Statistics CSV saved in: {csv_out}")

    make_figures(sub, COHORT_COL, groups, labels, fig_dir)

    docx_out = OUTPUT_DIR / "Table1.docx"
    make_docx_table(stats_df, groups, labels, n_total, docx_out)

    print("\nDone. Output in:", OUTPUT_DIR.resolve())


if __name__ == "__main__":
    main()

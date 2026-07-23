"""
Analisi G x E cromosoma-per-cromosoma con GENetLib.

Si integra nel resto della pipeline `gene_environment` (stesso config.py,
stesso logging, stesse funzioni di normalizzazione id e caricamento PCA di
build_dataset.py / pca_utils.py) invece di reinventarli.

--------------------------------------------------------------------------
DIFFERENZE RISPETTO AL RESTO DELLA PIPELINE (modeling.py / orchestrator.py)
--------------------------------------------------------------------------
Il resto della pipeline testa UNA variante alla volta con OLS + matching +
permutazioni, con la formula:

    onset_age ~ variant * (risaie_1500) + sex + PC1 + ... + PCk

(interazione variant x exposure, sex e PCA come covariate additive: vedi
modeling.build_formula). Questo modulo fa concettualmente la STESSA cosa
(PCA e "sex" additive, exposure interagisce con la genetica) ma con un
modello diverso: invece di N regressioni OLS separate (una per variante),
usa GENetLib (rete neurale con penalita' MCP + L2, package
XMU-Kuangnan-Fang-Team/GENetLib) per stimare TUTTI gli SNP di un cromosoma
insieme, con selezione automatica delle varianti importanti.

Per ottenere lo stesso "PCA/sex additive, mai interazione G x PCA" qui si
procede in due passaggi (coerente con l'approccio "correggi poi testa" gia'
usato altrove nel progetto per la struttura di popolazione):

  STEP A (in questo modulo, `compute_pca_corrected_residuals`):
      onset_age = [PCA_1..k, sex] . delta + eps      (OLS, per generazione)
      onset_age' = onset_age - [PCA, sex] . delta_hat

  STEP B (GENetLib, per cromosoma):
      onset_age' = G . beta + (G x E) . theta,   E = SOLO risaie_1500

GENetLib (funzione scalar_ge) non ha un parametro "solo additivo, non
interattivo" per una covariata: qualunque colonna passata come E viene
automaticamente incrociata con OGNI SNP. Per questo "sex" e le PCA NON
possono essere passate come E a GENetLib (violerebbe il vincolo "niente
G x PCA"): vengono invece rimosse dal fenotipo nello step di regressione
preliminare, esattamente come le PCA.

--------------------------------------------------------------------------
FORMATO REALE DI RAW_FILE (gen.parquet), come costruito da
vcf_pipeline/vcf_to_parquet.py:
  - formato WIDE: indice = id campione (gia' passato da clean_sample_id
    durante la conversione VCF->parquet, poi ripulito ANCHE qui per
    sicurezza), colonne = una per variante.
  - nome colonna variante: "{CHROM}_{POS}_{REF}_{ALT}" (build_variant_label
    usa lo stesso schema per il resto della pipeline). Il prefisso CHROM
    puo' essere "1".."22" oppure "chr1".."chr22" a seconda di come il VCF
    sorgente nominava i contig (vedi extract_matrix.resolve_chrom_name):
    qui viene rilevato automaticamente.
  - valori: int8 0/1 (presenza di almeno un allele mutato: la
    binarizzazione avviene gia' in vcf_to_parquet.merge_chromosome,
    "arr[arr > 0] = 1" -- NON e' un dosaggio 0/1/2).
  - NON esiste alcuna colonna "generazione" nel parquet: la generazione di
    ogni campione si ottiene SOLO da sample_generation_map.csv (prodotto da
    vcf_to_parquet.save_sample_generation_map), esattamente come fa
    build_dataset._build_narrow_covariates. Qui la stessa logica di
    risoluzione e' replicata in `resolve_generation_map`.
--------------------------------------------------------------------------
"""
from __future__ import annotations

import json
import os
import time
import traceback
from concurrent.futures import ProcessPoolExecutor, as_completed
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

import duckdb
import numpy as np
import pandas as pd
import pyarrow.parquet as pq
import torch
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from sklearn.linear_model import LinearRegression
from sklearn.preprocessing import StandardScaler

try:
    from tqdm import tqdm
    _HAS_TQDM = True
except ImportError:
    _HAS_TQDM = False

try:
    from docx import Document
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

from GENetLib.scalar_ge import scalar_ge

from gene_environment.config import Config, get_config, _env, _env_int, _env_float, _env_bool, _env_list
from gene_environment.logging_utils import configure_logging, get_logger
from gene_environment.utils.id_utils import clean_sample_id
from gene_environment.utils.pca_utils import load_pca_covariates, PCA_ID_COLUMN
from gene_environment.vcf_pipeline.vcf_to_parquet import CHROMOSOMES

log = get_logger(__name__)

# Codifica sesso identica a build_dataset.py (SEX_ENCODING), tenuta qui
# separata (non importata) perche' in build_dataset.py e' locale alla
# funzione e non esposta come costante di modulo.
SEX_ENCODING = {"M": 1, "F": 0}


# ============================================================================
# 1. CONFIGURAZIONE SPECIFICA DI GENetLib (variabili GXE_*)
# ============================================================================
# Riusa gli helper privati di gene_environment.config (_env, _env_int, ...)
# per restare nello stesso stile/convenzione del resto del progetto: tutti i
# valori vengono da variabili d'ambiente / .env, con default sensati.

def _env_float_optional(name: str) -> Optional[float]:
    val = os.environ.get(name)
    return float(val) if val not in (None, "") else None


def _env_int_list(name: str, default: str) -> list:
    return [int(v) for v in _env_list(name, default)]


@dataclass(frozen=True)
class GXEConfig:
    # Cromosomi da processare (default: tutti quelli prodotti da build-matrix)
    chromosomes: list = field(default_factory=lambda: _env_list("GXE_CHROMOSOMES", ",".join(CHROMOSOMES)) or CHROMOSOMES)

    # Generazioni da includere nel run. Se vuoto, usa SOLO cfg.generation
    # (comportamento identico al resto della pipeline, che processa una
    # generazione per run). Se valorizzato (es. "1,2"), le residualizza
    # PCA-per-generazione e le combina in un unico modello GENetLib per
    # cromosoma (analisi congiunta delle due coorti).
    generations: list = field(default_factory=lambda: _env_int_list("GXE_GENERATIONS", ""))

    # Iperparametri scalar_ge (vedi docstring di run_genetlib_scalar_ge per
    # la corrispondenza con i concetti "lam"/"alpha"/"max_iter")
    num_hidden_layers: int = field(default_factory=lambda: _env_int("GXE_NUM_HIDDEN_LAYERS", 2))
    nodes_hidden_layer: list = field(default_factory=lambda: _env_int_list("GXE_NODES_HIDDEN_LAYER", "64,16"))
    num_epochs: int = field(default_factory=lambda: _env_int("GXE_NUM_EPOCHS", 100))
    learning_rate1: float = field(default_factory=lambda: _env_float("GXE_LEARNING_RATE1", 0.02))
    learning_rate2: float = field(default_factory=lambda: _env_float("GXE_LEARNING_RATE2", 0.01))
    lambda1: Optional[float] = field(default_factory=lambda: _env_float_optional("GXE_LAMBDA1"))
    lambda2: float = field(default_factory=lambda: _env_float("GXE_LAMBDA2", 0.05))
    Lambda: float = field(default_factory=lambda: _env_float("GXE_LAMBDA_L2", 0.05))
    split_type: int = field(default_factory=lambda: _env_int("GXE_SPLIT_TYPE", 0))
    ratio: list = field(default_factory=lambda: _env_int_list("GXE_RATIO", "7,3"))

    # Soglia di significativita' (frazione del massimo |peso| assoluto,
    # convenzione nativa di GENetLib -- NON un p-value)
    significance_threshold: float = field(default_factory=lambda: _env_float("GXE_SIGNIFICANCE_THRESHOLD", 0.3))

    min_samples_required: int = field(default_factory=lambda: _env_int("GXE_MIN_SAMPLES", 30))

    output_dir: str = field(default_factory=lambda: _env("GXE_OUTPUT_DIR", "./output/gxe_genetlib"))
    save_plots: bool = field(default_factory=lambda: _env_bool("GXE_SAVE_PLOTS", True))
    save_word_summary: bool = field(default_factory=lambda: _env_bool("GXE_SAVE_WORD_SUMMARY", True))
    torch_threads_per_worker: int = field(default_factory=lambda: _env_int("GXE_TORCH_THREADS", 1))


# ============================================================================
# 2. STEP A - REGRESSIONE PRELIMINARE onset_age ~ PCA + covariate ADDITIVE
# ============================================================================

def resolve_generation_map(cfg: Config) -> Optional[pd.DataFrame]:
    """
    Replica la logica di priorita' usata in
    vcf_pipeline.build_dataset._build_narrow_covariates:
        1) sample_generation_map.csv (cfg.sample_generation_map, o
           <OUTPUT_FOLDER>/sample_generation_map.csv se non impostato)
        2) colonna cfg.env_generation_col nel file ambientale (legacy)
        3) nessun filtro possibile -> None (tutte le righe, generazione
           sconosciuta: NON utilizzabile per Step A, che richiede la
           generazione per scegliere il file PCA giusto)

    Ritorna un DataFrame [id, generation] oppure None se non c'e' modo di
    determinare la generazione dei campioni.
    """
    map_path = cfg.sample_generation_map or os.path.join(cfg.output_folder, "sample_generation_map.csv")
    if os.path.exists(map_path):
        gen_map = pd.read_csv(map_path, dtype={"id": str})
        gen_map["id"] = gen_map["id"].astype(str)
        log.info("Mappa id->generazione caricata da %s (%d campioni)", map_path, len(gen_map))
        return gen_map[["id", "generation"]]

    log.warning(
        "Nessuna mappa id->generazione trovata in %s. Se il file ambientale ha una "
        "colonna di generazione, impostala in ENV_GENERATION_COL; altrimenti Step A "
        "(regressione PCA, specifica per generazione) non puo' essere eseguito "
        "correttamente.", map_path,
    )
    return None


def _encode_covariates(df_env: pd.DataFrame, covariates: list) -> tuple[pd.DataFrame, list]:
    """Codifica le covariate additive (es. 'sex') esattamente come
    build_dataset._build_narrow_covariates, ritornando il dataframe con le
    colonne codificate e la lista dei nomi colonna effettivamente numerici
    e utilizzabili nella regressione."""
    df = df_env.copy()
    resolved = []
    for cov in covariates:
        if cov not in df.columns:
            log.warning("Covariata '%s' non trovata nel file ambientale: ignorata.", cov)
            continue
        if cov == "sex":
            unmapped = set(df["sex"].dropna().unique()) - set(SEX_ENCODING.keys())
            if unmapped:
                raise ValueError(f"'sex': valori non riconosciuti {unmapped}, aggiorna SEX_ENCODING")
            df["sex"] = df["sex"].map(SEX_ENCODING).astype(float)
        else:
            df[cov] = pd.to_numeric(df[cov], errors="coerce")
        resolved.append(cov)
    return df, resolved


def load_environment_and_phenotype(cfg: Config, logger) -> pd.DataFrame:
    """Carica ENV_FILE, valida le colonne richieste (SAMPLE_ID_COL, TARGET_COL)."""
    if not os.path.exists(cfg.env_file):
        raise FileNotFoundError(f"ENV_FILE non trovato: {cfg.env_file}")

    df_env = pd.read_csv(cfg.env_file, sep=cfg.sep, decimal=cfg.decimal)

    required = [cfg.sample_id_col, cfg.target_col, cfg.exposure]
    missing = [c for c in required if c not in df_env.columns]
    if missing:
        raise ValueError(f"Colonne mancanti in {cfg.env_file}: {missing}. Disponibili: {list(df_env.columns)}")

    df_env[cfg.sample_id_col] = df_env[cfg.sample_id_col].astype(str)
    df_env[cfg.target_col] = pd.to_numeric(df_env[cfg.target_col], errors="coerce")
    df_env[cfg.exposure] = pd.to_numeric(df_env[cfg.exposure], errors="coerce")

    n_before = len(df_env)
    df_env = df_env.dropna(subset=[cfg.target_col, cfg.exposure]).drop_duplicates(cfg.sample_id_col)
    if len(df_env) < n_before:
        logger.warning("[ENV] Rimossi %d pazienti con %s/%s mancante o id duplicato", n_before - len(df_env), cfg.target_col, cfg.exposure)

    logger.info("[ENV] Caricati %d pazienti da %s", len(df_env), cfg.env_file)
    return df_env


def build_exposure_column(cfg: Config, df: pd.DataFrame) -> str:
    """Standardizza EXPOSURE come in build_dataset.py (StandardScaler se
    cfg.standardize=True) e ritorna il nome della colonna da usare come E."""
    if cfg.standardize:
        col = f"{cfg.exposure}_std"
        df[col] = StandardScaler().fit_transform(df[[cfg.exposure]])
        return col
    return cfg.exposure


def compute_pca_corrected_residuals(
    cfg: Config, gxe_cfg: GXEConfig, df_env: pd.DataFrame, gen_map: Optional[pd.DataFrame], logger,
) -> pd.DataFrame:
    """
    STEP A: per ciascuna generazione richiesta,
        onset_age  = [PCA_1..k, covariate additive] . delta + eps
        onset_age' = onset_age - [PCA, covariate] . delta_hat

    Le PCA (e le covariate additive come 'sex') sono usate SOLO qui. Non
    vengono mai passate a GENetLib.
    """
    generations = gxe_cfg.generations or [cfg.generation]

    if gen_map is None and len(generations) > 0:
        raise RuntimeError(
            "Impossibile eseguire Step A: nessuna mappa id->generazione disponibile "
            "(vedi resolve_generation_map) e le PCA sono specifiche per generazione."
        )

    all_residuals = []
    for generation in generations:
        ids_this_gen = set(gen_map.loc[gen_map["generation"] == generation, "id"])
        sub = df_env[df_env[cfg.sample_id_col].isin(ids_this_gen)].copy()
        if sub.empty:
            logger.warning("[STEP A] generazione=%s: nessun paziente trovato, salto", generation)
            continue

        sub, covariate_cols = _encode_covariates(sub, cfg.covariates)

        pca_df = None
        pc_cols = []
        if cfg.use_pca_covariates:
            pca_df = load_pca_covariates(cfg.pca_covariates_path_template, generation, cfg.pca_n_components)
            pc_cols = [c for c in pca_df.columns if c != PCA_ID_COLUMN]
            n_before = len(sub)
            sub = sub.merge(pca_df, left_on=cfg.sample_id_col, right_on=PCA_ID_COLUMN, how="inner")
            if len(sub) < n_before:
                logger.warning(
                    "[STEP A] generazione=%s: %d pazienti persi nel merge con le PCA (id non in comune)",
                    generation, n_before - len(sub),
                )

        design_cols = pc_cols + covariate_cols
        sub = sub.dropna(subset=[cfg.target_col] + design_cols)
        if sub.empty or not design_cols:
            logger.warning(
                "[STEP A] generazione=%s: dataset vuoto o nessuna covariata di correzione "
                "disponibile (PCA=%s, covariate=%s), salto", generation, cfg.use_pca_covariates, cfg.covariates,
            )
            continue

        X = sub[design_cols].to_numpy(dtype=float)
        y = sub[cfg.target_col].to_numpy(dtype=float)

        reg = LinearRegression(fit_intercept=True)
        reg.fit(X, y)
        residuals = y - reg.predict(X)
        r2 = reg.score(X, y)
        logger.info(
            "[STEP A] generazione=%s: regressione %s ~ %s su %d pazienti, R2=%.4f",
            generation, cfg.target_col, design_cols, len(sub), r2,
        )

        out = sub[[cfg.sample_id_col]].copy()
        out["generation"] = generation
        out[cfg.target_col] = y
        out[f"{cfg.target_col}_resid"] = residuals
        all_residuals.append(out)

    if not all_residuals:
        raise RuntimeError("STEP A: nessun residuo calcolato per nessuna generazione richiesta.")

    result = pd.concat(all_residuals, ignore_index=True)
    logger.info("[STEP A] Totale pazienti con fenotipo corretto (Y'): %d", len(result))
    return result


# ============================================================================
# 3. LETTURA GENOTIPO PER CROMOSOMA (gen.parquet, formato WIDE)
# ============================================================================

def get_variant_schema(raw_file: str) -> list:
    """Legge SOLO lo schema (nomi colonna) del parquet, senza caricare i
    dati: efficiente anche con ~1.3M colonne (legge solo il footer)."""
    return pq.ParquetFile(raw_file).schema_arrow.names


def _chrom_prefix_candidates(chrom: str) -> set:
    """I contig VCF sorgente possono essere nominati '1' o 'chr1' (vedi
    extract_matrix.resolve_chrom_name): qui accettiamo entrambe le
    convenzioni senza doverle conoscere a priori."""
    return {str(chrom), f"chr{chrom}", f"Chr{chrom}", f"CHR{chrom}"}


def select_columns_for_chromosome(all_columns: list, chrom: str, id_col: str = "id") -> list:
    """Seleziona, fra tutte le colonne del parquet, quelle relative alle
    varianti del cromosoma richiesto (formato nome colonna:
    '{CHROM}_{POS}_{REF}_{ALT}', vedi id_utils.build_variant_label)."""
    candidates = _chrom_prefix_candidates(chrom)
    variant_cols = [
        c for c in all_columns
        if c != id_col and c.split("_", 1)[0] in candidates
    ]
    return variant_cols


def load_genotype_matrix_for_chromosome(
    cfg: Config, chromosome: str, all_columns: list, patient_ids: set, logger,
) -> pd.DataFrame:
    """
    Legge da gen.parquet, via DuckDB, SOLO la colonna id + le colonne SNP
    del cromosoma richiesto (column pruning: DuckDB/parquet leggono da
    disco solo le colonne selezionate, non l'intero file da ~1.3M colonne).
    """
    variant_cols = select_columns_for_chromosome(all_columns, chromosome, cfg.sample_id_col)
    if not variant_cols:
        logger.warning("[chr%s] Nessuna colonna SNP trovata per questo cromosoma nel Parquet", chromosome)
        return pd.DataFrame()

    con = duckdb.connect(database=":memory:")
    con.execute("PRAGMA threads=4")
    try:
        cols_sql = ", ".join([f'"{cfg.sample_id_col}"'] + [f'"{c}"' for c in variant_cols])
        query = f"SELECT {cols_sql} FROM read_parquet('{cfg.raw_file}')"
        g_wide = con.execute(query).fetchdf()
    finally:
        con.close()

    g_wide[cfg.sample_id_col] = g_wide[cfg.sample_id_col].astype(str).map(clean_sample_id)
    g_wide = g_wide.drop_duplicates(cfg.sample_id_col).set_index(cfg.sample_id_col)

    # Filtra subito ai soli pazienti che ci servono (riduce il costo delle
    # operazioni successive, in particolare per cromosomi molto popolati)
    g_wide = g_wide[g_wide.index.isin(patient_ids)]

    logger.info(
        "[chr%s] Genotipo caricato: %d pazienti x %d SNP (colonne lette selettivamente dal Parquet)",
        chromosome, len(g_wide), len(variant_cols),
    )
    return g_wide


# ============================================================================
# 4. COSTRUZIONE DATASET (G, E, Y') PER IL MODELLO
# ============================================================================

def build_chromosome_dataset(
    cfg: Config, gxe_cfg: GXEConfig, chromosome: str,
    g_wide: pd.DataFrame, residual_df: pd.DataFrame, exposure_col: str, logger,
):
    if g_wide.empty:
        return None

    resid_col = f"{cfg.target_col}_resid"
    narrow = residual_df.set_index(cfg.sample_id_col)[[resid_col, exposure_col]]

    merged = g_wide.join(narrow, how="inner").dropna()
    if len(merged) < gxe_cfg.min_samples_required:
        logger.warning(
            "[chr%s] Solo %d pazienti disponibili dopo il merge (minimo richiesto: %d), salto",
            chromosome, len(merged), gxe_cfg.min_samples_required,
        )
        return None

    snp_cols = list(g_wide.columns)
    G_df = merged[snp_cols].astype(float)
    E_df = merged[[exposure_col]].astype(float)
    y_resid = merged[resid_col].astype(float).to_numpy()

    variances = G_df.var(axis=0)
    zero_var = variances[variances == 0].index.tolist()
    if zero_var:
        logger.info("[chr%s] Rimossi %d SNP monomorfici sul campione corrente", chromosome, len(zero_var))
        G_df = G_df.drop(columns=zero_var)

    if G_df.shape[1] == 0:
        logger.warning("[chr%s] Nessuno SNP polimorfico rimasto, salto", chromosome)
        return None

    logger.info(
        "[chr%s] Dataset finale: %d pazienti, %d SNP, E=%s",
        chromosome, len(merged), G_df.shape[1], list(E_df.columns),
    )
    return G_df, E_df, y_resid, list(G_df.columns), list(merged.index)


# ============================================================================
# 5. TRAINING GENetLib (scalar_ge) ED ESTRAZIONE COEFFICIENTI
# ============================================================================

def run_genetlib_scalar_ge(gxe_cfg: GXEConfig, chromosome: str, G_df: pd.DataFrame, E_df: pd.DataFrame, y_resid: np.ndarray, logger):
    """
    Y' = G*beta + (G x E)*theta ,  E = SOLO risaie_1500 (mai PCA/sex).

    NOTA: GENetLib e' una rete neurale (MCP + L2), non una regressione
    lineare classica: "coefficienti" = pesi del layer sparso della rete
    dopo il training (net.sparse1 = effetto G principale, net.sparse2 =
    interazione G x E). "lam"/"alpha"/"max_iter" del task originale
    corrispondono a lambda2/Lambda/num_epochs.
    """
    torch.manual_seed(0)
    G = G_df.to_numpy(dtype=float)
    E = E_df.to_numpy(dtype=float)
    y = y_resid.reshape(-1, 1)

    result = scalar_ge(
        y=y, G=G, E=E, ytype="Continuous",
        num_hidden_layers=gxe_cfg.num_hidden_layers,
        nodes_hidden_layer=gxe_cfg.nodes_hidden_layer,
        num_epochs=gxe_cfg.num_epochs,
        learning_rate1=gxe_cfg.learning_rate1,
        learning_rate2=gxe_cfg.learning_rate2,
        lambda1=gxe_cfg.lambda1,
        lambda2=gxe_cfg.lambda2,
        Lambda=gxe_cfg.Lambda,
        threshold=gxe_cfg.significance_threshold,
        split_type=gxe_cfg.split_type,
        ratio=gxe_cfg.ratio,
        important_feature=True,
        plot=False,
    )

    train_res, ifs_g_idx, ifs_ge_idx = result
    train_loss, eval_loss, train_r2, eval_r2, net = train_res
    if torch.is_tensor(train_loss):
        train_loss = train_loss.detach().cpu().numpy()
    if torch.is_tensor(eval_loss):
        eval_loss = eval_loss.detach().cpu().numpy()
    if torch.is_tensor(train_r2):
        train_r2 = float(train_r2.detach().cpu().numpy())
    if torch.is_tensor(eval_r2):
        eval_r2 = float(eval_r2.detach().cpu().numpy())

    logger.info(
        "[chr%s] Training completato: MSE_train=%.4f MSE_valid=%.4f R2_train=%.4f R2_valid=%.4f",
        chromosome, float(np.asarray(train_loss).reshape(-1)[0]), float(np.asarray(eval_loss).reshape(-1)[0]), train_r2, eval_r2,
    )

    n_snps = G.shape[1]
    n_env = E.shape[1]
    coef_main = net.sparse1.weight.data.detach().cpu().numpy().reshape(-1)
    coef_inter_matrix = net.sparse2.weight.data.detach().cpu().numpy().reshape(-1).reshape(n_env, n_snps)

    metrics = {
        "train_loss": float(np.asarray(train_loss).reshape(-1)[0]),
        "eval_loss": float(np.asarray(eval_loss).reshape(-1)[0]),
        "train_r2": float(train_r2),
        "eval_r2": float(eval_r2),
        "n_samples": int(G.shape[0]),
        "n_snps": int(n_snps),
        "n_env": int(n_env),
    }
    return {
        "coef_main": coef_main,
        "coef_inter_matrix": coef_inter_matrix,
        "important_snp_idx": set(ifs_g_idx),
        "important_interaction_idx": set(ifs_ge_idx),
        "metrics": metrics,
    }


def build_results_table(snp_names, env_names, coef_main, coef_inter_matrix, important_snp_idx, important_interaction_idx) -> pd.DataFrame:
    n_snps = len(snp_names)
    max_abs_inter = np.max(np.abs(coef_inter_matrix), axis=0)
    argmax_env = np.argmax(np.abs(coef_inter_matrix), axis=0)
    sig_inter_snp_idx = {idx % n_snps for idx in important_interaction_idx}

    rows = []
    for i, snp in enumerate(snp_names):
        sig_main = i in important_snp_idx
        sig_inter = i in sig_inter_snp_idx
        rows.append({
            "snp_id": snp,
            "coef_main": coef_main[i],
            "abs_coef_main": abs(coef_main[i]),
            "significant_main": sig_main,
            "max_abs_coef_interaction": max_abs_inter[i],
            "top_interacting_env": env_names[argmax_env[i]] if env_names else None,
            "significant_interaction": sig_inter,
            "significant_overall": sig_main or sig_inter,
        })
    return pd.DataFrame(rows).sort_values("abs_coef_main", ascending=False).reset_index(drop=True)


# ============================================================================
# 6. SALVATAGGIO OUTPUT PER CROMOSOMA
# ============================================================================

def save_chromosome_outputs(gxe_cfg: GXEConfig, chromosome: str, results_df, coef_main, coef_inter_matrix, snp_names, env_names, metrics, logger) -> Path:
    chrom_dir = Path(gxe_cfg.output_dir) / f"chr{chromosome}"
    chrom_dir.mkdir(parents=True, exist_ok=True)

    np.save(chrom_dir / "coef_main.npy", coef_main)
    np.save(chrom_dir / "coef_interaction_matrix.npy", coef_inter_matrix)
    np.save(chrom_dir / "snp_names.npy", np.array(snp_names))
    np.save(chrom_dir / "env_names.npy", np.array(env_names))
    results_df.to_csv(chrom_dir / "snp_results.csv", index=False)
    with open(chrom_dir / "metrics.json", "w") as f:
        json.dump(metrics, f, indent=2)

    if gxe_cfg.save_plots:
        try:
            top = results_df.reindex(results_df["abs_coef_main"].sort_values(ascending=False).index).head(25)
            fig, ax = plt.subplots(figsize=(9, 6))
            colors = ["#d62728" if s else "#1f77b4" for s in top["significant_overall"]]
            ax.barh(top["snp_id"].astype(str), top["coef_main"], color=colors)
            ax.set_xlabel("Coefficiente effetto principale G (peso sparse1)")
            ax.set_title(f"Chr{chromosome}: top 25 SNP per |coefficiente| (rosso = significativo)")
            ax.invert_yaxis()
            fig.tight_layout()
            fig.savefig(chrom_dir / "top_snp_coefficients.png", dpi=150)
            plt.close(fig)
        except Exception:
            logger.warning("[chr%s] Impossibile generare il grafico: %s", chromosome, traceback.format_exc())

    logger.info("[chr%s] Risultati salvati in %s", chromosome, chrom_dir)
    return chrom_dir


# ============================================================================
# 7. PIPELINE PER SINGOLO CROMOSOMA (worker parallelo)
# ============================================================================

def process_chromosome(args) -> dict:
    (chromosome, all_columns, residual_records, exposure_col, gxe_cfg_dict, cfg_dict) = args

    cfg = Config(**cfg_dict) if not isinstance(cfg_dict, Config) else cfg_dict
    gxe_cfg = GXEConfig(**gxe_cfg_dict) if not isinstance(gxe_cfg_dict, GXEConfig) else gxe_cfg_dict

    configure_logging(cfg.log_dir)
    logger = get_logger(f"{__name__}.chr{chromosome}")
    torch.set_num_threads(gxe_cfg.torch_threads_per_worker)

    residual_df = pd.DataFrame(residual_records)
    t0 = time.time()
    status = {"chromosome": chromosome, "status": "unknown", "error": None}

    try:
        logger.info("===== INIZIO elaborazione cromosoma %s =====", chromosome)
        patient_ids = set(residual_df[cfg.sample_id_col])

        g_wide = load_genotype_matrix_for_chromosome(cfg, chromosome, all_columns, patient_ids, logger)
        dataset = build_chromosome_dataset(cfg, gxe_cfg, chromosome, g_wide, residual_df, exposure_col, logger)
        if dataset is None:
            status["status"] = "skipped_insufficient_data"
            return status

        G_df, E_df, y_resid, snp_names, patient_ids_used = dataset
        model_out = run_genetlib_scalar_ge(gxe_cfg, chromosome, G_df, E_df, y_resid, logger)

        results_df = build_results_table(
            snp_names, list(E_df.columns), model_out["coef_main"], model_out["coef_inter_matrix"],
            model_out["important_snp_idx"], model_out["important_interaction_idx"],
        )
        n_sig = int(results_df["significant_overall"].sum())
        logger.info("[chr%s] SNP significativi: %d / %d", chromosome, n_sig, len(results_df))

        chrom_dir = save_chromosome_outputs(
            gxe_cfg, chromosome, results_df, model_out["coef_main"], model_out["coef_inter_matrix"],
            snp_names, list(E_df.columns), model_out["metrics"], logger,
        )

        status.update({
            "status": "success", "n_snps": len(snp_names), "n_samples": len(patient_ids_used),
            "n_significant": n_sig, "output_dir": str(chrom_dir), "metrics": model_out["metrics"],
            "elapsed_sec": round(time.time() - t0, 1),
        })
        logger.info("===== FINE cromosoma %s in %.1fs =====", chromosome, status["elapsed_sec"])
        return status

    except Exception as e:
        logger.error("[chr%s] ERRORE: %s\n%s", chromosome, e, traceback.format_exc())
        status.update({"status": "error", "error": str(e), "elapsed_sec": round(time.time() - t0, 1)})
        return status


# ============================================================================
# 8. REPORT WORD RIASSUNTIVO (opzionale)
# ============================================================================

def write_word_summary(gxe_cfg: GXEConfig, run_results: list, logger) -> Optional[Path]:
    if not gxe_cfg.save_word_summary:
        return None
    if not _HAS_DOCX:
        logger.warning("python-docx non installato: report Word non generato")
        return None

    doc = Document()
    doc.add_heading("Pipeline G x E (GENetLib) - Report riassuntivo", level=1)
    doc.add_paragraph(f"Generato il {datetime.now():%Y-%m-%d %H:%M:%S}")
    doc.add_paragraph(
        "PCA e covariate additive (sex) rimosse dal fenotipo in uno step di regressione "
        "preliminare (onset_age ~ PCA + sex); GENetLib stimato sui residui, con E = "
        "sola esposizione (nessuna interazione G x PCA/sex)."
    )
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(["Chr", "Stato", "N pazienti", "N SNP", "N sig.", "R2 valid", "Tempo (s)"]):
        table.rows[0].cells[i].text = h
    for r in run_results:
        row = table.add_row().cells
        row[0].text = str(r.get("chromosome", ""))
        row[1].text = str(r.get("status", ""))
        row[2].text = str(r.get("n_samples", "-"))
        row[3].text = str(r.get("n_snps", "-"))
        row[4].text = str(r.get("n_significant", "-"))
        m = r.get("metrics") or {}
        row[5].text = f"{m.get('eval_r2', float('nan')):.3f}" if m else "-"
        row[6].text = str(r.get("elapsed_sec", "-"))

    out_path = Path(gxe_cfg.output_dir) / "summary_report.docx"
    doc.save(out_path)
    logger.info("Report Word salvato in %s", out_path)
    return out_path


# ============================================================================
# 9. ORCHESTRAZIONE PRINCIPALE
# ============================================================================

def run_gxe_genetlib_pipeline(cfg: Optional[Config] = None, gxe_cfg: Optional[GXEConfig] = None) -> None:
    cfg = cfg or get_config()
    gxe_cfg = gxe_cfg or GXEConfig()

    configure_logging(cfg.log_dir)
    logger = get_logger(__name__)
    logger.info("========== AVVIO PIPELINE G x E (GENetLib) ==========")
    logger.info("RAW_FILE=%s ENV_FILE=%s EXPOSURE=%s TARGET=%s", cfg.raw_file, cfg.env_file, cfg.exposure, cfg.target_col)
    logger.info("Cromosomi: %s | Generazioni: %s", gxe_cfg.chromosomes, gxe_cfg.generations or [cfg.generation])

    for path, label in [(cfg.raw_file, "RAW_FILE"), (cfg.env_file, "ENV_FILE")]:
        if not os.path.exists(path):
            raise FileNotFoundError(f"{label} non trovato: {path}")

    all_columns = get_variant_schema(cfg.raw_file)
    if cfg.sample_id_col not in all_columns:
        raise ValueError(
            f"Colonna id campione '{cfg.sample_id_col}' non trovata nello schema di {cfg.raw_file}. "
            f"Colonne (prime 20 di {len(all_columns)}): {all_columns[:20]}"
        )
    logger.info("Schema Parquet letto: %d colonne totali (%d varianti attese)", len(all_columns), len(all_columns) - 1)

    df_env = load_environment_and_phenotype(cfg, logger)
    exposure_col = build_exposure_column(cfg, df_env)
    gen_map = resolve_generation_map(cfg)

    residual_df = compute_pca_corrected_residuals(cfg, gxe_cfg, df_env, gen_map, logger)
    # exposure_col e' stato calcolato su df_env (prima del merge con le PCA in
    # Step A): lo riportiamo dentro residual_df per il resto della pipeline.
    residual_df = residual_df.merge(df_env[[cfg.sample_id_col, exposure_col]], on=cfg.sample_id_col, how="inner")

    residual_records = residual_df.to_dict("records")

    from dataclasses import asdict
    cfg_dict = asdict(cfg) if not isinstance(cfg, dict) else cfg
    # DBConfig annidata non serve ai worker e puo' contenere credenziali:
    # non la propaghiamo ai processi figli.
    cfg_dict.pop("db", None)
    cfg_for_workers = Config(**{k: v for k, v in cfg_dict.items()})
    gxe_cfg_dict = asdict(gxe_cfg)

    tasks = [
        (chrom, all_columns, residual_records, exposure_col, gxe_cfg_dict, asdict(cfg_for_workers))
        for chrom in gxe_cfg.chromosomes
    ]

    logger.info("Avvio elaborazione parallela di %d cromosomi con %d worker", len(tasks), cfg.max_workers)
    run_results = []
    with ProcessPoolExecutor(max_workers=cfg.max_workers) as executor:
        futures = {executor.submit(process_chromosome, t): t[0] for t in tasks}
        iterator = as_completed(futures)
        if _HAS_TQDM:
            iterator = tqdm(iterator, total=len(futures), desc="Cromosomi processati")
        for future in iterator:
            chrom = futures[future]
            try:
                res = future.result()
            except Exception as e:
                logger.error("[chr%s] Eccezione non gestita nel worker: %s", chrom, e)
                res = {"chromosome": chrom, "status": "error", "error": str(e)}
            run_results.append(res)

    n_ok = sum(1 for r in run_results if r["status"] == "success")
    n_err = sum(1 for r in run_results if r["status"] == "error")
    n_skip = sum(1 for r in run_results if r["status"] == "skipped_insufficient_data")
    logger.info("===== PIPELINE COMPLETATA: %d ok, %d errori, %d saltati =====", n_ok, n_err, n_skip)

    os.makedirs(gxe_cfg.output_dir, exist_ok=True)
    pd.DataFrame(run_results).to_csv(Path(gxe_cfg.output_dir) / "run_summary.csv", index=False)
    write_word_summary(gxe_cfg, run_results, logger)


if __name__ == "__main__":
    run_gxe_genetlib_pipeline()

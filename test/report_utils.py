"""
Shared export utilities (CSV + Word) and automated (assert-based) checks
used by test_vqtl_pipeline.py (report per single se_method) and by
run_scenarios.py (recap per single scenario + aggregate report across all
scenarios).

Contents, in order:
  1) export_csv                                  -- generic CSV export
  2) CheckResult / CheckSuite / run_checks        -- PASS/WARN/FAIL checks
     on the vQTL pipeline outcome (used by test_vqtl_pipeline.py)
  3) export_docx                                  -- Word report for a
     single se_method of the vQTL pipeline (Step 3-7)
  4) build_recap / generate_recap                 -- cross ground_truth vs
     pipeline_results.csv (gene-environment part) -> summary + detail + docx
  5) generate_multi_scenario_recap                -- aggregation of the
     recaps of several scenarios into a single Word report
  6) build_scenario_comparison_table /
     write_manuscript_scenario_report              -- manuscript-ready
     scenario-by-scenario comparison table (G×E + vQTL metrics together),
     with an optional embedded figure (e.g. the isolated-test power curve)
"""
from __future__ import annotations

import json
import os
from dataclasses import dataclass, field
from datetime import datetime

import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CAUSAL_TYPES = {"gxe_meanshift", "pure_variance"}

COLOR_HEADER_BG = "2F5496"
COLOR_HEADER_TXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_CAUSAL_BG = "FFF2CC"


# ============================================================
# CSV export
# ============================================================

def export_csv(df: pd.DataFrame, tables_dir: str, name: str) -> str:
    os.makedirs(tables_dir, exist_ok=True)
    path = os.path.join(tables_dir, f"{name}.csv")
    df.to_csv(path, index=False)
    print(f"[export] {path}")
    return path


# ============================================================
# Automated checks (PASS / WARN / FAIL) -- not just printed for human
# review: this is where pass/fail is actually decided.
# ============================================================

@dataclass
class CheckResult:
    name: str
    level: str  # "PASS" | "WARN" | "FAIL"
    detail: str


@dataclass
class CheckSuite:
    results: list = field(default_factory=list)

    def add(self, name: str, ok: bool, detail: str, warn_only: bool = False) -> None:
        if ok:
            level = "PASS"
        elif warn_only:
            level = "WARN"
        else:
            level = "FAIL"
        self.results.append(CheckResult(name, level, detail))

    def skip(self, name: str, detail: str) -> None:
        self.results.append(CheckResult(name, "SKIP", detail))

    @property
    def has_failures(self) -> bool:
        return any(r.level == "FAIL" for r in self.results)

    def print_report(self) -> None:
        print("\n" + "-" * 78)
        print(f"{'CHECK':<45}{'OUTCOME':<8}DETAIL")
        print("-" * 78)
        for r in self.results:
            print(f"{r.name:<45}{r.level:<8}{r.detail}")
        print("-" * 78)
        n_fail = sum(1 for r in self.results if r.level == "FAIL")
        n_warn = sum(1 for r in self.results if r.level == "WARN")
        n_pass = sum(1 for r in self.results if r.level == "PASS")
        print(f"PASS={n_pass}  WARN={n_warn}  FAIL={n_fail}")

    def to_list(self) -> list[dict]:
        return [{"name": r.name, "level": r.level, "detail": r.detail} for r in self.results]


def run_checks(
    lambda_gc: float,
    all_causal: set,
    found_causal: set,
    candidates: pd.DataFrame,
    interaction_df_display: pd.DataFrame,
    perm_df_display: pd.DataFrame,
    alpha: float = 0.05,
) -> CheckSuite:
    suite = CheckSuite()

    # 1. lambda_GC reasonably close to 1. The asymptotic P is known to be
    #    anti-conservative (see the WARNING in filter_candidates), so a
    #    strong deviation is treated here as WARN, not FAIL: it is a signal
    #    to check in filter_candidates/scan, not a failure of the test
    #    itself.
    suite.add(
        "lambda_GC close to 1",
        0.8 <= lambda_gc <= 1.5,
        f"lambda_GC={lambda_gc:.3f} (expected in [0.8, 1.5])",
        warn_only=True,
    )

    # 2. At least one causal variant recovered among the candidates: if
    #    ZERO, this is a real failure (the scan/filter is not working at
    #    all).
    suite.add(
        "At least 1 causal variant among the candidates",
        len(found_causal) >= 1,
        f"{len(found_causal)}/{len(all_causal)} causal variants recovered: {sorted(found_causal) or '[]'}",
    )

    # 3. Majority of causal variants recovered (indicative threshold, not
    #    100%: this is a statistical test on synthetic data, not
    #    deterministic).
    frac = len(found_causal) / len(all_causal) if all_causal else 0.0
    suite.add(
        "Causal recovery >= 50%",
        frac >= 0.5,
        f"{frac:.0%} of the {len(all_causal)} causal variants recovered among the candidates",
        warn_only=True,
    )

    if candidates.empty or interaction_df_display.empty:
        suite.skip("Step 5+: interaction/permutation", "no candidates available, steps 5-7 not run")
        return suite

    gxe_rows = interaction_df_display[interaction_df_display["effect_type"] == "gxe_meanshift"]
    pv_rows = interaction_df_display[interaction_df_display["effect_type"] == "pure_variance"]

    # 4. Causal G×E variants among the candidates must have a significant
    #    interaction with a consistent sign.
    if gxe_rows.empty:
        suite.skip("G×E: significant interaction with consistent sign", "no G×E variant among the candidates in this run")
    else:
        sign_ok = (
            (gxe_rows["pval"] < alpha)
            & (pd_sign(gxe_rows["beta_I"]) == pd_sign(gxe_rows["true_beta_interaction"]))
        )
        n_ok = int(sign_ok.sum())
        suite.add(
            "G×E: significant interaction with consistent sign",
            n_ok == len(gxe_rows),
            f"{n_ok}/{len(gxe_rows)} G×E with pval<{alpha} and consistent sign",
        )

    # 5. Pure vQTL variants must NOT show a significant interaction (false
    #    positive of the interaction test).
    if pv_rows.empty:
        suite.skip("Pure vQTL: no interaction false positives", "no pure vQTL variant among the candidates in this run")
    else:
        n_falsepos = int((pv_rows["pval"] < alpha).sum())
        suite.add(
            "Pure vQTL: no interaction false positives",
            n_falsepos == 0,
            f"{n_falsepos}/{len(pv_rows)} pure vQTL with a falsely significant interaction",
        )

    # 6. Low empirical pval (Step 7) for causal G×E variants among the top
    #    permuted loci -- independent confirmation of Step 5.
    perm_gxe = perm_df_display[perm_df_display["effect_type"] == "gxe_meanshift"] if not perm_df_display.empty else perm_df_display
    if perm_gxe is None or perm_gxe.empty:
        suite.skip("G×E: low empirical_pval (Step 7)", "no G×E variant among the top permuted loci in this run")
    else:
        n_ok = int((perm_gxe["empirical_pval"] < alpha).sum())
        suite.add(
            "G×E: low empirical_pval (Step 7)",
            n_ok == len(perm_gxe),
            f"{n_ok}/{len(perm_gxe)} G×E with empirical_pval<{alpha}",
            warn_only=True,  # 500 permutations: limited power, not deterministic
        )

    # 7. Low Levene p-value for pure vQTL variants among the top permuted
    #    loci (heteroscedasticity signal even without interaction).
    perm_pv = perm_df_display[perm_df_display["effect_type"] == "pure_variance"] if not perm_df_display.empty else perm_df_display
    if perm_pv is None or perm_pv.empty:
        suite.skip("Pure vQTL: low levene_pval (Step 7)", "no pure vQTL variant among the top permuted loci in this run")
    else:
        n_ok = int((perm_pv["levene_pval"] < 0.1).sum())
        suite.add(
            "Pure vQTL: low levene_pval (Step 7)",
            n_ok == len(perm_pv),
            f"{n_ok}/{len(perm_pv)} pure vQTL with levene_pval<0.10",
            warn_only=True,
        )

    return suite


def pd_sign(series: pd.Series) -> pd.Series:
    return series.apply(np.sign)


# ============================================================
# Word export (python-docx)
# ============================================================

def _set_cell_shading(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_df_table(doc: Document, df: pd.DataFrame, highlight_col: str | None = "effect_type") -> None:
    if df.empty:
        doc.add_paragraph("(no data)").italic = True
        return
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = COLOR_HEADER_TXT
                r.font.size = Pt(9)
        _set_cell_shading(hdr_cells[i], COLOR_HEADER_BG)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        is_causal = highlight_col in df.columns and row.get(highlight_col) in CAUSAL_TYPES
        for i, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                text = f"{val:.3g}"
            else:
                text = str(val)
            cells[i].text = text
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if is_causal:
                _set_cell_shading(cells[i], COLOR_CAUSAL_BG)


def export_docx(
    out_path: str,
    generation: int,
    summary: dict,
    check_suite: CheckSuite,
    tables: list[tuple[str, str, pd.DataFrame]],
) -> str:
    """tables: list of (title, note, dataframe) in the desired order."""
    doc = Document()

    title = doc.add_heading(f"vQTL pipeline report — generation {generation}", level=0)

    # ---- Summary ----
    doc.add_heading("Summary", level=1)
    for line in [
        f"lambda_GC: {summary['lambda_gc']}",
        f"Causal variants recovered as candidates: {summary['n_found_causal']}/{summary['n_causal_total']} "
        f"({', '.join(summary['found_causal']) or 'none'})",
        f"False positives among the candidates: {summary['n_false_positives']}/{summary['n_null_truth']}",
        f"G×E with significant interaction (Step 5): {summary['n_gxe_sig']}/{summary['n_gxe_total']}",
        f"Pure vQTL with a falsely significant interaction (Step 5, expected 0): "
        f"{summary['n_pv_falsepos']}/{summary['n_pv_total']}",
    ]:
        doc.add_paragraph(line)

    # ---- Automated check outcome ----
    doc.add_heading("Automated check outcome", level=1)
    overall = "FAILED" if check_suite.has_failures else "PASSED"
    p = doc.add_paragraph()
    run = p.add_run(f"Overall outcome: {overall}")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00) if check_suite.has_failures else RGBColor(0x00, 0x80, 0x00)

    check_df = pd.DataFrame(check_suite.to_list())
    if not check_df.empty:
        check_df = check_df.rename(columns={"name": "check", "level": "outcome", "detail": "detail"})
        _add_df_table(doc, check_df, highlight_col=None)

    # ---- Per-step tables ----
    for title_text, note_text, df in tables:
        doc.add_heading(title_text, level=1)
        if note_text:
            note_p = doc.add_paragraph(note_text)
            note_p.runs[0].italic = True
            note_p.runs[0].font.size = Pt(9)
        _add_df_table(doc, df)
        doc.add_paragraph("")

    doc.save(out_path)
    print(f"[export] {out_path}")
    return out_path


# ============================================================
# Gene-environment recap: crosses ground_truth.csv (known truth, from
# gen_fake_data.py) with the REAL pipeline output (pipeline_results.csv,
# from modeling.process_single_variant) and always produces 3 files in the
# given folder:
#   - recap_summary.json  -> aggregate numbers (power by sign/magnitude,
#                             false positives on the null variants, etc.)
#   - recap_detail.csv    -> one row per variant, with outcome (TP/FN/FP/TN)
#   - recap_report.docx   -> same content in Word tables, readable without
#                             opening the CSV/JSON
# Called automatically at the end of every scenario/isolated variant by the
# orchestration scripts (run_scenarios.py, run_isolated_casual_test.py) --
# no need to run it by hand. Requires python-docx.
# ============================================================

MAGNITUDE_BINS = [0, 1, 2, 3, 4, 5, 6, 7, 8, 9, np.inf]
MAGNITUDE_LABELS = ["0-1", "1-2", "2-3", "3-4", "4-5", "5-6", "6-7", "7-8", "8-9", "9+"]


# ============================================================
# 1) Cross ground_truth / pipeline_results -> detail + summary
# ============================================================

def _classify_row(row: pd.Series, pvalue_threshold: float) -> dict:
    effect_type = row["effect_type"]
    true_inter = row.get("true_beta_interaction", 0.0)
    true_main = row.get("true_beta_main", 0.0)
    p_emp = row.get("p_emp", np.nan)
    obs_coef = row.get("obs_coef", np.nan)

    detected = bool(pd.notna(p_emp) and p_emp < pvalue_threshold)

    if effect_type == "gxe_meanshift" and true_inter != 0:
        category, sign = "gxe_interaction", ("pos" if true_inter > 0 else "neg")
        outcome = "TP (found)" if detected else "FN (missed)"
    elif effect_type == "gxe_meanshift" and true_inter == 0:
        # main effect only, true interaction = 0 -> false-positive control
        category, sign = "main_only_control", "zero"
        outcome = "FP (spurious interaction)" if detected else "TN (correct)"
    elif effect_type == "pure_variance":
        category, sign = "pure_variance", "n/a"
        outcome = "n/a (see vQTL step)"
    else:
        category, sign = "null", "zero"
        outcome = "FP (false positive)" if detected else "TN (correct)"

    return dict(
        category=category, sign=sign, magnitude=abs(true_inter), detected=detected,
        outcome=outcome, true_beta_interaction=true_inter, true_beta_main=true_main,
        obs_coef=obs_coef, p_emp=p_emp,
    )


def build_recap(ground_truth_df: pd.DataFrame, pipeline_results_df: pd.DataFrame,
                 pvalue_threshold: float = 0.05) -> tuple[pd.DataFrame, dict]:
    m = ground_truth_df.merge(pipeline_results_df, on="variant", how="left", suffixes=("", "_pr"))
    classified = m.apply(lambda r: _classify_row(r, pvalue_threshold), axis=1, result_type="expand")
    detail = pd.concat([m[["variant"]], classified], axis=1)
    detail["magnitude_bin"] = pd.cut(detail["magnitude"], bins=MAGNITUDE_BINS,
                                      labels=MAGNITUDE_LABELS, right=False)

    gxe = detail[detail.category == "gxe_interaction"]
    main_only = detail[detail.category == "main_only_control"]
    null_ = detail[detail.category == "null"]
    pv = detail[detail.category == "pure_variance"]

    def rate(df, mask=None):
        d = df if mask is None else df[mask]
        return None if len(d) == 0 else round(float(d["detected"].mean()), 4)

    by_sign = {
        s: dict(n=int((gxe.sign == s).sum()), n_detected=int(gxe[gxe.sign == s]["detected"].sum()),
                power=rate(gxe, gxe.sign == s))
        for s in ["pos", "neg"]
    }

    by_magnitude = {}
    for lab in MAGNITUDE_LABELS:
        sub = gxe[gxe.magnitude_bin == lab]
        if len(sub) == 0:
            continue
        by_magnitude[lab] = dict(
            n=int(len(sub)), n_detected=int(sub["detected"].sum()), power=rate(sub),
            n_pos=int((sub.sign == "pos").sum()), power_pos=rate(sub, sub.sign == "pos"),
            n_neg=int((sub.sign == "neg").sum()), power_neg=rate(sub, sub.sign == "neg"),
        )

    summary = dict(
        generated_at=datetime.now().isoformat(timespec="seconds"),
        pvalue_threshold=pvalue_threshold,
        n_variants_total=int(len(detail)),
        gxe_interaction=dict(
            n=int(len(gxe)), n_detected=int(gxe["detected"].sum()), power_overall=rate(gxe),
            by_sign=by_sign, by_magnitude=by_magnitude,
        ),
        main_only_control=dict(
            n=int(len(main_only)), n_false_interaction=int(main_only["detected"].sum()),
            false_positive_rate=rate(main_only),
        ),
        null_genomewide=dict(
            n=int(len(null_)), n_false_positive=int(null_["detected"].sum()),
            false_positive_rate=rate(null_),
        ),
        pure_variance=dict(
            n=int(len(pv)),
            note="Not tested by the G×E model (obs_coef/p_emp not applicable); "
                 "see the vQTL step (step3-7) for the outcome on these variants.",
        ),
    )
    return detail, summary


# ============================================================
# 2) JSON + CSV writing
# ============================================================

def write_json(summary: dict, path: str) -> None:
    with open(path, "w") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)


def write_csv(detail: pd.DataFrame, path: str) -> None:
    detail.to_csv(path, index=False)


# ============================================================
# 3) DOCX writing (python-docx, no dependency on Node/docx-js)
# ============================================================

def _add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if val is None else str(val)
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                row.cells[i].width = Cm(w)
    return table


def _pct(x) -> str:
    return "n/a" if x is None else f"{x * 100:.1f}%"


def write_docx(summary: dict, detail: pd.DataFrame, path: str, title: str = "Validation report — G×E pipeline") -> None:
    doc = Document()

    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {summary['generated_at']}    ").italic = True
    meta.add_run(f"P-value threshold: {summary['pvalue_threshold']}").italic = True

    # ---- General summary ----
    doc.add_heading("General summary", level=1)
    gxe, mo, nu, pv = (summary["gxe_interaction"], summary["main_only_control"],
                       summary["null_genomewide"], summary["pure_variance"])
    _add_table(
        doc,
        ["Category", "N", "N detected/FP", "Rate"],
        [
            ["Causal G×E variants (power)", gxe["n"], gxe["n_detected"], _pct(gxe["power_overall"])],
            ["Main-only control (expected 0 false positives)", mo["n"], mo["n_false_interaction"], _pct(mo["false_positive_rate"])],
            ["Genome-wide null variants (expected false positives ≈ threshold)", nu["n"], nu["n_false_positive"], _pct(nu["false_positive_rate"])],
            ["Pure variance (vQTL, not in the G×E model)", pv["n"], "-", "see vQTL step"],
        ],
        col_widths_cm=[8, 2, 3, 3],
    )

    # ---- Power by sign ----
    doc.add_heading("Power by interaction sign", level=1)
    rows = []
    for s, lab in [("pos", "Positive"), ("neg", "Negative")]:
        d = gxe["by_sign"][s]
        rows.append([lab, d["n"], d["n_detected"], _pct(d["power"])])
    _add_table(doc, ["Sign", "N", "N detected", "Power"], rows, col_widths_cm=[5, 3, 3, 3])

    # ---- Power by magnitude ----
    doc.add_heading("Power by |beta_interaction| magnitude", level=1)
    rows = []
    for lab, d in gxe["by_magnitude"].items():
        rows.append([lab, d["n"], _pct(d["power"]), d["n_pos"], _pct(d["power_pos"]), d["n_neg"], _pct(d["power_neg"])])
    _add_table(
        doc,
        ["|beta| bin", "N total", "Power total", "N pos", "Power pos", "N neg", "Power neg"],
        rows, col_widths_cm=[3, 2, 3, 2, 3, 2, 3],
    )

    # ---- Detail of causal G×E variants ----
    doc.add_heading("Detail of causal G×E variants", level=1)
    gxe_detail = detail[detail.category == "gxe_interaction"].sort_values("magnitude", ascending=False)
    rows = []
    for _, r in gxe_detail.iterrows():
        rows.append([
            r["variant"], f"{r['true_beta_interaction']:.2f}",
            "" if pd.isna(r["obs_coef"]) else f"{r['obs_coef']:.2f}",
            "" if pd.isna(r["p_emp"]) else f"{r['p_emp']:.3f}",
            r["outcome"],
        ])
    _add_table(doc, ["Variant", "true beta", "observed beta", "p_emp", "outcome"], rows,
               col_widths_cm=[4, 2.5, 2.5, 2, 4])

    # ---- Main-only false-positive control ----
    if mo["n"] > 0:
        doc.add_heading("False-positive control (main effect only, true interaction = 0)", level=1)
        mo_detail = detail[detail.category == "main_only_control"]
        rows = [[r["variant"], f"{r['true_beta_main']:.2f}",
                 "" if pd.isna(r["p_emp"]) else f"{r['p_emp']:.3f}", r["outcome"]]
                for _, r in mo_detail.iterrows()]
        _add_table(doc, ["Variant", "true beta_main", "p_emp", "outcome"], rows, col_widths_cm=[4, 3, 3, 5])

    # ---- False positives on null variants ----
    doc.add_heading("False positives on genome-wide null variants", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"{nu['n_false_positive']} / {nu['n']} null variants came out significant "
        f"(p_emp < {summary['pvalue_threshold']}) — rate {_pct(nu['false_positive_rate'])}."
    )
    fp_null = detail[(detail.category == "null") & (detail.detected)]
    if len(fp_null) > 0:
        rows = [[r["variant"], "" if pd.isna(r["obs_coef"]) else f"{r['obs_coef']:.2f}",
                 "" if pd.isna(r["p_emp"]) else f"{r['p_emp']:.3f}"]
                for _, r in fp_null.iterrows()]
        _add_table(doc, ["Variant", "observed beta", "p_emp"], rows, col_widths_cm=[5, 4, 4])

    for run in doc.paragraphs[0].runs:
        run.font.size = Pt(20)

    doc.save(path)


# ============================================================
# 4) "All-in-one" function, to be called at the end of the pipeline
# ============================================================

def generate_recap(ground_truth_path: str, pipeline_results_path: str, out_dir: str,
                    pvalue_threshold: float = 0.05) -> dict:
    """Reads the two CSVs, writes recap_summary.json / recap_detail.csv /
    recap_report.docx to out_dir, and returns the summary dict (useful for
    logging or for the multi-scenario summary report, see
    generate_multi_scenario_recap below). Also returns 'detail' inside the
    dict (private key "_detail") so run_scenarios.py does not have to
    re-read the CSV from disk for the aggregation."""
    os.makedirs(out_dir, exist_ok=True)
    gt = pd.read_csv(ground_truth_path)
    pr = pd.read_csv(pipeline_results_path)
    detail, summary = build_recap(gt, pr, pvalue_threshold=pvalue_threshold)

    write_json(summary, os.path.join(out_dir, "recap_summary.json"))
    write_csv(detail, os.path.join(out_dir, "recap_detail.csv"))
    write_docx(summary, detail, os.path.join(out_dir, "recap_report.docx"))

    summary["_detail"] = detail  # handy for aggregation, does not end up in the JSON (see write_json)
    return summary


# ============================================================
# 5) SINGLE final report across all scenarios
# ============================================================

def load_scenario_recap(recap_dir: str) -> tuple[pd.DataFrame, dict]:
    """Re-reads generate_recap()'s output from disk for a scenario (useful
    if the dict/detail is no longer in memory, e.g. separate runs)."""
    detail = pd.read_csv(os.path.join(recap_dir, "recap_detail.csv"))
    with open(os.path.join(recap_dir, "recap_summary.json")) as f:
        summary = json.load(f)
    return detail, summary


def _rate(df: pd.DataFrame, mask=None):
    d = df if mask is None else df[mask]
    return None if len(d) == 0 else round(float(d["detected"].mean()), 4)


def generate_multi_scenario_recap(scenario_summaries: dict[str, dict], out_dir: str,
                                   pvalue_threshold: float = 0.05,
                                   title: str = "Validation report — summary across all scenarios") -> dict:
    """
    scenario_summaries: {scenario_name: summary_dict}, where summary_dict is
    what generate_recap() returned for THAT scenario (must have the
    "_detail" key with the DataFrame -- if instead you are re-reading it
    from disk later, use load_scenario_recap() and build the
    {name: {**summary, "_detail": detail}} dict yourself before calling
    this function).

    Writes to out_dir:
      - all_scenarios_summary.json  (per-scenario + aggregated across all)
      - all_scenarios_detail.csv    (concatenated, with a 'scenario' column)
      - all_scenarios_report.docx   (comparison table + aggregate power)
    """
    os.makedirs(out_dir, exist_ok=True)

    combined_parts = []
    for name, summary in scenario_summaries.items():
        d = summary["_detail"].copy()
        d["scenario"] = name
        combined_parts.append(d)
    combined = pd.concat(combined_parts, ignore_index=True)

    gxe = combined[combined.category == "gxe_interaction"]
    null_ = combined[combined.category == "null"]
    mo = combined[combined.category == "main_only_control"]

    # ---- per-scenario comparison ----
    per_scenario_rows = []
    for name, summary in scenario_summaries.items():
        g = summary["gxe_interaction"]
        n = summary["null_genomewide"]
        per_scenario_rows.append(dict(
            scenario=name,
            n_causal=g["n"], power_overall=g["power_overall"],
            power_pos=g["by_sign"]["pos"]["power"], power_neg=g["by_sign"]["neg"]["power"],
            fp_rate_null=n["false_positive_rate"],
        ))

    # ---- aggregate across all scenarios together ----
    by_sign_agg = {
        s: dict(n=int((gxe.sign == s).sum()), n_detected=int(gxe[gxe.sign == s]["detected"].sum()),
                power=_rate(gxe, gxe.sign == s))
        for s in ["pos", "neg"]
    }
    by_magnitude_agg = {}
    for lab in MAGNITUDE_LABELS:
        sub = gxe[gxe.magnitude_bin == lab]
        if len(sub) == 0:
            continue
        by_magnitude_agg[lab] = dict(
            n=int(len(sub)), power=_rate(sub),
            n_pos=int((sub.sign == "pos").sum()), power_pos=_rate(sub, sub.sign == "pos"),
            n_neg=int((sub.sign == "neg").sum()), power_neg=_rate(sub, sub.sign == "neg"),
        )

    agg_summary = dict(
        generated_at=datetime.now().isoformat(timespec="seconds"),
        pvalue_threshold=pvalue_threshold,
        n_scenarios=len(scenario_summaries),
        per_scenario=per_scenario_rows,
        aggregate=dict(
            n_gxe_observations=int(len(gxe)),  # note: same variants replicated across scenarios, not independent
            power_overall=_rate(gxe),
            by_sign=by_sign_agg,
            by_magnitude=by_magnitude_agg,
            fp_rate_null_pooled=_rate(null_),
            fp_rate_main_only_pooled=_rate(mo) if len(mo) else None,
        ),
    )

    with open(os.path.join(out_dir, "all_scenarios_summary.json"), "w") as f:
        json.dump(agg_summary, f, indent=2, ensure_ascii=False)
    combined.to_csv(os.path.join(out_dir, "all_scenarios_detail.csv"), index=False)

    _write_multi_docx(agg_summary, os.path.join(out_dir, "all_scenarios_report.docx"), title=title)

    return agg_summary


def _write_multi_docx(agg_summary: dict, path: str, title: str) -> None:
    doc = Document()
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {agg_summary['generated_at']}    ").italic = True
    meta.add_run(f"Scenarios: {agg_summary['n_scenarios']}    ").italic = True
    meta.add_run(f"P-value threshold: {agg_summary['pvalue_threshold']}").italic = True

    doc.add_heading("Comparison across scenarios", level=1)
    rows = []
    for r in agg_summary["per_scenario"]:
        rows.append([
            r["scenario"], r["n_causal"], _pct(r["power_overall"]),
            _pct(r["power_pos"]), _pct(r["power_neg"]), _pct(r["fp_rate_null"]),
        ])
    _add_table(
        doc,
        ["Scenario", "N causal", "Power total", "Power pos", "Power neg", "FP rate (null)"],
        rows, col_widths_cm=[5, 2, 2.5, 2.5, 2.5, 3],
    )

    doc.add_heading("Aggregate power across all scenarios — by sign", level=1)
    a = agg_summary["aggregate"]
    rows = []
    for s, lab in [("pos", "Positive"), ("neg", "Negative")]:
        d = a["by_sign"][s]
        rows.append([lab, d["n"], d["n_detected"], _pct(d["power"])])
    _add_table(doc, ["Sign", "N", "N detected", "Power"], rows, col_widths_cm=[5, 3, 3, 3])

    doc.add_heading("Aggregate power across all scenarios — by magnitude", level=1)
    rows = []
    for lab, d in a["by_magnitude"].items():
        rows.append([lab, d["n"], _pct(d["power"]), d["n_pos"], _pct(d["power_pos"]), d["n_neg"], _pct(d["power_neg"])])
    _add_table(
        doc,
        ["|beta| bin", "N total", "Power total", "N pos", "Power pos", "N neg", "Power neg"],
        rows, col_widths_cm=[3, 2, 3, 2, 3, 2, 3],
    )

    p = doc.add_paragraph()
    p.add_run(
        f"Pooled false positives on null variants: {_pct(a['fp_rate_null_pooled'])}. "
        + (f"Pooled false positives on the main-only control: {_pct(a['fp_rate_main_only_pooled'])}."
           if a["fp_rate_main_only_pooled"] is not None else "")
    )
    note = doc.add_paragraph()
    note.add_run(
        "Note: 'N' in the aggregate power counts variant×scenario observations, "
        "not independent variants (if DEFAULT_CAUSAL_VARIANTS are the same in every "
        "scenario, with the same seed they are not independent replicates — see the "
        "analysis in the conversation)."
    ).italic = True

    for run in doc.paragraphs[0].runs:
        run.font.size = Pt(20)

    doc.save(path)


# ============================================================
# 6) Manuscript-ready scenario comparison table (G×E + vQTL metrics
# together) -- built from the list of per-scenario result dicts produced
# by run_scenarios.run_all_scenarios() (same schema as
# scenarios/all_scenarios_summary.json). This is the compact, one-row-per-
# scenario table meant to go directly into the manuscript/supplementary,
# as opposed to _write_multi_docx() above (G×E-only, more granular,
# intended for internal review of the recap/ folders).
# ============================================================

SCENARIO_LABELS_DEFAULT: dict[str, str] = {
    "baseline": "Baseline",
    "population_stratification": "Population stratification",
    "nonrandom_missing_carriers": "Non-random missingness",
    "small_sample": "Reduced sample",
    "high_zero_inflation_exposure": "Zero-inflated exposure",
    "high_sample": "Enlarged sample",
}


def build_scenario_comparison_table(all_results: list[dict],
                                     labels: dict[str, str] | None = None) -> pd.DataFrame:
    """Builds the compact, manuscript-ready scenario comparison table from
    `all_results` (the list returned by run_scenarios.run_all_scenarios()
    under the "all_results" key, i.e. the same content written to
    scenarios/all_scenarios_summary.json).

    One row per scenario, with:
      Scenario, G×E power, FP rate (G×E, genome-wide), λGC (vQTL),
      vQTL candidate recovery, Specificity FP (pure-variance)

    Scenarios with status != "ok" are still included, with "n/a" in the
    numeric columns, so a failed run is visible in the table rather than
    silently dropped.
    """
    labels = labels or SCENARIO_LABELS_DEFAULT
    rows = []
    for r in all_results:
        name = r.get("scenario", "?")
        label = labels.get(name, name)

        if r.get("status") != "ok":
            rows.append({
                "Scenario": label,
                "G×E power": "n/a",
                "FP rate (G×E, genome-wide)": "n/a",
                "λGC (vQTL)": "n/a",
                "vQTL candidate recovery": "n/a",
                "Specificity FP (pure-variance)": "n/a",
            })
            continue

        ge_recap = ((r.get("ge_interaction") or {}).get("recap") or {})
        ge = ge_recap.get("gxe_interaction", {}) or {}
        nf = ge_recap.get("null_genomewide", {}) or {}
        va = r.get("vqtl_asymptotic", {}) or {}

        ge_power = ge.get("power_overall")
        ge_n, ge_det = ge.get("n"), ge.get("n_detected")
        fp_rate, fp_n, fp_det = nf.get("false_positive_rate"), nf.get("n"), nf.get("n_false_positive")
        lam = va.get("lambda_gc")
        found, total = va.get("n_found_causal"), va.get("n_causal_total")
        pv_fp, pv_n = va.get("n_pv_falsepos"), va.get("n_pv_total")

        rows.append({
            "Scenario": label,
            "G×E power": (f"{ge_power * 100:.0f}% ({ge_det}/{ge_n})" if ge_power is not None else "n/a"),
            "FP rate (G×E, genome-wide)": (f"{fp_rate * 100:.1f}% ({fp_det}/{fp_n})" if fp_rate is not None else "n/a"),
            "λGC (vQTL)": (f"{lam:.3f}" if lam is not None else "n/a"),
            "vQTL candidate recovery": (f"{found}/{total} ({found / total * 100:.0f}%)" if total else "n/a"),
            "Specificity FP (pure-variance)": (f"{pv_fp}/{pv_n}" if pv_n is not None else "n/a"),
        })
    return pd.DataFrame(rows)


def write_manuscript_scenario_report(
    table_df: pd.DataFrame,
    out_path: str,
    figure_path: str | None = None,
    table_title: str = "Supplementary Table — Pipeline robustness across simulated scenarios",
    table_caption: str = (
        "Each row corresponds to a full re-run of the analytical pipeline on an "
        "independently simulated dataset. Specificity FP reports the number of "
        "pure-variance control variants incorrectly flagged as showing an "
        "interaction effect."
    ),
    figure_title: str = "Supplementary Figure — Detection power as a function of simulated effect size",
    figure_caption: str = (
        "Power to detect simulated effects in single-variant isolation tests, shown "
        "separately for the G×E interaction test and the vQTL pure-variance test."
    ),
    figure_width_dxa: int = 600,
) -> str:
    """Writes a single manuscript-ready docx containing the scenario
    comparison table (see build_scenario_comparison_table) and, optionally,
    an embedded figure (e.g. a power-curve PNG produced by
    run_isolated_casual_test.generate_manuscript_power_curve). Meant to be
    the one file you actually paste into the manuscript/supplementary,
    as opposed to the more granular per-scenario/isolated reports."""
    doc = Document()

    doc.add_heading(table_title, level=1)
    if table_caption:
        cap = doc.add_paragraph(table_caption)
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)

    headers = list(table_df.columns)
    rows = table_df.astype(str).values.tolist()
    _add_table(doc, headers, rows)

    if figure_path and os.path.exists(figure_path):
        doc.add_paragraph("")
        doc.add_heading(figure_title, level=1)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(figure_path, width=Pt(figure_width_dxa))
        if figure_caption:
            cap = doc.add_paragraph(figure_caption)
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)

    doc.save(out_path)
    print(f"[export] {out_path}")
    return out_path

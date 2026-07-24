"""
Isolated power test: EVERY variant (from either default pool) is tested with
BOTH the G×E mechanism and the vQTL (pure_variance) mechanism, one variant at
a time against an otherwise-null panel.

WHY BOTH TESTS FOR EVERY VARIANT
---------------------------------
For every variant label in the union of DEFAULT_CAUSAL_VARIANTS and
DEFAULT_PURE_VARIANCE_VARIANTS, we generate ONE dataset in which that single
variant is simultaneously declared causal-for-G×E (real or null beta) and
causal-for-vQTL (real or flat-null sd_by_dosage). This isolates exactly the
contribution of the single variant, on both mechanisms simultaneously, in a
single dataset. Both run_ge_interaction() and run_vqtl_debug() are then run
on that same dataset, and the label is looked up in both result tables.

============================================================================
WHY "G×E-only variant significant on vQTL" IS NOT NECESSARILY A FALSE
POSITIVE (read before touching PVALUE_THRESHOLD / the zone logic below)
============================================================================
gen_fake_data.py builds the phenotype for a causal_variants entry as:

    onset_age += beta_main * dosage_bin + beta_inter * dosage_bin * exposure_std_approx

dosage_bin is binary (carrier/non-carrier) and exposure_std_approx is a
per-subject random variable with variance ~1. This means:

    Var(onset_age | non-carrier) ~= noise_sd^2
    Var(onset_age | carrier)     ~= noise_sd^2 + beta_inter^2

i.e. a "pure" G×E variant with no declared vQTL effect STILL mechanically
induces real heteroscedasticity by dosage, proportional to beta_inter^2. The
vQTL test correctly detects this when beta_inter is large -- that is NOT a
specificity failure of the vQTL test, it is the vQTL test doing its job on
a phenotype that genuinely has unequal variance across dosage groups. It IS,
however, a sign that "is this variant declared as G×E-only" is the wrong
ground-truth label to test vQTL specificity against.

The predicted variance ratio induced by a given beta_inter is:

    predicted_var_ratio = 1 + (beta_inter / noise_sd)^2

We use this every run to split G×E-only variants into three zones instead
of a single "expect not significant" bucket:
  - "pure_null"          : predicted_var_ratio is small -> a real vQTL
                            specificity test (expect NOT significant).
  - "confounded_expected": predicted_var_ratio is large enough that a
                            vQTL hit is the MECHANICALLY EXPECTED outcome,
                            not an error (expect significant, and if it
                            IS significant that's counted as a correct
                            outcome, not a false positive).
  - "borderline"          : in between -- excluded from the pass/fail
                            counts entirely (neither zone's expectation is
                            reliable there), reported separately.

The two thresholds (ratio_low, ratio_high) are NOT hard-coded: they are
re-fit EVERY RUN from the data itself (logistic regression of
"significant_vqtl" on log(predicted_var_ratio) over the G×E-only variants),
because they depend on run-specific parameters (noise_sd, sample size,
vQTL test power at that n, ...) that can change between runs. See
fit_vqtl_confound_thresholds().

DESIGN, CACHING, DISK-BASED AGGREGATION: unchanged from the previous version
of this script -- see individual function docstrings below.

OUTPUT: isolated/<variant_label>/... (data + raw results) and a single combined
recap isolated/isolated_power_curve.csv + .json (now including
predicted_var_ratio, vqtl_zone, expected_sig_vqtl, match_vqtl, expected_sig_ge,
match_ge, as_expected for every variant), plus isolated/plots/*.png and
isolated/isolated_causal_test_report.docx (all in English).

============================================================================
SINGLE ENTRY POINT FOR THE WHOLE TEST BATTERY
============================================================================
Running this file runs THREE phases in sequence (each one individually
skippable, see --skip-* below):

  1) isolated   -- everything described above: every variant, both
                    mechanisms, one at a time. Answers "does the pipeline
                    recover each variant's effect at the right magnitude,
                    with the right sign, and no crosstalk between the two
                    mechanisms?".
  2) scenarios  -- the small-set robustness battery from run_scenarios.py
                    (population stratification, missingness, small/high
                    sample size, zero-inflated exposure, ...). Answers
                    "does the pipeline still work under conditions other
                    than the clean default one?". Reuses run_scenarios.py
                    as-is (run_all_scenarios()), not duplicated here.
  3) pure_null  -- exact binomial confidence interval on the false-positive
                    rate observed in the isolated run's "pure_null" vQTL
                    zone (phase 1's cleanest false-positive check), and --
                    only if the nominal 5% falls outside (or close to the
                    edge of) that CI -- a second check using the bootstrap
                    p-value (P_boot, K=200) already computed for every
                    variant by run_vqtl_debug, to tell apart "estimation
                    noise at K=50" from "a real bias in the variance test
                    itself". Needs phase 1 to have produced
                    isolated_power_curve.csv at least once (from this run,
                    or a previous one on disk).

HOW TO RUN IT (run from YOU, not from this chat, same folder as
run_scenarios.py, gen_fake_data.py, fake_vqtl_repository.py, report_utils.py
and test_vqtl_pipeline.py):
    python run_isolated_casual_test.py                     # all 3 phases
    python run_isolated_casual_test.py --workers 4 --output-dir         # parallel within each phase
    python run_isolated_casual_test.py --force              # ignore isolated_summary.json cache
    python run_isolated_casual_test.py --skip-scenarios     # phases 1 + 3 only
    python run_isolated_casual_test.py --skip-isolated --skip-pure-null   # phase 2 only
    python run_isolated_casual_test.py --scenarios baseline small_sample  # phase 2, subset only
"""
from __future__ import annotations

import argparse
import inspect
import json
import os
import sys
import time
import traceback
from concurrent.futures import ProcessPoolExecutor, as_completed
from datetime import datetime

import numpy as np
import pandas as pd
import statsmodels.api as sm
from scipy.stats import binomtest
import matplotlib
matplotlib.use("Agg")  # headless: no display required
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)

# Direct reuse of the logic already written in run_scenarios.py -- no
# duplication: same run_ge_interaction/run_vqtl_debug/_set_common_env used
# there.
import run_scenarios as rs
from gen_fake_data import (
    generate_dataset,
    DEFAULT_CAUSAL_VARIANTS,
    DEFAULT_PURE_VARIANCE_VARIANTS,
)

ISOLATED_ROOT = os.path.join(SCRIPT_DIR, "isolated")
PLOTS_DIR = os.path.join(ISOLATED_ROOT, "plots")
PVALUE_THRESHOLD = 0.05

# noise_sd actually used to generate the phenotype -- read from
# generate_dataset()'s own default instead of hard-coding 8.5 here, so this
# stays correct automatically if the default in gen_fake_data.py ever
# changes. We ALSO pass it explicitly to every generate_dataset() call below
# (instead of relying on the implicit default) and store it in each
# variant's result, so predicted_var_ratio is always computed with the
# value that was actually used to generate that variant's data, even if
# someone changes the default between runs for a subset of cached variants.
DEFAULT_NOISE_SD = inspect.signature(generate_dataset).parameters["noise_sd"].default

# Fallback "flat null" sd_by_dosage used for variants that are NOT in
# DEFAULT_PURE_VARIANCE_VARIANTS, so they can still be explicitly declared
# causal-for-vQTL-with-null-effect and therefore still appear in the vQTL
# comparison table. Derived from an existing pool entry if possible so the
# dosage keys match; otherwise falls back to {0,1,2}: 1.0.
def _derive_flat_null_sd() -> dict:
    if DEFAULT_PURE_VARIANCE_VARIANTS:
        sample = next(iter(DEFAULT_PURE_VARIANCE_VARIANTS.values()))
        dosages = list(sample.keys())
    else:
        dosages = [0, 1, 2]
    return {d: 1.0 for d in dosages}


FLAT_NULL_SD = _derive_flat_null_sd()
NULL_GE_BETAS = (0.0, 0.0)


def section(title: str) -> None:
    print("\n" + "#" * 88)
    print(title)
    print("#" * 88)


def build_variant_plan() -> list[dict]:
    """Union of both default pools -> one plan entry per unique label, with
    both mechanisms' parameters (real or null) and both roles."""
    labels = sorted(set(DEFAULT_CAUSAL_VARIANTS) | set(DEFAULT_PURE_VARIANCE_VARIANTS))
    plan = []
    for label in labels:
        in_ge_pool = label in DEFAULT_CAUSAL_VARIANTS
        in_vqtl_pool = label in DEFAULT_PURE_VARIANCE_VARIANTS

        beta_inter, beta_main = DEFAULT_CAUSAL_VARIANTS.get(label, NULL_GE_BETAS)
        sd_by_dosage = DEFAULT_PURE_VARIANCE_VARIANTS.get(label, FLAT_NULL_SD)

        if not in_ge_pool:
            ge_role = "crosstalk_check"
        elif beta_inter == 0.0:
            ge_role = "fpr_control"
        else:
            ge_role = "power"

        vqtl_role = "power" if in_vqtl_pool else "crosstalk_check"

        plan.append({
            "variant": label,
            "in_ge_pool": in_ge_pool,
            "in_vqtl_pool": in_vqtl_pool,
            "beta_inter": beta_inter,
            "beta_main": beta_main,
            "ge_role": ge_role,
            "sd_by_dosage": sd_by_dosage,
            "vqtl_role": vqtl_role,
        })
    return plan


# ============================================================
# Cache: if a variant already has an isolated_summary.json with status "ok",
# skip re-running it -- useful to resume after an error/interruption without
# redoing every variant. Bypassable with --force.
# ============================================================

def _summary_path(label: str) -> str:
    return os.path.join(ISOLATED_ROOT, label, "isolated_summary.json")


def _load_cached_result(var_dir: str, force: bool = False) -> dict | None:
    if force:
        return None
    summary_path = os.path.join(var_dir, "isolated_summary.json")
    if not os.path.isfile(summary_path):
        return None
    try:
        with open(summary_path) as f:
            cached = json.load(f)
    except (json.JSONDecodeError, OSError):
        return None  # corrupted/incomplete file: recompute
    if cached.get("status") != "ok":
        return None  # a previously failed run should always be retried
    return cached


# ============================================================
# One variant at a time, both mechanisms
# ============================================================

def run_isolated_variant(plan_entry: dict, n_workers: int = 1, force: bool = False) -> dict:
    label = plan_entry["variant"]
    var_dir = os.path.join(ISOLATED_ROOT, label)
    fake_dir = os.path.join(var_dir, "fake_data")

    cached = _load_cached_result(var_dir, force=force)
    if cached is not None:
        print(f"[isolated] {label}: cached result found (status ok), skipping recompute.")
        return cached

    section(
        f"[isolated] {label}  "
        f"(GxE: beta_inter={plan_entry['beta_inter']}, beta_main={plan_entry['beta_main']}, "
        f"role={plan_entry['ge_role']} | vQTL: sd_by_dosage={plan_entry['sd_by_dosage']}, "
        f"role={plan_entry['vqtl_role']})"
    )
    os.makedirs(fake_dir, exist_ok=True)

    result: dict = {
        "variant": label,
        "in_ge_pool": plan_entry["in_ge_pool"],
        "in_vqtl_pool": plan_entry["in_vqtl_pool"],
        "beta_inter": plan_entry["beta_inter"],
        "beta_main": plan_entry["beta_main"],
        "ge_role": plan_entry["ge_role"],
        "sd_by_dosage": plan_entry["sd_by_dosage"],
        "vqtl_role": plan_entry["vqtl_role"],
        "noise_sd_used": DEFAULT_NOISE_SD,
        "status": "ok", "error": None,
    }
    try:
        generate_dataset(
            out_dir=fake_dir, verbose=True, noise_sd=DEFAULT_NOISE_SD,
            causal_variants={label: (plan_entry["beta_inter"], plan_entry["beta_main"])},
            pure_variance_variants={label: plan_entry["sd_by_dosage"]},
        )
        vqtl_n_jobs = max(1, (os.cpu_count() or 2) // max(1, n_workers))
        rs._set_common_env(fake_dir, var_dir, vqtl_n_jobs=vqtl_n_jobs)

        # ---- G×E side ----
        ge_res = rs.run_ge_interaction(fake_dir, var_dir)
        result["ge_result"] = ge_res
        res_df = pd.read_csv(ge_res["results_csv"])
        row = res_df.loc[res_df["variant"] == label]
        if row.empty:
            result["ge_found_in_results"] = False
            result["p_ge"] = None
            result["significant_ge"] = False
        else:
            p_emp = float(row["p_emp"].iloc[0]) if "p_emp" in row.columns else None
            result["ge_found_in_results"] = True
            result["p_ge"] = p_emp
            result["significant_ge"] = (p_emp is not None) and (p_emp < PVALUE_THRESHOLD)

        if not result["ge_found_in_results"]:
            result["recovered_ge"] = False
        elif plan_entry["ge_role"] == "power":
            result["recovered_ge"] = result["significant_ge"]
        else:  # fpr_control or crosstalk_check: expect NOT significant
            result["recovered_ge"] = not result["significant_ge"]

        # ---- vQTL side ----
        debug_res = rs.run_vqtl_debug(fake_dir, var_dir)
        result["vqtl_debug"] = debug_res
        comparison = pd.read_csv(debug_res["comparison_csv"])
        vrow = comparison.loc[comparison["SNP"] == label]
        if vrow.empty:
            result["vqtl_found_in_results"] = False
            result["p_vqtl"] = None
            result["significant_vqtl"] = False
            result["p_vqtl_boot"] = None
            result["significant_vqtl_boot"] = False
        else:
            p_asym = float(vrow["P_asym"].iloc[0]) if "P_asym" in vrow.columns else None
            p_boot = float(vrow["P_boot"].iloc[0]) if "P_boot" in vrow.columns else None
            result["vqtl_found_in_results"] = True
            result["p_vqtl"] = p_asym
            result["significant_vqtl"] = (p_asym is not None) and (p_asym < PVALUE_THRESHOLD)
            result["p_vqtl_boot"] = p_boot
            result["significant_vqtl_boot"] = (p_boot is not None) and (p_boot < PVALUE_THRESHOLD)

        if not result["vqtl_found_in_results"]:
            result["recovered_vqtl"] = False
        elif plan_entry["vqtl_role"] == "power":
            result["recovered_vqtl"] = result["significant_vqtl"]
        else:  # crosstalk_check: expect NOT significant (naive label -- refined
               # per-run into pure_null/confounded_expected/borderline, see
               # classify_vqtl_crosstalk_zones() below; recovered_vqtl here is
               # kept only as the raw naive flag, for backward compatibility)
            result["recovered_vqtl"] = not result["significant_vqtl"]

        # predicted_var_ratio can be computed right away (only needs
        # beta_inter and the noise_sd actually used for THIS variant), no
        # aggregate/other-variant information required.
        result["predicted_var_ratio"] = 1.0 + (plan_entry["beta_inter"] / DEFAULT_NOISE_SD) ** 2

    except Exception as exc:
        result["status"] = "FAILED"
        result["error"] = f"{type(exc).__name__}: {exc}"
        result["traceback"] = traceback.format_exc()
        print(f"\n*** VARIANT '{label}' FAILED: {result['error']} ***")
        traceback.print_exc()

    with open(os.path.join(var_dir, "isolated_summary.json"), "w") as f:
        json.dump(result, f, indent=2, default=str)
    return result


def _worker(plan_entry: dict, n_workers: int, force: bool) -> dict:
    return run_isolated_variant(plan_entry, n_workers=n_workers, force=force)


# ============================================================
# Final aggregation, ALWAYS read back from disk.
#
# Individual variants may have been skipped this run (cache hit) or computed
# by a different process (--workers) or in a previous invocation of the
# script entirely. Rather than trusting the in-memory list accumulated
# during this particular run, we re-read every isolated_summary.json from
# disk for the variants in the current plan. This is the single source of
# truth for the combined recap, and makes it correct even if some variants
# were skipped this run (cache hit), the process crashed/was interrupted
# partway through a previous run, or variants were computed across multiple
# separate invocations.
# ============================================================

def collect_results_from_disk(plan: list[dict]) -> pd.DataFrame:
    rows = []
    missing = []
    for entry in plan:
        label = entry["variant"]
        summary_path = _summary_path(label)
        if not os.path.isfile(summary_path):
            missing.append(label)
            continue
        try:
            with open(summary_path) as f:
                rows.append(json.load(f))
        except (json.JSONDecodeError, OSError) as exc:
            print(f"[warn] could not read {summary_path}: {exc}")
            missing.append(label)

    if missing:
        print(
            f"\n[warn] {len(missing)} variant(s) have no readable isolated_summary.json on disk "
            f"and are excluded from the combined recap: {missing}"
        )

    df = pd.DataFrame(rows)
    if not df.empty and "predicted_var_ratio" not in df.columns:
        # backward compatibility with isolated_summary.json files written by
        # the previous version of this script (no predicted_var_ratio saved)
        df["predicted_var_ratio"] = np.nan
    if not df.empty and "noise_sd_used" not in df.columns:
        df["noise_sd_used"] = np.nan
    # recompute defensively for any row missing it (cached from an older run,
    # or noise_sd_used present but predicted_var_ratio wasn't saved)
    need_fill = df["predicted_var_ratio"].isna() & df["noise_sd_used"].notna() & df["beta_inter"].notna()
    if need_fill.any():
        df.loc[need_fill, "predicted_var_ratio"] = 1.0 + (
            df.loc[need_fill, "beta_inter"] / df.loc[need_fill, "noise_sd_used"]
        ) ** 2
    return df


# ============================================================
# G×E-confound-aware vQTL specificity classification.
#
# Re-fit EVERY run from the data itself -- see module docstring for why a
# fixed threshold is wrong (it depends on noise_sd, sample size, vQTL test
# power at that n, etc., all of which can change between runs).
# ============================================================

FALLBACK_RATIO_LOW = 1.05
FALLBACK_RATIO_HIGH = 1.35


def fit_vqtl_confound_thresholds(df: pd.DataFrame) -> dict:
    """Logistic regression of significant_vqtl ~ log(predicted_var_ratio)
    over the G×E-only variants (vqtl_role == 'crosstalk_check'), used to
    find:
      - ratio_low: predicted_var_ratio at which P(significant) is back down
        to the empirical baseline rate (observed on the beta_inter=0
        controls, i.e. ratio==1) -- below this, a hit is genuine noise.
      - ratio_high: predicted_var_ratio at which P(significant) reaches 90%
        -- above this, a hit is the mechanically expected outcome.
    Falls back to fixed constants (with a printed warning) if the fit fails
    or isn't well-behaved (e.g. too few points, no variation in the outcome,
    non-monotonic/negative slope)."""
    sub = df[
        (df["vqtl_role"] == "crosstalk_check")
        & (df["status"] == "ok")
        & df["p_vqtl"].notna()
        & df["predicted_var_ratio"].notna()
    ].copy()

    result = {
        "method": "fallback_fixed",
        "ratio_low": FALLBACK_RATIO_LOW,
        "ratio_high": FALLBACK_RATIO_HIGH,
        "baseline_rate": None,
        "n_points": len(sub),
        "note": None,
    }

    if len(sub) < 8:
        result["note"] = f"Only {len(sub)} G×E-only variants available; using fallback fixed thresholds."
        print(f"[warn] {result['note']}")
        return result

    sub["log_ratio"] = np.log(sub["predicted_var_ratio"].clip(lower=1.0 + 1e-9))
    y = sub["significant_vqtl"].astype(float).to_numpy()

    baseline_mask = sub["predicted_var_ratio"] <= (sub["predicted_var_ratio"].min() + 1e-6)
    baseline_rate = float(y[baseline_mask.to_numpy()].mean()) if baseline_mask.any() else float(y.mean())
    baseline_rate = min(max(baseline_rate, 0.01), 0.20)  # keep it sane: between 1% and 20%
    result["baseline_rate"] = baseline_rate

    if y.min() == y.max():
        result["note"] = "No variation in significant_vqtl outcome across G×E-only variants; using fallback thresholds."
        print(f"[warn] {result['note']}")
        return result

    try:
        X = sm.add_constant(sub["log_ratio"].to_numpy())
        model = sm.Logit(y, X).fit(disp=0)
        b0, b1 = model.params
        if not np.isfinite(b0) or not np.isfinite(b1) or b1 <= 0:
            result["note"] = f"Logistic fit degenerate or non-increasing (b1={b1:.4g}); using fallback thresholds."
            print(f"[warn] {result['note']}")
            return result

        def _logit(p: float) -> float:
            p = min(max(p, 1e-6), 1 - 1e-6)
            return np.log(p / (1 - p))

        target_high = 0.90
        x_low = (_logit(baseline_rate) - b0) / b1
        x_high = (_logit(target_high) - b0) / b1

        ratio_low = max(1.0, float(np.exp(x_low)))
        ratio_high = float(np.exp(x_high))
        if not np.isfinite(ratio_low) or not np.isfinite(ratio_high) or ratio_high <= ratio_low:
            result["note"] = "Fitted thresholds not well-ordered/finite; using fallback thresholds."
            print(f"[warn] {result['note']}")
            return result

        result.update({
            "method": "logistic_fit",
            "ratio_low": ratio_low,
            "ratio_high": ratio_high,
            "b0": float(b0), "b1": float(b1),
            "note": None,
        })
        print(
            f"[thresholds] Fitted from data: baseline_rate={baseline_rate:.3f}, "
            f"ratio_low={ratio_low:.3f}, ratio_high={ratio_high:.3f} "
            f"(logistic b0={b0:.3f}, b1={b1:.3f}, n={len(sub)})"
        )
        return result
    except Exception as exc:
        result["note"] = f"Logistic fit raised {type(exc).__name__}: {exc}; using fallback thresholds."
        print(f"[warn] {result['note']}")
        return result


def add_expectation_columns(df: pd.DataFrame, thresholds: dict) -> pd.DataFrame:
    """Adds, for every variant:
      expected_sig_ge, match_ge   (unchanged logic: G×E test has no known
                                    confound, so 'power' -> expect significant,
                                    everything else -> expect not significant)
      vqtl_zone                   ('power' | 'pure_null' | 'confounded_expected'
                                    | 'borderline')
      expected_sig_vqtl           (True / False / NaN for 'borderline')
      match_vqtl                  (bool, or NaN when vqtl_zone == 'borderline')
      as_expected                 (match_ge AND match_vqtl, NaN if match_vqtl
                                    is NaN -- i.e. undetermined, not "failed")
    """
    df = df.copy()

    df["expected_sig_ge"] = df["ge_role"] == "power"
    df["match_ge"] = df["significant_ge"] == df["expected_sig_ge"]

    ratio_low, ratio_high = thresholds["ratio_low"], thresholds["ratio_high"]

    def _zone(row):
        if row["vqtl_role"] == "power":
            return "power"
        r = row.get("predicted_var_ratio")
        if r is None or pd.isna(r):
            return "borderline"
        if r <= ratio_low:
            return "pure_null"
        if r >= ratio_high:
            return "confounded_expected"
        return "borderline"

    df["vqtl_zone"] = df.apply(_zone, axis=1)

    def _expected_sig_vqtl(row):
        if row["vqtl_zone"] == "power":
            return True
        if row["vqtl_zone"] == "pure_null":
            return False
        if row["vqtl_zone"] == "confounded_expected":
            return True
        return np.nan  # borderline: undetermined by design

    df["expected_sig_vqtl"] = df.apply(_expected_sig_vqtl, axis=1)

    def _match_vqtl(row):
        if pd.isna(row["expected_sig_vqtl"]):
            return np.nan
        return bool(row["significant_vqtl"] == row["expected_sig_vqtl"])

    df["match_vqtl"] = df.apply(_match_vqtl, axis=1)

    def _as_expected(row):
        if pd.isna(row["match_vqtl"]):
            return np.nan
        return bool(row["match_ge"] and row["match_vqtl"])

    df["as_expected"] = df.apply(_as_expected, axis=1)
    return df


# ============================================================
# Reporting: plots (matplotlib) + Word report (python-docx), always written
# under ISOLATED_ROOT/plots/ and ISOLATED_ROOT/isolated_causal_test_report.docx.
# Requires: pip install matplotlib python-docx statsmodels
# ============================================================

def _fmt_p(p) -> str:
    if p is None or pd.isna(p):
        return "n/a"
    return f"{p:.4f}" if p >= 0.0005 else f"{p:.2e}"


def _sd_ratio(d):
    if not isinstance(d, dict) or not d:
        return None
    vals = list(d.values())
    return max(vals) / min(vals) if min(vals) > 0 else None


def generate_plots(df: pd.DataFrame, thresholds: dict) -> dict:
    os.makedirs(PLOTS_DIR, exist_ok=True)
    paths = {}

    ge_power = df[df["ge_role"] == "power"].copy()
    ge_fpr = df[df["ge_role"] == "fpr_control"].copy()
    vqtl_power = df[df["vqtl_role"] == "power"].copy()
    ge_crosstalk = df[df["ge_role"] == "crosstalk_check"].copy()   # vQTL-only variants tested on GxE
    vqtl_crosstalk = df[df["vqtl_role"] == "crosstalk_check"].copy()  # GxE-only variants tested on vQTL

    # ---- 1) GxE power scatter: |beta_inter| vs outcome ----
    if not ge_power.empty:
        fig, ax = plt.subplots(figsize=(7, 4.2))
        abs_beta = ge_power["beta_inter"].abs()
        colors = ge_power["recovered_ge"].map({True: "#1D9E75", False: "#D85A30"})
        pos = ge_power["beta_inter"] > 0
        ax.scatter(abs_beta[pos], [1] * pos.sum(), c=colors[pos], marker="o", s=70,
                   edgecolor="white", linewidth=0.5, label="positive sign")
        ax.scatter(abs_beta[~pos], [0] * (~pos).sum(), c=colors[~pos], marker="s", s=70,
                   edgecolor="white", linewidth=0.5, label="negative sign")
        ax.set_yticks([0, 1])
        ax.set_yticklabels(["negative beta_inter", "positive beta_inter"])
        ax.set_xlabel("|beta_inter|")
        ax.set_title("GxE isolated power test — outcome by effect magnitude and sign")
        legend_elems = [
            Line2D([0], [0], marker="o", color="w", markerfacecolor="#1D9E75", markersize=9, label="recovered (significant)"),
            Line2D([0], [0], marker="o", color="w", markerfacecolor="#D85A30", markersize=9, label="missed (not significant)"),
        ]
        ax.legend(handles=legend_elems, loc="lower right", frameon=False)
        ax.grid(axis="x", alpha=0.3)
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "ge_power_scatter.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["ge_power_scatter"] = p

        bins = list(range(0, 10))
        ge_power["bin"] = pd.cut(abs_beta, bins=bins, right=True)
        by_bin = ge_power.groupby("bin", observed=True)["recovered_ge"].agg(["sum", "count"])
        by_bin["power_pct"] = 100 * by_bin["sum"] / by_bin["count"]
        by_bin = by_bin[by_bin["count"] > 0]
        fig, ax = plt.subplots(figsize=(7, 4))
        ax.bar([str(b) for b in by_bin.index], by_bin["power_pct"], color="#2a78d6")
        for i in range(len(by_bin)):
            ax.text(i, by_bin["power_pct"].iloc[i] + 2,
                     f"{int(by_bin['sum'].iloc[i])}/{int(by_bin['count'].iloc[i])}", ha="center", fontsize=8)
        ax.set_ylim(0, 110)
        ax.set_ylabel("power (%)")
        ax.set_xlabel("|beta_inter| bin")
        ax.set_title("GxE detection rate by effect-size bin")
        ax.grid(axis="y", alpha=0.3)
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "ge_power_by_bin.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["ge_power_by_bin"] = p

    # ---- 2) FPR controls (main-effect only) ----
    if not ge_fpr.empty:
        fig, ax = plt.subplots(figsize=(7, 4))
        colors = ge_fpr["recovered_ge"].map({True: "#1D9E75", False: "#E24B4A"})
        ax.bar(ge_fpr["variant"], ge_fpr["p_ge"].fillna(1.0), color=colors)
        ax.axhline(0.05, color="#898781", linestyle="--", linewidth=1, label="p = 0.05 threshold")
        ax.set_ylabel("p_emp (interaction test)")
        ax.set_title("Main-effect-only controls (beta_inter = 0) — false-positive check")
        ax.tick_params(axis="x", rotation=60, labelsize=7)
        ax.legend(frameon=False, loc="upper right")
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "ge_fpr_controls.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["ge_fpr_controls"] = p

    # ---- 3) vQTL power scatter ----
    if not vqtl_power.empty:
        vqtl_power["sd_ratio"] = vqtl_power["sd_by_dosage"].apply(_sd_ratio)
        fig, ax = plt.subplots(figsize=(7, 4.2))
        colors = vqtl_power["recovered_vqtl"].map({True: "#1D9E75", False: "#D85A30"})
        ax.scatter(vqtl_power["sd_ratio"], vqtl_power["p_vqtl"].fillna(1.0), c=colors, s=70,
                   edgecolor="white", linewidth=0.5)
        ax.axhline(0.05, color="#898781", linestyle="--", linewidth=1, label="p = 0.05 threshold")
        ax.set_xlabel("sd ratio (max/min across dosage)")
        ax.set_ylabel("p (asymptotic)")
        ax.set_title("vQTL isolated power test — pure_variance variants")
        ax.legend(frameon=False, loc="upper right")
        ax.grid(alpha=0.3)
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "vqtl_power_scatter.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["vqtl_power_scatter"] = p

    # ---- 4) GxE crosstalk on the GxE side (unchanged: vQTL-only variants tested on GxE) ----
    if not ge_crosstalk.empty:
        fp = int((~ge_crosstalk["recovered_ge"]).sum())
        fig, ax = plt.subplots(figsize=(5.5, 4))
        pct = 100 * fp / len(ge_crosstalk)
        ax.bar(["vQTL variants tested on GxE\n(expect: not significant)"], [pct], color="#8a4fd6")
        ax.text(0, pct + 1.5, f"{fp}/{len(ge_crosstalk)}", ha="center", fontsize=9)
        ax.set_ylabel("spurious GxE detection rate (%)")
        ax.set_ylim(0, max(pct, 10) * 1.3)
        ax.set_title("GxE specificity check (vQTL-only variants)")
        ax.grid(axis="y", alpha=0.3)
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "ge_specificity.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["ge_specificity"] = p

    # ---- 5) vQTL confound-aware curve: predicted_var_ratio vs significance,
    #         with fitted thresholds and zone coloring. This REPLACES the old
    #         single-number "crosstalk specificity" bar for the vQTL side,
    #         which conflated genuine specificity failures with mechanically
    #         expected hits. ----
    if not vqtl_crosstalk.empty and vqtl_crosstalk["predicted_var_ratio"].notna().any():
        zone_colors = {"pure_null": "#2a78d6", "confounded_expected": "#d68a2a", "borderline": "#9a9a9a"}
        fig, ax = plt.subplots(figsize=(7.5, 4.5))
        for zone, color in zone_colors.items():
            sub = vqtl_crosstalk[vqtl_crosstalk["vqtl_zone"] == zone] if "vqtl_zone" in vqtl_crosstalk.columns else pd.DataFrame()
            if sub.empty:
                continue
            marker_colors = sub["significant_vqtl"].map({True: color, False: "white"})
            ax.scatter(sub["predicted_var_ratio"], sub["p_vqtl"].fillna(1.0), c=marker_colors,
                       edgecolor=color, linewidth=1.3, s=60, label=zone, zorder=3)
        ax.axhline(0.05, color="#898781", linestyle="--", linewidth=1, zorder=1)
        ax.axvline(thresholds["ratio_low"], color="#2a78d6", linestyle=":", linewidth=1.2,
                   label=f"ratio_low={thresholds['ratio_low']:.2f}", zorder=1)
        ax.axvline(thresholds["ratio_high"], color="#d68a2a", linestyle=":", linewidth=1.2,
                   label=f"ratio_high={thresholds['ratio_high']:.2f}", zorder=1)
        ax.set_xscale("log")
        ax.set_yscale("log")
        ax.set_xlabel("predicted_var_ratio = 1 + (beta_inter / noise_sd)^2  [log scale]")
        ax.set_ylabel("p_vqtl  [log scale]")
        ax.set_title(f"vQTL outcome on G×E-only variants vs. predicted confound\n(thresholds: {thresholds['method']})")
        ax.legend(fontsize=7, frameon=False, loc="lower left")
        ax.grid(alpha=0.3, which="both")
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "vqtl_confound_curve.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["vqtl_confound_curve"] = p

    # ---- 6) Overlap: power variants recovered by BOTH mechanisms ----
    both_defined = df[df["in_ge_pool"] & df["in_vqtl_pool"]].copy()
    if not both_defined.empty:
        fig, ax = plt.subplots(figsize=(6, 4.5))
        cats = ["GxE only", "vQTL only", "Both", "Neither"]
        ge_ok = both_defined["recovered_ge"].fillna(False)
        vq_ok = both_defined["recovered_vqtl"].fillna(False)
        counts = [
            int((ge_ok & ~vq_ok).sum()),
            int((~ge_ok & vq_ok).sum()),
            int((ge_ok & vq_ok).sum()),
            int((~ge_ok & ~vq_ok).sum()),
        ]
        ax.bar(cats, counts, color=["#2a78d6", "#d68a2a", "#1D9E75", "#D85A30"])
        for i, c in enumerate(counts):
            ax.text(i, c + 0.1, str(c), ha="center", fontsize=9)
        ax.set_ylabel("number of variants")
        ax.set_title("Variants designed for BOTH mechanisms — detection overlap")
        ax.grid(axis="y", alpha=0.3)
        fig.tight_layout()
        p = os.path.join(PLOTS_DIR, "overlap_both_mechanisms.png")
        fig.savefig(p, dpi=150); plt.close(fig)
        paths["overlap_both_mechanisms"] = p

    return paths


def generate_word_report(df: pd.DataFrame, plot_paths: dict, thresholds: dict) -> str:
    ge_power = df[df["ge_role"] == "power"]
    ge_fpr = df[df["ge_role"] == "fpr_control"]
    vqtl_power = df[df["vqtl_role"] == "power"]
    ge_crosstalk = df[df["ge_role"] == "crosstalk_check"]
    vqtl_crosstalk = df[df["vqtl_role"] == "crosstalk_check"]
    pure_null = vqtl_crosstalk[vqtl_crosstalk["vqtl_zone"] == "pure_null"] if "vqtl_zone" in vqtl_crosstalk.columns else pd.DataFrame()
    confounded = vqtl_crosstalk[vqtl_crosstalk["vqtl_zone"] == "confounded_expected"] if "vqtl_zone" in vqtl_crosstalk.columns else pd.DataFrame()
    borderline = vqtl_crosstalk[vqtl_crosstalk["vqtl_zone"] == "borderline"] if "vqtl_zone" in vqtl_crosstalk.columns else pd.DataFrame()

    doc = Document()
    doc.add_heading("Isolated causal variant test — combined results report", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True
    doc.add_paragraph(
        "Every variant from either default pool was tested one at a time against an "
        "otherwise-null panel, on BOTH the G×E and the vQTL mechanism simultaneously "
        "(using a null/flat configuration on whichever mechanism it was not designed for). "
        "Combined results below are assembled directly from the per-variant result files "
        "saved on disk."
    )
    doc.add_paragraph(
        "Note on vQTL specificity: a G×E-only variant with interaction term "
        "beta_inter*dosage*exposure mechanically induces heteroscedasticity by dosage "
        "(predicted variance ratio = 1 + (beta_inter/noise_sd)^2), independent of any "
        "declared vQTL effect. A vQTL hit on such a variant is therefore only a genuine "
        "specificity failure when the predicted ratio is small; when it is large, "
        "significance is the mechanically expected outcome. G×E-only variants are "
        f"classified into three zones EVERY RUN (not fixed thresholds) via a logistic fit "
        f"of vQTL significance on the predicted ratio -- see 'vQTL specificity' section below "
        f"(method used this run: {thresholds['method']}, ratio_low={thresholds['ratio_low']:.3f}, "
        f"ratio_high={thresholds['ratio_high']:.3f})."
    )

    doc.add_heading("Summary", level=1)
    if not ge_power.empty:
        n_rec, n_tot = int(ge_power["recovered_ge"].sum()), len(ge_power)
        doc.add_paragraph(
            f"GxE power test: {n_rec}/{n_tot} causal variants recovered "
            f"(overall power {100 * n_rec / n_tot:.1f}%)."
        )
    if not ge_fpr.empty:
        n_fp, n_tot = int((~ge_fpr["recovered_ge"]).sum()), len(ge_fpr)
        doc.add_paragraph(
            f"GxE false-positive controls (main effect only, beta_inter = 0): "
            f"{n_fp}/{n_tot} incorrectly flagged as significant "
            f"(false-positive rate {100 * n_fp / n_tot:.1f}%)."
        )
    if not vqtl_power.empty:
        n_rec, n_tot = int(vqtl_power["recovered_vqtl"].sum()), len(vqtl_power)
        doc.add_paragraph(
            f"vQTL power test (pure_variance): {n_rec}/{n_tot} causal variants recovered "
            f"(overall power {100 * n_rec / n_tot:.1f}%)."
        )
    if not ge_crosstalk.empty:
        n_fp, n_tot = int((~ge_crosstalk["recovered_ge"]).sum()), len(ge_crosstalk)
        doc.add_paragraph(
            f"Cross-mechanism check, vQTL-only variants tested on G×E: {n_fp}/{n_tot} "
            f"showed a spurious G×E signal ({100 * n_fp / n_tot:.1f}%)."
        )

    doc.add_heading("vQTL specificity (G×E-confound-aware)", level=1)
    doc.add_paragraph(
        f"Thresholds this run: ratio_low={thresholds['ratio_low']:.3f}, "
        f"ratio_high={thresholds['ratio_high']:.3f} (method: {thresholds['method']}"
        + (f", {thresholds['note']}" if thresholds.get("note") else "") + ")."
    )
    if not pure_null.empty:
        n_fp = int((~pure_null["match_vqtl"].astype(bool)).sum())
        doc.add_paragraph(
            f"Genuine specificity zone (pure_null, predicted_var_ratio <= {thresholds['ratio_low']:.3f}): "
            f"{n_fp}/{len(pure_null)} showed a spurious vQTL signal "
            f"({100 * n_fp / len(pure_null):.1f}%) -- these are real specificity failures / noise."
        )
    if not confounded.empty:
        n_sig = int(confounded["significant_vqtl"].sum())
        doc.add_paragraph(
            f"Mechanically-confounded zone (confounded_expected, predicted_var_ratio >= "
            f"{thresholds['ratio_high']:.3f}): {n_sig}/{len(confounded)} were significant on vQTL "
            f"({100 * n_sig / len(confounded):.1f}%) -- this is the EXPECTED outcome given the "
            f"induced heteroscedasticity, not a specificity failure."
        )
    if not borderline.empty:
        n_sig = int(borderline["significant_vqtl"].sum())
        doc.add_paragraph(
            f"Borderline zone ({thresholds['ratio_low']:.3f} < predicted_var_ratio < "
            f"{thresholds['ratio_high']:.3f}): {len(borderline)} variants excluded from the pass/fail "
            f"counts above ({n_sig} were significant) -- reported for transparency, not scored."
        )

    failed = df[df["status"] == "FAILED"]
    if not failed.empty:
        p = doc.add_paragraph()
        run = p.add_run(f"{len(failed)} variant run(s) raised an exception and could not be evaluated — see the detail table.")
        run.font.color.rgb = RGBColor(0xE2, 0x4B, 0x4A)

    if plot_paths:
        doc.add_heading("Plots", level=1)
        captions = {
            "ge_power_scatter": "GxE isolated power test — outcome by effect magnitude and sign.",
            "ge_power_by_bin": "GxE detection rate by effect-size bin.",
            "ge_fpr_controls": "Main-effect-only controls — false-positive check.",
            "vqtl_power_scatter": "vQTL isolated power test — pure_variance variants.",
            "ge_specificity": "GxE specificity check (vQTL-only variants tested on GxE).",
            "vqtl_confound_curve": "vQTL outcome on G×E-only variants vs. predicted variance-ratio confound, with data-driven zone thresholds.",
            "overlap_both_mechanisms": "Detection overlap for variants designed on both mechanisms.",
        }
        for key in ["ge_power_scatter", "ge_power_by_bin", "ge_fpr_controls",
                    "vqtl_power_scatter", "ge_specificity", "vqtl_confound_curve", "overlap_both_mechanisms"]:
            if key in plot_paths:
                doc.add_picture(plot_paths[key], width=Inches(6))
                cap = doc.add_paragraph(captions[key])
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    run.italic = True
                    run.font.size = Pt(9)

    doc.add_heading("Detailed results", level=1)
    cols = ["variant", "in_ge_pool", "in_vqtl_pool",
            "ge_role", "beta_inter", "beta_main", "p_ge", "significant_ge", "expected_sig_ge", "match_ge",
            "vqtl_role", "sd_by_dosage", "predicted_var_ratio", "vqtl_zone",
            "p_vqtl", "significant_vqtl", "expected_sig_vqtl", "match_vqtl",
            "as_expected", "status"]
    cols = [c for c in cols if c in df.columns]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = c
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    ordered = df.sort_values(["ge_role", "vqtl_role", "variant"])
    for _, row in ordered.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            val = row[c]
            if c in ("p_ge", "p_vqtl"):
                val = _fmt_p(val)
            elif c == "predicted_var_ratio" and pd.notna(val):
                val = f"{val:.3f}"
            cells[i].text = "" if pd.isna(val) else str(val)
        for c in table.rows[-1].cells:
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(8)

    out_path = os.path.join(ISOLATED_ROOT, "isolated_causal_test_report.docx")
    doc.save(out_path)
    return out_path


# ============================================================
# PHASE 1: isolated per-variant test (everything above this point).
# Runs every variant in the plan (parallel or sequential), rebuilds the
# combined table from disk, classifies vQTL zones, writes CSV/JSON/plots/
# docx, and returns a result dict summarizing pass/fail for main() to
# combine with the other two phases' outcomes.
# ============================================================

def run_isolated_phase(n_workers: int = 1, force: bool = False) -> dict:
    os.makedirs(ISOLATED_ROOT, exist_ok=True)
    plan = build_variant_plan()
    print(f"{len(plan)} unique variants to test (union of both default pools), both mechanisms each.")

    if n_workers == 1:
        for entry in plan:
            run_isolated_variant(entry, n_workers=1, force=force)
    else:
        print(f"Running {len(plan)} isolated variants with {n_workers} parallel processes...")
        with ProcessPoolExecutor(max_workers=n_workers) as pool:
            futures = {pool.submit(_worker, entry, n_workers, force): entry["variant"] for entry in plan}
            for fut in as_completed(futures):
                label = futures[fut]
                try:
                    res = fut.result()
                    status = res.get("status", "?")
                except Exception as exc:
                    status = "FAILED"
                    print(f"\n*** VARIANT '{label}' FAILED in worker process: {exc} ***")
                print(f"[{label}] done ({status}).")

    section("COMBINED RESULTS (assembled from disk)")
    # Always rebuild the combined table from the per-variant isolated_summary.json
    # files on disk, rather than from whatever this run's execution happened to
    # accumulate in memory (see collect_results_from_disk docstring).
    df = collect_results_from_disk(plan)
    if df.empty:
        print("No results found on disk -- nothing to report.")
        return {"df": pd.DataFrame(), "has_failures": True}

    # Zone thresholds are refit from THIS run's data every time (see module
    # docstring for why they can't be hard-coded), then applied to classify
    # every variant and compute expected_sig_ge/vqtl, match_ge/vqtl, as_expected.
    thresholds = fit_vqtl_confound_thresholds(df)
    df = add_expectation_columns(df, thresholds)

    print(df.to_string(index=False))

    summary_csv = os.path.join(ISOLATED_ROOT, "isolated_power_curve.csv")
    df.to_csv(summary_csv, index=False)
    with open(os.path.join(ISOLATED_ROOT, "isolated_power_curve.json"), "w") as f:
        json.dump(df.to_dict(orient="records"), f, indent=2, default=str)
    with open(os.path.join(ISOLATED_ROOT, "vqtl_confound_thresholds.json"), "w") as f:
        json.dump(thresholds, f, indent=2, default=str)
    print(f"\n[export] {summary_csv}")
    print(f"[export] {os.path.join(ISOLATED_ROOT, 'vqtl_confound_thresholds.json')}")

    section("REPORT — plots and Word")
    plot_paths = generate_plots(df, thresholds)
    for name, p in plot_paths.items():
        print(f"[plot] {name}: {p}")
    report_path = generate_word_report(df, plot_paths, thresholds)
    print(f"[export] {report_path}")

    ok = df[df["status"] == "ok"]
    missed_power_ge = ok.loc[(ok["ge_role"] == "power") & (~ok["match_ge"]), "variant"].tolist()
    fp_ge = ok.loc[(ok["ge_role"] != "power") & (~ok["match_ge"]), "variant"].tolist()
    missed_power_vqtl = ok.loc[(ok["vqtl_role"] == "power") & (~ok["match_vqtl"].astype("boolean").fillna(True)), "variant"].tolist()
    # genuine vQTL specificity failures: only in the pure_null zone (see report)
    fp_vqtl_genuine = ok.loc[
        (ok["vqtl_zone"] == "pure_null") & (~ok["match_vqtl"].astype("boolean").fillna(True)), "variant"
    ].tolist()
    failed = df.loc[df["status"] == "FAILED", "variant"].tolist()

    if missed_power_ge:
        print(f"\n*** CAUSAL VARIANTS NOT RECOVERED (GxE power, isolated): {missed_power_ge} ***")
    if fp_ge:
        print(f"*** SPURIOUS GxE SIGNIFICANCE (fpr_control or crosstalk): {fp_ge} ***")
    if missed_power_vqtl:
        print(f"\n*** CAUSAL VARIANTS NOT RECOVERED (vQTL power, isolated): {missed_power_vqtl} ***")
    if fp_vqtl_genuine:
        print(f"*** GENUINE SPURIOUS vQTL SIGNIFICANCE (pure_null zone only): {fp_vqtl_genuine} ***")
    if failed:
        print(f"*** FAILED VARIANTS (exception): {failed} ***")

    has_failures = bool(missed_power_ge or fp_ge or missed_power_vqtl or fp_vqtl_genuine or failed)
    if not has_failures:
        print("\n*** All isolated causal variants recovered correctly on both mechanisms "
              "(vQTL judged with G×E-confound-aware zones). ***")

    return {
        "df": df,
        "report_path": report_path,
        "missed_power_ge": missed_power_ge,
        "fp_ge": fp_ge,
        "missed_power_vqtl": missed_power_vqtl,
        "fp_vqtl_genuine": fp_vqtl_genuine,
        "failed": failed,
        "has_failures": has_failures,
    }


# ============================================================
# PHASE 3: pure_null zone confidence interval (ported from the former
# analize_pure_null.py -- it only made sense as a follow-up to phase 1, so
# it now lives here as the phase run right after it by default).
#
# Exact binomial CI on the observed false-positive rate in the isolated
# run's "pure_null" vQTL zone (predicted_var_ratio small enough that a vQTL
# hit is genuine noise, not G×E-induced heteroscedasticity -- see module
# docstring). Only if the nominal 5% falls outside (or close to the edge
# of) that CI do we spend the extra comparison against the bootstrap
# p-value (K=200, steadier SE estimate than the asymptotic one at K=50) to
# tell apart "estimation noise at K=50" from "a real bias in the test".
# This needs p_vqtl_boot/significant_vqtl_boot, which run_isolated_variant()
# now saves for every variant (previously missing: run_vqtl_debug already
# computes both se_methods for every variant, so this costs nothing extra).
# ============================================================

NOMINAL_VQTL_RATE = 0.05


def _binomial_ci_block(n_sig: int, n_tot: int, alpha: float) -> dict:
    if n_tot == 0:
        return {"n": 0, "n_significant": 0, "observed_rate": None,
                "ci_low": None, "ci_high": None, "nominal_in_ci": None}
    res = binomtest(n_sig, n_tot, p=NOMINAL_VQTL_RATE, alternative="two-sided")
    ci = res.proportion_ci(confidence_level=1 - alpha, method="exact")
    return {
        "n": n_tot,
        "n_significant": n_sig,
        "observed_rate": round(n_sig / n_tot, 4),
        "ci_low": round(ci.low, 4),
        "ci_high": round(ci.high, 4),
        "confidence_level": 1 - alpha,
        "nominal_in_ci": bool(ci.low <= NOMINAL_VQTL_RATE <= ci.high),
        "binom_test_pvalue": round(res.pvalue, 4),
    }


def run_pure_null_ci_check(df: pd.DataFrame, alpha: float = 0.05, near_edge_margin: float = 0.02) -> dict:
    if "vqtl_zone" not in df.columns:
        raise SystemExit(
            "isolated_power_curve.csv doesn't have a vqtl_zone column: was it written by an "
            "old version of this script? Re-run the isolated phase (drop --skip-isolated)."
        )

    pure_null = df[(df["vqtl_zone"] == "pure_null") & (df["status"] == "ok")].copy()
    n_tot = len(pure_null)
    n_sig = int(pure_null["significant_vqtl"].sum()) if n_tot else 0
    if n_tot:
        print(f"pure_null zone: n={n_tot} variants, {n_sig} significant on P_asym "
              f"(observed rate {100 * n_sig / n_tot:.1f}%).")
    else:
        print("pure_null zone is empty: nothing to check.")

    ci_asym = _binomial_ci_block(n_sig, n_tot, alpha)
    if n_tot:
        print(f"Exact {int(ci_asym['confidence_level'] * 100)}% CI on P_asym: "
              f"[{ci_asym['ci_low']}, {ci_asym['ci_high']}]")

    report = {"alpha": alpha, "n_pure_null_total_rows": n_tot, "asym": ci_asym}

    proceed_to_boot = False
    if n_tot == 0:
        pass
    elif not ci_asym["nominal_in_ci"]:
        print("The nominal 5% falls OUTSIDE the P_asym CI -> comparing against P_boot.")
        proceed_to_boot = True
    else:
        dist_to_edge = min(abs(NOMINAL_VQTL_RATE - ci_asym["ci_low"]), abs(ci_asym["ci_high"] - NOMINAL_VQTL_RATE))
        if dist_to_edge < near_edge_margin:
            print(f"The nominal 5% is INSIDE the P_asym CI but close to the edge "
                  f"(distance {dist_to_edge:.3f} < margin {near_edge_margin}) -> "
                  f"comparing against P_boot anyway, to be safe.")
            proceed_to_boot = True
        else:
            print("The nominal 5% is comfortably inside the P_asym CI: no evidence of a problem. "
                  "No need for the P_boot comparison.")

    report["proceeded_to_boot_comparison"] = proceed_to_boot
    flag_for_review = False

    if proceed_to_boot:
        has_boot = "p_vqtl_boot" in pure_null.columns and pure_null["p_vqtl_boot"].notna().any()
        if not has_boot:
            print("\n[warn] No p_vqtl_boot available for the pure_null zone (isolated_power_curve.csv "
                  "was produced before run_isolated_variant() started saving it). Re-run the isolated "
                  "phase -- even just with --force on the affected variants is enough to repopulate "
                  "p_vqtl_boot without regenerating the datasets, since run_vqtl_debug always runs "
                  "both se_methods anyway.")
            report["boot"] = None
        else:
            if "significant_vqtl_boot" in pure_null.columns:
                n_sig_boot = int(pure_null["significant_vqtl_boot"].fillna(False).sum())
            else:
                n_sig_boot = int((pure_null["p_vqtl_boot"] < NOMINAL_VQTL_RATE).sum())
            ci_boot = _binomial_ci_block(n_sig_boot, n_tot, alpha)
            print(f"Same subset on P_boot (K=200): {n_sig_boot}/{n_tot} significant "
                  f"(observed rate {100 * n_sig_boot / n_tot:.1f}%), CI: [{ci_boot['ci_low']}, {ci_boot['ci_high']}]")
            report["boot"] = ci_boot

            if ci_boot["nominal_in_ci"] and not ci_asym["nominal_in_ci"]:
                print("\n-> P_boot (K=200) falls back within nominal while P_asym (K=50) doesn't: "
                      "consistent with SE estimation noise at low K, not a real bias. Worth considering "
                      "a higher default VQTL_ASYMPTOTIC_BOOTSTRAP_K, or se_method='bootstrap' for "
                      "false-positive screening.")
            elif not ci_boot["nominal_in_ci"] and not ci_asym["nominal_in_ci"]:
                print("\n-> P_boot (K=200, the production method) ALSO shows the nominal 5% outside the "
                      "CI: not just a K issue. Likely a systematic bias in the variance test itself, "
                      "independent of the bootstrap replicate count -- worth a dedicated check (e.g. a "
                      "QQ-plot of vQTL p-values under pure H0, with no G×E component at all).")
                flag_for_review = True
            else:
                print("\n-> Both methods fall within nominal on this subset: the signal initially seen "
                      "on P_asym was compatible with sampling noise at this n.")

    report["flag_for_review"] = flag_for_review
    out_path = os.path.join(ISOLATED_ROOT, "pure_null_ci_report.json")
    os.makedirs(ISOLATED_ROOT, exist_ok=True)
    with open(out_path, "w") as f:
        json.dump(report, f, indent=2, default=str)
    print(f"\n[export] {out_path}")
    return report


# ============================================================
# PHASE 2: robustness scenario battery, delegated to run_scenarios.py
# (run_all_scenarios()) -- not duplicated here, see run_scenarios.py's own
# docstring for what each scenario covers.
# ============================================================

def run_scenario_phase(n_workers: int = 1, names: list[str] | None = None) -> dict:
    return rs.run_all_scenarios(names=names, n_workers=n_workers)


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Full vQTL/G×E pipeline test battery: isolated per-variant test, "
                     "robustness scenario battery, and pure_null zone CI check."
    )
    parser.add_argument("--workers", type=int, default=1,
                         help="Parallel processes used within each phase (isolated variants, then "
                              "scenarios). Default 1 (sequential).")
    parser.add_argument("--force", action="store_true",
                         help="Isolated phase only: recompute even variants that already have an "
                              "isolated_summary.json with status ok (default: cached result reused).")
    parser.add_argument("--skip-isolated", action="store_true", help="Skip phase 1 (isolated per-variant test).")
    parser.add_argument("--skip-scenarios", action="store_true", help="Skip phase 2 (robustness scenario battery).")
    parser.add_argument("--skip-pure-null", action="store_true", help="Skip phase 3 (pure_null zone CI check).")
    parser.add_argument("--scenarios", nargs="*", default=None,
                         help="Phase 2 only: specific scenario names to run (default: all of them).")
    parser.add_argument("--pure-null-alpha", type=float, default=0.05,
                         help="Phase 3: 1 - confidence level for the binomial CI (default 0.05 -> 95%% CI).")
    parser.add_argument("--pure-null-near-edge-margin", type=float, default=0.02,
                         help="Phase 3: if the nominal 5%% is INSIDE the CI but closer than this to its "
                              "edge, run the P_boot comparison anyway, out of caution (default 0.02).")
    parser.add_argument("--output-dir", default=None,
                         help="Cartella dove scrivere isolated/ e scenarios/ (default: la cartella di "
                              "questo script). Vale per tutte e 3 le fasi.")
    args = parser.parse_args()

    if args.output_dir:
        global ISOLATED_ROOT, PLOTS_DIR
        output_root = os.path.abspath(args.output_dir)
        os.makedirs(output_root, exist_ok=True)
        ISOLATED_ROOT = os.path.join(output_root, "isolated")
        PLOTS_DIR = os.path.join(ISOLATED_ROOT, "plots")
        rs.SCENARIOS_ROOT = os.path.join(output_root, "scenarios")
        print(f"[config] Output root: {output_root} "
              f"(isolated: {ISOLATED_ROOT} | scenarios: {rs.SCENARIOS_ROOT})")

    n_workers = max(1, args.workers)
    t0 = time.time()
    phase_results = {}

    if not args.skip_isolated:
        section("PHASE 1/3 — ISOLATED PER-VARIANT TEST")
        phase_results["isolated"] = run_isolated_phase(n_workers=n_workers, force=args.force)
    else:
        print("[skip] phase 1 (isolated) skipped by --skip-isolated.")

    if not args.skip_scenarios:
        section("PHASE 2/3 — ROBUSTNESS SCENARIO BATTERY")
        phase_results["scenarios"] = run_scenario_phase(n_workers=n_workers, names=args.scenarios)
    else:
        print("[skip] phase 2 (scenarios) skipped by --skip-scenarios.")

    if not args.skip_pure_null:
        section("PHASE 3/3 — PURE_NULL ZONE CONFIDENCE INTERVAL")
        isolated_df = phase_results.get("isolated", {}).get("df")
        if isolated_df is None or isolated_df.empty:
            # Phase 1 wasn't run this time (or produced nothing): fall back to
            # whatever isolated_power_curve.csv is already on disk from a
            # previous run, same as the old standalone analize_pure_null.py.
            csv_path = os.path.join(ISOLATED_ROOT, "isolated_power_curve.csv")
            if os.path.isfile(csv_path):
                isolated_df = pd.read_csv(csv_path)
            else:
                isolated_df = None
        if isolated_df is None:
            print(f"[skip] phase 3 (pure_null) skipped: no isolated_power_curve.csv found "
                  f"(run phase 1 at least once, or drop --skip-isolated).")
        else:
            phase_results["pure_null"] = run_pure_null_ci_check(
                isolated_df, alpha=args.pure_null_alpha, near_edge_margin=args.pure_null_near_edge_margin)
    else:
        print("[skip] phase 3 (pure_null) skipped by --skip-pure-null.")

    section("FINAL SUMMARY — all phases")
    print(f"Completed in {time.time() - t0:.0f}s.")

    isolated_failed = phase_results.get("isolated", {}).get("has_failures", False)
    scenarios_failed = phase_results.get("scenarios", {}).get("has_failures", False)
    pure_null_flag = phase_results.get("pure_null", {}).get("flag_for_review", False)

    if isolated_failed:
        print("*** Phase 1 (isolated): FAILURES found, see above. ***")
    if scenarios_failed:
        print(f"*** Phase 2 (scenarios): failed scenarios "
              f"{phase_results['scenarios'].get('failed')} / "
              f"vQTL-check failures {phase_results['scenarios'].get('vqtl_failed')}. ***")
    if pure_null_flag:
        print("*** Phase 3 (pure_null): possible systematic bias flagged for manual review, see above. ***")

    if isolated_failed or scenarios_failed:
        sys.exit(1)
    if not (isolated_failed or scenarios_failed or pure_null_flag):
        print("*** All executed phases completed with no failures. ***")


if __name__ == "__main__":
    main()
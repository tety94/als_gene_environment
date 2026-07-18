#!/usr/bin/env python3
"""
build_supplementary_report.py
==============================
Assembles a single Word (.docx) supplementary document — tables and
figures ready to paste into a paper's Supplementary Materials, all in
English — by reading the CSV/PNG/log outputs already produced by:

    00_run_plink_qc.sh
    01_run_extra_qc_checks.sh
    qc_report.py
    interpret_plink_output.py
    qc_attrition_summary.py
    qc_supplementary_plots.py

This script does NOT recompute anything. It only reads existing output
files and formats them. If an expected input file is missing, the
corresponding table/figure is skipped with a placeholder note instead of
crashing the whole build — run the missing upstream step and re-run this
script to fill the gap.

USAGE
-----
python3 build_supplementary_report.py \
    --qc-dir /path/to/qc_output \
    --kinship-report-dir /path/to/qc_output/qc_report \
    --diagnostics-dir /path/to/qc_output/diagnostics_output \
    --attrition-csv /path/to/qc_output/qc_attrition.csv \
    --supp-plots-dir /path/to/qc_output/supplementary_plots \
    --out /path/to/qc_output/Supplementary_QC_Report.docx

All directory/file arguments are optional and default to the standard
subfolder names used in the pipeline's own examples, relative to
--qc-dir. Pass them explicitly if you used different output paths.
"""

import argparse
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------

class FigureCounter:
    def __init__(self):
        self.table_n = 0
        self.figure_n = 0

    def table(self):
        self.table_n += 1
        return f"Table S{self.table_n}"

    def figure(self):
        self.figure_n += 1
        return f"Figure S{self.figure_n}"


def set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_caption(doc, label, text, above=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{label}. ")
    run.bold = True
    run2 = p.add_run(text)
    run2.italic = not above  # figure captions (below) italic; table captions (above) not
    return p


def add_missing_note(doc, what, expected_path):
    p = doc.add_paragraph()
    run = p.add_run(f"[Not available: {what} — expected file not found at {expected_path}. "
                     f"Run the corresponding upstream script and re-run this report.]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)


def add_df_table(doc, df: pd.DataFrame, col_widths_in=None):
    n_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(hdr_cells[i], "D9E2F3")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = "" if pd.isna(val) else str(val)

    if col_widths_in:
        for row in table.rows:
            for i, w in enumerate(col_widths_in):
                row.cells[i].width = Inches(w)

    return table


def add_figure(doc, img_path: Path, width_in=6.0):
    doc.add_picture(str(img_path), width=Inches(width_in))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Readers for raw plink2 / pipeline output formats
# ---------------------------------------------------------------------------

def read_plink_table(path: Path) -> pd.DataFrame:
    df = pd.read_csv(path, sep=r"\s+")
    df.columns = [c.lstrip("#") for c in df.columns]
    return df


def parse_diagnostics_report(path: Path) -> dict:
    """
    Pulls the key numbers out of interpret_plink_output.py's
    diagnostics_report.txt via regex on the fixed log lines it writes.
    """
    out = {}
    if not path.exists():
        return out
    text = path.read_text(encoding="utf-8")

    m = re.search(r"Coppie totali riportate da plink2:\s*(\d+)", text)
    if m:
        out["n_pairs_total"] = int(m.group(1))

    m = re.search(r"Soglia PI_HAT:\s*([\d.]+)", text)
    if m:
        out["pi_hat_threshold"] = float(m.group(1))

    m = re.search(r"Coppie sopra soglia:\s*(\d+)\s*\(([\d.]+)% del totale\)", text)
    if m:
        out["n_pairs_above_threshold"] = int(m.group(1))
        out["pct_pairs_above_threshold"] = float(m.group(2))

    m = re.search(r"R\^2 \([^)]*\):\s*([\d.]+)\s*\(([\d.]+)%", text)
    if m:
        out["pc_exposure_r2"] = float(m.group(1))
        out["pc_exposure_r2_pct"] = float(m.group(2))

    m = re.search(r"N test:\s*(\d+)\s*\|\s*Lambda GC\s*=\s*([\d.]+)", text)
    if m:
        out["lambda_gc_n_tests"] = int(m.group(1))
        out["lambda_gc"] = float(m.group(2))

    m = re.search(r"Campioni in eigenvec:\s*(\d+)\s*\|\s*in metadata:\s*(\d+)\s*\|\s*matchati:\s*(\d+)", text)
    if m:
        out["n_eigenvec"] = int(m.group(1))
        out["n_metadata"] = int(m.group(2))
        out["n_matched"] = int(m.group(3))

    return out


def parse_run_metadata(path: Path) -> dict:
    out = {}
    if not path.exists():
        return out
    lines = path.read_text(encoding="utf-8").splitlines()
    for i, line in enumerate(lines):
        if line.startswith("Data esecuzione extra checks:"):
            out["run_date"] = line.split(":", 1)[1].strip()
        elif line.startswith("Host:"):
            out["host"] = line.split(":", 1)[1].strip()
        elif line.strip() == "plink2 --version:" and i + 1 < len(lines):
            out["plink2_version"] = lines[i + 1].strip()
        elif line.strip() == "bcftools --version:" and i + 1 < len(lines):
            out["bcftools_version"] = lines[i + 1].strip()
    return out


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def build_title(doc, qc_dir):
    title = doc.add_heading("Supplementary Material: Genomic Quality Control", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"QC pipeline output directory: {qc_dir}")
    run.italic = True
    run.font.size = Pt(10)
    doc.add_paragraph()


def build_intro(doc):
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "This document summarizes the genotype quality-control (QC) pipeline applied to the "
        "study cohort prior to downstream gene-environment interaction analysis. Variant "
        "calling batches were merged, filtered for missingness, LD-pruned, and screened for "
        "cryptic relatedness and population stratification using PLINK2. The tables and "
        "figures below are generated automatically from the pipeline's own output files and "
        "are intended as a starting point for the Methods and Supplementary Materials "
        "sections; verdicts and thresholds should be reviewed before submission."
    )


def build_reproducibility(doc, qc_dir, counters):
    doc.add_heading("2. Reproducibility Information", level=1)
    meta_path = qc_dir / "run_metadata.txt"
    meta = parse_run_metadata(meta_path)
    if not meta:
        add_missing_note(doc, "software version / reproducibility metadata", meta_path)
        return
    label = counters.table()
    add_caption(doc, label, "Software versions and run metadata.")
    df = pd.DataFrame({
        "Field": ["Run date", "Host", "PLINK2 version", "bcftools version"],
        "Value": [
            meta.get("run_date", "n/a"),
            meta.get("host", "n/a"),
            meta.get("plink2_version", "n/a"),
            meta.get("bcftools_version", "n/a"),
        ],
    })
    add_df_table(doc, df, col_widths_in=[2.2, 4.0])
    doc.add_paragraph()


def build_attrition(doc, attrition_csv, counters):
    doc.add_heading("3. Sample and Variant Attrition", level=1)
    if not attrition_csv.exists():
        add_missing_note(doc, "sample/variant attrition table", attrition_csv)
        return
    df = pd.read_csv(attrition_csv)
    # Translate column headers to English for the paper table.
    rename = {
        "stage": "QC stage",
        "n_samples": "N samples",
        "n_variants": "N variants",
        "samples_dropped_step": "Samples dropped (this step)",
        "variants_dropped_step": "Variants dropped (this step)",
        "pct_samples_remaining": "% samples remaining",
        "pct_variants_remaining": "% variants remaining",
    }
    df = df.rename(columns=rename)
    stage_translate = {
        "post-merge (pre-QC)": "Post-merge (pre-QC)",
        "post --geno": "Post variant-missingness filter (--geno)",
        "post --mind": "Post sample-missingness filter (--mind)",
        "post LD pruning": "Post LD pruning",
    }
    if "QC stage" in df.columns:
        df["QC stage"] = df["QC stage"].replace(stage_translate)

    label = counters.table()
    add_caption(doc, label, "Sample and variant counts across successive QC filtering steps.")
    add_df_table(doc, df)
    doc.add_paragraph()

    png = attrition_csv.with_suffix(".png")
    if png.exists():
        flabel = counters.figure()
        add_figure(doc, png)
        add_caption(doc, flabel, "Sample and variant counts across QC stages.", above=False)
    doc.add_paragraph()


def build_relatedness(doc, kinship_dir, diagnostics_dir, counters):
    doc.add_heading("4. Relatedness (Kinship) Screening", level=1)

    diag_report = diagnostics_dir / "diagnostics_report.txt"
    diag = parse_diagnostics_report(diag_report)
    if diag:
        p = doc.add_paragraph()
        p.add_run(
            f"A total of {diag.get('n_pairs_total', 'n/a'):,} sample pairs were evaluated using "
            f"KING-robust kinship coefficients (PLINK2 --make-king-table). Using a PI_HAT "
            f"threshold of {diag.get('pi_hat_threshold', 'n/a')}, "
            f"{diag.get('n_pairs_above_threshold', 'n/a'):,} pairs "
            f"({diag.get('pct_pairs_above_threshold', 'n/a')}% of all pairs) exceeded the "
            f"threshold." if isinstance(diag.get('n_pairs_total'), int) else ""
        )

    counts_csv = kinship_dir / "kinship_category_counts.csv"
    if counts_csv.exists():
        df = pd.read_csv(counts_csv)
        df = df.rename(columns={
            "categoria": "Relatedness category",
            "n_coppie": "N pairs",
            "pct_coppie": "% of pairs",
        })
        translate = {
            "duplicato/gemello monozigote": "Duplicate / monozygotic twin",
            "parente di 1o grado": "1st-degree relative",
            "parente di 2o grado": "2nd-degree relative",
            "parente di 3o grado": "3rd-degree relative",
            "non imparentati": "Unrelated",
        }
        if "Relatedness category" in df.columns:
            df["Relatedness category"] = df["Relatedness category"].replace(translate)
        if "% of pairs" in df.columns:
            df["% of pairs"] = df["% of pairs"].round(3)
        label = counters.table()
        add_caption(doc, label, "Distribution of pairwise relatedness categories (KING thresholds).")
        add_df_table(doc, df)
        doc.add_paragraph()
    else:
        add_missing_note(doc, "kinship category counts table", counts_csv)

    cross_batch_csv = kinship_dir / "kinship_cross_batch_duplicates_suspect.csv"
    if cross_batch_csv.exists():
        n_cross = sum(1 for _ in open(cross_batch_csv)) - 1
        if n_cross > 0:
            p = doc.add_paragraph()
            run = p.add_run(
                f"Note: {n_cross} pairs showed duplicate/1st-degree-level kinship despite "
                f"having different sample IDs across genotyping batches, consistent with the "
                f"same individual having been genotyped more than once under different IDs. "
                f"These were flagged for manual review (see kinship_cross_batch_duplicates_suspect.csv)."
            )
            run.italic = True

    fig_candidates = [kinship_dir.parent / "kinship_distribution.png",
                       kinship_dir / "kinship_distribution.png",
                       diagnostics_dir / "pi_hat_distribution.png"]
    fig_path = next((f for f in fig_candidates if f.exists()), None)
    if fig_path:
        flabel = counters.figure()
        add_figure(doc, fig_path)
        add_caption(doc, flabel, "Distribution of pairwise kinship coefficients.", above=False)
    else:
        add_missing_note(doc, "kinship distribution figure", fig_candidates[0])
    doc.add_paragraph()


def build_sex_check(doc, qc_dir, supp_plots_dir, counters):
    doc.add_heading("5. Sex Check", level=1)
    sexcheck_path = qc_dir / "sex_check.sexcheck"
    if not sexcheck_path.exists():
        add_missing_note(doc, "sex-check table", sexcheck_path)
        return

    df = read_plink_table(sexcheck_path)
    n_total = len(df)
    n_problem = int((df["STATUS"] == "PROBLEM").sum()) if "STATUS" in df.columns else None

    label = counters.table()
    add_caption(doc, label, "Genetic sex check summary (PLINK2 --check-sex).")
    summary_df = pd.DataFrame({
        "Metric": ["Samples evaluated", "Sex mismatches flagged (STATUS = PROBLEM)"],
        "Value": [n_total, n_problem if n_problem is not None else "n/a"],
    })
    add_df_table(doc, summary_df, col_widths_in=[3.5, 2.5])
    doc.add_paragraph()

    if n_problem:
        flagged_csv = supp_plots_dir / "sex_check_flagged_samples.csv"
        if flagged_csv.exists():
            p = doc.add_paragraph()
            run = p.add_run(
                f"Flagged sample IDs are listed in {flagged_csv.name}; these should be "
                f"cross-checked against reported sex in the metadata and excluded or resolved "
                f"before downstream analysis."
            )
            run.italic = True

    fig_path = supp_plots_dir / "sex_check_distribution.png"
    if fig_path.exists():
        flabel = counters.figure()
        add_figure(doc, fig_path)
        add_caption(doc, flabel, "Distribution of chrX heterozygosity F-statistic by sex-check status.", above=False)
    else:
        add_missing_note(doc, "sex-check distribution figure", fig_path)
    doc.add_paragraph()


def build_heterozygosity(doc, qc_dir, supp_plots_dir, counters, sd_threshold=3.0):
    doc.add_heading("6. Heterozygosity Outlier Check", level=1)
    het_path = qc_dir / "heterozygosity.het"
    if not het_path.exists():
        add_missing_note(doc, "heterozygosity table", het_path)
        return

    df = read_plink_table(het_path)
    n_total = len(df)
    n_outliers = "n/a"
    if "F" in df.columns:
        mean_f, sd_f = df["F"].mean(), df["F"].std()
        lower, upper = mean_f - sd_threshold * sd_f, mean_f + sd_threshold * sd_f
        n_outliers = int(((df["F"] < lower) | (df["F"] > upper)).sum())

    label = counters.table()
    add_caption(doc, label, "Heterozygosity outlier summary (LD-pruned autosomal SNPs).")
    summary_df = pd.DataFrame({
        "Metric": ["Samples evaluated", f"Outliers beyond +/-{sd_threshold:.0f} SD of mean F"],
        "Value": [n_total, n_outliers],
    })
    add_df_table(doc, summary_df, col_widths_in=[3.5, 2.5])
    doc.add_paragraph()

    fig_path = supp_plots_dir / "heterozygosity_distribution.png"
    if fig_path.exists():
        flabel = counters.figure()
        add_figure(doc, fig_path)
        add_caption(doc, flabel, "Distribution of the heterozygosity F-statistic across samples.", above=False)
    else:
        add_missing_note(doc, "heterozygosity distribution figure", fig_path)
    doc.add_paragraph()


def build_population_structure(doc, kinship_dir, counters):
    doc.add_heading("7. Population Structure and Batch Effects", level=1)

    eta_csv = kinship_dir / "pca_batch_eta2.csv"
    if eta_csv.exists():
        df = pd.read_csv(eta_csv)
        df = df.rename(columns={
            "PC": "Principal component",
            "eta_squared": "Eta-squared (variance explained by batch)",
            "batch_effect_alto": "Substantial batch effect (eta^2 > 0.1)",
        })
        label = counters.table()
        add_caption(doc, label, "Fraction of principal-component variance explained by genotyping batch.")
        add_df_table(doc, df)
        doc.add_paragraph()
    else:
        add_missing_note(doc, "PC-batch eta-squared table", eta_csv)

    for fname, desc in [
        ("pca_scatter_by_batch.png", "Top two principal components, colored by genotyping batch."),
        ("pca_scree_plot.png", "Scree plot of variance explained by each principal component."),
    ]:
        fig_path = kinship_dir / fname
        if fig_path.exists():
            flabel = counters.figure()
            add_figure(doc, fig_path)
            add_caption(doc, flabel, desc, above=False)
            doc.add_paragraph()
        else:
            add_missing_note(doc, desc, fig_path)


def build_pc_exposure(doc, diagnostics_dir, counters):
    doc.add_heading("8. Principal Components vs. Exposure", level=1)

    diag_report = diagnostics_dir / "diagnostics_report.txt"
    diag = parse_diagnostics_report(diag_report)
    if diag.get("pc_exposure_r2") is not None:
        p = doc.add_paragraph()
        p.add_run(
            f"Regressing the exposure variable on the top principal components yielded "
            f"R\u00b2 = {diag['pc_exposure_r2']:.4f} ({diag['pc_exposure_r2_pct']:.1f}% of "
            f"exposure variance explained by genetic ancestry), based on "
            f"{diag.get('n_matched', 'n/a')} matched samples."
        )

    corr_csv = diagnostics_dir / "pc_exposure_correlation.csv"
    if corr_csv.exists():
        df = pd.read_csv(corr_csv)
        df["p_value"] = df["p_value"].apply(lambda x: f"{x:.3g}")
        df = df.rename(columns={
            "PC": "Principal component",
            "pearson_r": "Pearson r",
            "p_value": "p-value",
            "notevole": "Notable (|r|>0.2, p<0.05)",
        })
        label = counters.table()
        add_caption(doc, label, "Correlation between each principal component and the exposure variable.")
        add_df_table(doc, df)
        doc.add_paragraph()
    else:
        add_missing_note(doc, "PC-exposure correlation table", corr_csv)

    fig_path = diagnostics_dir / "pca_vs_exposure.png"
    if fig_path.exists():
        flabel = counters.figure()
        add_figure(doc, fig_path)
        add_caption(doc, flabel, "PC1 vs. PC2, colored by exposure value.", above=False)
        doc.add_paragraph()
    else:
        add_missing_note(doc, "PC1 vs PC2 by exposure figure", fig_path)


def build_maf_and_missingness(doc, supp_plots_dir, counters):
    doc.add_heading("9. Allele Frequency Spectrum and Missingness", level=1)

    for fname, desc in [
        ("maf_spectrum.png", "Minor allele frequency (MAF) spectrum after filtering."),
        ("missingness_distributions.png", "Per-variant and per-sample missingness distributions, pre-filter, with QC thresholds indicated."),
    ]:
        fig_path = supp_plots_dir / fname
        if fig_path.exists():
            flabel = counters.figure()
            add_figure(doc, fig_path)
            add_caption(doc, flabel, desc, above=False)
            doc.add_paragraph()
        else:
            add_missing_note(doc, desc, fig_path)


def build_genomic_inflation(doc, diagnostics_dir, counters):
    doc.add_heading("10. Genomic Inflation", level=1)

    diag_report = diagnostics_dir / "diagnostics_report.txt"
    diag = parse_diagnostics_report(diag_report)
    if diag.get("lambda_gc") is not None:
        label = counters.table()
        add_caption(doc, label, "Genomic inflation factor summary.")
        df = pd.DataFrame({
            "Metric": ["N tests", "Lambda GC"],
            "Value": [f"{int(diag['lambda_gc_n_tests']):,}", f"{diag['lambda_gc']:.4f}"],
        })
        add_df_table(doc, df, col_widths_in=[3.0, 3.0])
        doc.add_paragraph()
    else:
        add_missing_note(doc, "genomic inflation (lambda GC) summary", diag_report)

    fig_path = diagnostics_dir / "qq_plot.png"
    if fig_path.exists():
        flabel = counters.figure()
        add_figure(doc, fig_path)
        add_caption(doc, flabel, "Quantile-quantile (QQ) plot of observed vs. expected -log10(p) values.", above=False)
    else:
        add_missing_note(doc, "QQ-plot figure", fig_path)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Build an English Word supplementary QC report from pipeline output files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--qc-dir", required=True, type=Path,
                         help="main pipeline out_dir (00_run_plink_qc.sh / 01_run_extra_qc_checks.sh output)")
    parser.add_argument("--kinship-report-dir", type=Path, default=None,
                         help="out-dir used for qc_report.py (default: <qc-dir>/qc_report)")
    parser.add_argument("--diagnostics-dir", type=Path, default=None,
                         help="out-dir used for interpret_plink_output.py (default: <qc-dir>/diagnostics_output)")
    parser.add_argument("--attrition-csv", type=Path, default=None,
                         help="CSV produced by qc_attrition_summary.py (default: <qc-dir>/qc_attrition.csv)")
    parser.add_argument("--supp-plots-dir", type=Path, default=None,
                         help="out-dir used for qc_supplementary_plots.py (default: <qc-dir>/supplementary_plots)")
    parser.add_argument("--out", required=True, type=Path, help="output .docx path")
    args = parser.parse_args()

    qc_dir = args.qc_dir
    kinship_dir = args.kinship_report_dir or (qc_dir / "qc_report")
    diagnostics_dir = args.diagnostics_dir or (qc_dir / "diagnostics_output")
    attrition_csv = args.attrition_csv or (qc_dir / "qc_attrition.csv")
    supp_plots_dir = args.supp_plots_dir or (qc_dir / "supplementary_plots")

    doc = Document()
    counters = FigureCounter()

    build_title(doc, qc_dir)
    build_intro(doc)
    build_reproducibility(doc, qc_dir, counters)
    build_attrition(doc, attrition_csv, counters)
    build_relatedness(doc, kinship_dir, diagnostics_dir, counters)
    build_sex_check(doc, qc_dir, supp_plots_dir, counters)
    build_heterozygosity(doc, qc_dir, supp_plots_dir, counters)
    build_population_structure(doc, kinship_dir, counters)
    build_pc_exposure(doc, diagnostics_dir, counters)
    build_maf_and_missingness(doc, supp_plots_dir, counters)
    build_genomic_inflation(doc, diagnostics_dir, counters)

    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(args.out)
    print(f"Supplementary report saved to: {args.out}")
    print(f"Tables inserted: {counters.table_n} | Figures inserted: {counters.figure_n}")


if __name__ == "__main__":
    main()

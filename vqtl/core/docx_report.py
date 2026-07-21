"""
Step 9 - Word (.docx) export: Results + Supplementary Material tables/figures
ready to paste into a manuscript.

Nota sulla lingua: a differenza del resto del repo (commenti in italiano,
per coerenza con lo stile di gene_environment/quality_control), il
CONTENUTO di questo documento -- titoli, didascalie, intestazioni di
colonna, note a pie' di tabella -- e' interamente in INGLESE, perche' e' il
deliverable che finisce davvero nel paper/supplementary material, non
codice interno. I commenti del codice restano in italiano come nel resto
del progetto.

Cosa genera, in un unico .docx per coorte (VQTL_RESULTS_DIR/gen<N>/report.docx):

  RESULTS (corpo del paper)
    Table 1.  Top genome-wide vQTL loci (variance-effect screen)
    Table 2.  SNP x exposure interaction tests reaching nominal significance
              (candidate loci; il set completo testato e' in Supplementary Table S2)
    Table 3.  Permutation-based validation of prioritized loci (interazione +
              test di Levene permutazionale sulla varianza per genotipo)
    Figure 1. Manhattan plot
    Figure 2. QQ plot
    Figure 3. Forest plot of prioritized interaction loci

  SUPPLEMENTARY MATERIAL
    Table S1. Full genome-wide vQTL scan (troncata a docx_supp_max_rows righe,
              con nota che rimanda alla tabella DB vqtl_scan_results se lo
              scan e' piu' lungo)
    Table S2. Full SNP x exposure interaction results (TUTTI i candidati
              testati, non solo quelli significativi di Table 2)
    Table S3. Reactive gene-environment correlation (rGE) and heteroscedasticity screen
    Table S4. Robustness of prioritized loci across phenotype transformations
    Supplementary Figures S1..Sn: per-locus genotype/exposure plots (top loci)

Stile tabelle: "three-line table" (bordo sopra, bordo sotto l'header, bordo
in fondo, nessuna riga/colonna interna) -- la convenzione piu' diffusa nelle
riviste scientifiche (Nature, Cell, JAMA, ecc.), niente sfondo colorato.
"""
from __future__ import annotations

import os

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from gene_environment.logging_utils import get_logger

from vqtl.config import VqtlConfig

log = get_logger(__name__)

FONT_NAME = "Times New Roman"
BODY_SIZE = 10
CAPTION_SIZE = 10
TABLE_FONT_SIZE = 9

# Documento in landscape con margini ridotti: le tabelle "paper-ready" hanno
# spesso 8-11 colonne (SNP, Chr, Position, N, MAF, beta, SE, Z, P, P_gc,
# FDR...) che in portrait/margini standard costringono Word a spezzare le
# parole a meta' (osservato es. "risaie_1500" -> "risaie_150" + "0",
# "corrected" -> "correcte" + "d") -- inaccettabile per un documento
# destinato a un manoscritto. Landscape US Letter + margini 0.6in da'
# ~9.8in di larghezza utile, sufficiente per le tabelle di questa pipeline
# senza spezzare i token.
PAGE_MARGIN_IN = 0.6
PAGE_WIDTH_IN = 11.0
PAGE_HEIGHT_IN = 8.5
USABLE_WIDTH_IN = PAGE_WIDTH_IN - 2 * PAGE_MARGIN_IN

# Larghezza "preferita" (pollici) per ciascuna colonna, sufficiente a
# contenere il contenuto piu' lungo atteso senza andare a capo a meta'
# parola. Se la somma delle colonne di una tabella supera USABLE_WIDTH_IN,
# _add_data_table la scala proporzionalmente (vedi _set_column_widths).
_COLUMN_WIDTHS_IN: dict[str, float] = {
    "SNP": 1.15, "CHR": 0.5, "POS": 0.95, "N": 0.5, "MAF": 0.6,
    "beta_QI": 0.75, "SE": 0.75, "Z": 0.6, "P": 0.75, "P_gc": 0.95, "fdr_gc": 0.75,
    "exposure": 1.15, "beta_I": 0.85, "pval": 0.75,
    "beta_I_observed": 1.05, "n_perm_valid": 0.95, "empirical_pval": 0.85, "asymptotic_pval": 0.85,
    "levene_stat_observed": 0.95, "levene_pval": 0.85,
    "rGE_beta_exposure_on_snp": 1.15, "rGE_SE": 0.65, "rGE_pval": 0.75, "rGE_flag": 0.85,
    "het_BP_lm_pvalue": 0.95, "heteroscedasticity_flag": 1.05,
    "variant": 1.35,
}
_DEFAULT_COLUMN_WIDTH_IN = 0.85


# ------------------------------------------------------------------
# Helper di basso livello per lo stile "three-line table"
# (python-docx non ha un'API di alto livello per i bordi delle celle)
# ------------------------------------------------------------------
def _set_cell_border(cell, edge: str, sz: int = 8, color: str = "000000") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    tag = f"w:{edge}"
    el = tc_borders.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        tc_borders.append(el)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)


def _apply_three_line_style(table) -> None:
    """Bordo doppio-spesso sopra la tabella e sotto l'ultima riga, bordo
    sottile sotto l'header -- niente altre righe/griglie/sfondi."""
    n_rows = len(table.rows)
    for cell in table.rows[0].cells:
        _set_cell_border(cell, "top", sz=12)
    for cell in table.rows[0].cells:
        _set_cell_border(cell, "bottom", sz=6)
    for cell in table.rows[n_rows - 1].cells:
        _set_cell_border(cell, "bottom", sz=12)


def _set_font(run, bold: bool = False, italic: bool = False, size: float = BODY_SIZE) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # forza il font anche per il rendering "East Asian"/complex script di Word
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), FONT_NAME)


def _add_paragraph(doc, text: str = "", bold: bool = False, italic: bool = False,
                    size: float = BODY_SIZE, space_after: float = 6, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _set_font(run, bold=bold, italic=italic, size=size)
    return p


def _add_heading(doc, text: str, level: int = 1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    _set_font(run, bold=True, size=14 if level == 1 else 12)
    return h


def _add_table_caption(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}. ")
    _set_font(r1, bold=True, size=CAPTION_SIZE)
    r2 = p.add_run(text)
    _set_font(r2, bold=False, italic=True, size=CAPTION_SIZE)
    return p


def _add_figure_caption(doc, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    r1 = p.add_run(f"{label}. ")
    _set_font(r1, bold=True, size=CAPTION_SIZE)
    r2 = p.add_run(text)
    _set_font(r2, bold=False, size=CAPTION_SIZE)
    return p


def _add_footnote_text(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(text)
    _set_font(run, size=8, italic=True)
    return p


# ------------------------------------------------------------------
# Formattazione numerica in stile paper
# ------------------------------------------------------------------
def _fmt_p(p) -> str:
    if p is None or (isinstance(p, float) and np.isnan(p)):
        return "NA"
    p = float(p)
    if p < 0.001:
        return f"{p:.2e}"
    return f"{p:.3f}"


def _fmt_float(x, nd: int = 3) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "NA"
    return f"{float(x):.{nd}f}"


def _fmt_int(x) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "NA"
    return str(int(x))


# Colonna sorgente -> (intestazione paper-ready, funzione di formattazione)
_COLUMN_SPECS: dict[str, tuple[str, callable]] = {
    "SNP": ("Variant", str),
    "CHR": ("Chr", str),
    "POS": ("Position (bp)", _fmt_int),
    "N": ("N", _fmt_int),
    "MAF": ("MAF", lambda v: _fmt_float(v, 3)),
    "beta_QI": ("\u03b2_QI", lambda v: _fmt_float(v, 4)),
    "SE": ("SE", lambda v: _fmt_float(v, 4)),
    "Z": ("Z", lambda v: _fmt_float(v, 3)),
    "P": ("P", _fmt_p),
    "P_gc": ("P (GC-corrected)", _fmt_p),
    "fdr_gc": ("FDR (BH)", _fmt_p),
    "exposure": ("Exposure", str),
    "beta_I": ("\u03b2_interaction", lambda v: _fmt_float(v, 4)),
    "pval": ("P", _fmt_p),
    "beta_I_observed": ("Observed \u03b2_interaction", lambda v: _fmt_float(v, 4)),
    "n_perm_valid": ("N permutations", _fmt_int),
    "empirical_pval": ("Empirical P", _fmt_p),
    "asymptotic_pval": ("Asymptotic P", _fmt_p),
    "levene_stat_observed": ("Levene stat", lambda v: _fmt_float(v, 3)),
    "levene_pval": ("Levene empirical P", _fmt_p),
    "rGE_beta_exposure_on_snp": ("\u03b2 (exposure \u2192 genotype)", lambda v: _fmt_float(v, 4)),
    "rGE_SE": ("SE", lambda v: _fmt_float(v, 4)),
    "rGE_pval": ("rGE P", _fmt_p),
    "rGE_flag": ("rGE flagged", lambda v: "Yes" if bool(v) else "No"),
    "het_BP_lm_pvalue": ("Breusch-Pagan P", _fmt_p),
    "heteroscedasticity_flag": ("Heteroscedasticity flagged", lambda v: "Yes" if bool(v) else "No"),
    "variant": ("Phenotype variant", lambda v: {
        "original": "Original", "log_transform": "Log-transformed",
        "rank_inverse_normal": "Rank inverse-normal", "outliers_removed": "Outliers removed (|z|>3)",
    }.get(v, str(v))),
}


def _set_column_widths(table, widths_in: list[float]) -> None:
    """Larghezze di colonna FISSE (non autofit): necessario per evitare che
    Word/LibreOffice spezzino a meta' i token che non contengono spazi
    (id varianti, nomi esposizione) quando la colonna e' troppo stretta.
    Va impostata sia a livello di tabella (tblLayout=fixed) sia su ogni
    singola cella della colonna (Word ignora larghezze parziali/incoerenti
    tra celle della stessa colonna)."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)


def _column_widths_for(cols: list[str]) -> list[float]:
    raw = [_COLUMN_WIDTHS_IN.get(c, _DEFAULT_COLUMN_WIDTH_IN) for c in cols]
    total = sum(raw)
    if total <= USABLE_WIDTH_IN or total == 0:
        return raw
    # scala proporzionalmente se la somma eccede la larghezza utile della pagina
    scale = USABLE_WIDTH_IN / total
    return [w * scale for w in raw]


def _paper_table(df: pd.DataFrame, columns: list[str]) -> tuple[list[str], list[list[str]], list[float]]:
    """Ritorna (intestazioni_inglesi, righe_formattate, larghezze_pollici)
    per le sole colonne richieste che esistono davvero in df, nell'ordine
    richiesto."""
    cols = [c for c in columns if c in df.columns]
    headers = [_COLUMN_SPECS.get(c, (c, str))[0] for c in cols]
    rows = []
    for _, row in df.iterrows():
        rows.append([_COLUMN_SPECS.get(c, (c, str))[1](row[c]) for c in cols])
    widths = _column_widths_for(cols)
    return headers, rows, widths


def _add_data_table(doc, headers: list[str], rows: list[list[str]], col_widths_in: list[float] | None = None) -> None:
    if not rows:
        _add_paragraph(doc, "No rows to display.", italic=True, size=BODY_SIZE)
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_font(run, bold=True, size=TABLE_FONT_SIZE)

    for row_vals in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_vals):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            _set_font(run, size=TABLE_FONT_SIZE)

    _apply_three_line_style(table)
    if col_widths_in:
        _set_column_widths(table, col_widths_in)


def _add_image_full_width(doc, path: str, width_in: float = 6.0) -> bool:
    if not os.path.exists(path):
        log.warning("Immagine non trovata, saltata nel docx: %s", path)
        return False
    doc.add_picture(path, width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def _add_image_pair(doc, left_path: str, right_path: str, width_in: float = 3.1) -> bool:
    """Due immagini affiancate in una tabella senza bordi (boxplot + scatter
    dello stesso locus), per tenerle vicine nello stesso paragrafo/figura."""
    if not (os.path.exists(left_path) and os.path.exists(right_path)):
        log.warning("Coppia di immagini incompleta, saltata: %s / %s", left_path, right_path)
        return False
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    for cell, path in zip(table.rows[0].cells, [left_path, right_path]):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
    return True


# ------------------------------------------------------------------
# Costruzione del documento
# ------------------------------------------------------------------
def build_docx_report(
    vcfg: VqtlConfig, cohort_dir: str, generation: int, n_samples: int,
    exposures: list[str], target_col: str,
    vqtl_df: pd.DataFrame, candidates: pd.DataFrame, interaction_df: pd.DataFrame,
    interaction_significant_df: pd.DataFrame,
    rge_df: pd.DataFrame, perm_df: pd.DataFrame, robustness_df: pd.DataFrame,
) -> str:
    fig_dir = os.path.join(cohort_dir, "figures")
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.left_margin = section.right_margin = Inches(PAGE_MARGIN_IN)
    section.top_margin = section.bottom_margin = Inches(PAGE_MARGIN_IN)
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(BODY_SIZE)

    # ---- Title ----
    _add_paragraph(
        doc, f"vQTL / Gene-by-Environment Interaction Analysis \u2014 Cohort gen{generation}",
        bold=True, size=15, space_after=2,
    )
    _add_paragraph(
        doc,
        f"Phenotype: {target_col}. Exposure(s) tested: {', '.join(exposures)}. "
        f"Analysis-ready sample size: n = {n_samples}.",
        italic=True, size=10, space_after=16,
    )

    # ================= RESULTS =================
    _add_heading(doc, "Results", level=1)

    _add_paragraph(
        doc,
        "A genome-wide variance-effect (vQTL) screen was performed to prioritize variants "
        "whose dosage modulates the dispersion of the phenotype, followed by SNP-by-exposure "
        "interaction testing for the prioritized loci, and permutation-based validation of the "
        "top interaction estimates. Full methodological detail is provided in the Methods "
        "section; abbreviations are defined in the notes below each table.",
        size=10, space_after=12,
    )

    top_scan = vqtl_df.sort_values("P_gc" if "P_gc" in vqtl_df.columns else "P").head(vcfg.docx_top_n_scan) if not vqtl_df.empty else vqtl_df
    _add_table_caption(doc, "Table 1", f"Top genome-wide vQTL loci ranked by GC-corrected p-value (top {len(top_scan)} of {len(vqtl_df)} variants tested).")
    headers, rows, widths = _paper_table(top_scan, ["SNP", "CHR", "POS", "N", "MAF", "beta_QI", "SE", "Z", "P", "P_gc", "fdr_gc"])
    _add_data_table(doc, headers, rows, widths)
    _add_footnote_text(
        doc,
        "Abbreviations: Chr, chromosome; MAF, minor allele frequency; \u03b2_QI, quantile-integrated "
        "variance effect (QUAIL-style scan); SE, standard error (asymptotic); P, asymptotic p-value; "
        "P (GC-corrected), p-value after genomic-control correction for inflation; FDR (BH), "
        "Benjamini-Hochberg false discovery rate across all variants tested.",
    )

    _add_table_caption(
        doc, "Table 2",
        f"SNP \u00d7 exposure interaction tests reaching nominal significance "
        f"(P < {vcfg.interaction_sig_threshold}) among loci prioritized by the genome-wide vQTL screen "
        f"({len(interaction_significant_df)} of {len(interaction_df)} candidate SNP \u00d7 exposure pairs tested; "
        f"the full set is reported in Supplementary Table S2).",
    )
    headers, rows, widths = _paper_table(interaction_significant_df, ["SNP", "CHR", "POS", "exposure", "beta_I", "SE", "pval", "N", "MAF"])
    _add_data_table(doc, headers, rows, widths)
    _add_footnote_text(
        doc,
        "Model: phenotype ~ SNP + exposure + SNP\u00d7exposure + covariates, fit by ordinary least "
        "squares (or logistic regression for a binary phenotype) with heteroscedasticity-robust "
        "standard errors (HC3/HC1). \u03b2_interaction is the SNP\u00d7exposure interaction coefficient. "
        "P-values are nominal (not corrected for multiple testing across loci/exposures).",
    )

    if not perm_df.empty:
        _add_table_caption(
            doc, "Table 3",
            "Permutation-based validation of the top prioritized loci: SNP\u00d7exposure interaction "
            "(Freedman-Lane permutation) and genotype variance heterogeneity (permutation-based Levene test).",
        )
        headers, rows, widths = _paper_table(
            perm_df,
            ["SNP", "CHR", "POS", "exposure", "beta_I_observed", "n_perm_valid", "empirical_pval",
             "asymptotic_pval", "levene_stat_observed", "levene_pval"],
        )
        _add_data_table(doc, headers, rows, widths)
        _add_footnote_text(
            doc,
            "Interaction empirical P = (1 + number of permutations with |\u03b2_interaction| \u2265 observed) / "
            "(N permutations + 1). Permutations preserve the SNP, exposure and covariate main-effect "
            "structure (Freedman & Lane, 1983), testing the null hypothesis of no interaction effect "
            "specifically. Levene stat / Levene empirical P: variance-heterogeneity test for the "
            "covariate-residualized phenotype across genotype groups (0/1/2), median-centered "
            "(Brown-Forsythe); the null distribution is built by permuting genotype labels (not "
            "residuals) and recomputing the statistic, so it makes no distributional assumption beyond "
            "exchangeability under the null -- an assumption-light confirmation of the variance effect "
            "detected by the genome-wide scan (Table 1) for this locus. One-sided (larger statistic = "
            "more evidence of heteroscedasticity). Independent of exposure: identical across rows for "
            "the same variant.",
        )

    fig_added = False
    if _add_image_full_width(doc, os.path.join(fig_dir, "manhattan_vqtl.png"), width_in=8.5):
        _add_figure_caption(doc, "Figure 1", "Manhattan plot of the genome-wide vQTL scan. Dashed lines mark genome-wide (P = 5\u00d710\u207b\u2078) and suggestive (P = 1\u00d710\u207b\u2075) significance thresholds.")
        fig_added = True
    if _add_image_full_width(doc, os.path.join(fig_dir, "qq_vqtl.png"), width_in=4.5):
        _add_figure_caption(doc, "Figure 2", "Quantile-quantile plot of the genome-wide vQTL scan p-values, with the genomic inflation factor (\u03bb_GC).")
        fig_added = True
    if _add_image_full_width(doc, os.path.join(fig_dir, "forest_top_interactions.png"), width_in=7.0):
        _add_figure_caption(doc, "Figure 3", "Forest plot of SNP\u00d7exposure interaction effect estimates (\u03b2_interaction) with 95% confidence intervals for the top prioritized loci.")
        fig_added = True
    if not fig_added:
        _add_paragraph(doc, "No figures were available for this cohort.", italic=True)

    # ================= SUPPLEMENTARY MATERIAL =================
    doc.add_page_break()
    _add_heading(doc, "Supplementary Material", level=1)

    n_scan_total = len(vqtl_df)
    scan_supp = vqtl_df.sort_values("P_gc" if "P_gc" in vqtl_df.columns else "P").head(vcfg.docx_supp_max_rows) if not vqtl_df.empty else vqtl_df
    truncated_note = (
        f" Showing the top {len(scan_supp)} of {n_scan_total} variants tested, ranked by GC-corrected "
        f"p-value; the complete table is available in the vqtl_scan_results database table."
        if n_scan_total > len(scan_supp) else ""
    )
    _add_table_caption(doc, "Supplementary Table S1", f"Genome-wide vQTL scan results.{truncated_note}")
    headers, rows, widths = _paper_table(scan_supp, ["SNP", "CHR", "POS", "N", "MAF", "beta_QI", "SE", "Z", "P", "P_gc", "fdr_gc"])
    _add_data_table(doc, headers, rows, widths)
    _add_footnote_text(doc, "See Table 1 note for column definitions.")

    _add_table_caption(
        doc, "Supplementary Table S2",
        "Full SNP \u00d7 exposure interaction test results for all candidate loci (Table 1 / Supplementary "
        "Table S1 candidates), including pairs not reaching nominal significance (shown in Table 2).",
    )
    headers, rows, widths = _paper_table(interaction_df, ["SNP", "CHR", "POS", "exposure", "beta_I", "SE", "pval", "N", "MAF"])
    _add_data_table(doc, headers, rows, widths)
    _add_footnote_text(doc, "See Table 2 note for model and column definitions.")

    if not rge_df.empty:
        _add_table_caption(doc, "Supplementary Table S3", "Reactive gene-environment correlation (rGE) and heteroscedasticity screen for candidate SNP \u00d7 exposure pairs.")
        headers, rows, widths = _paper_table(
            rge_df,
            ["SNP", "CHR", "POS", "exposure", "rGE_beta_exposure_on_snp", "rGE_SE", "rGE_pval", "rGE_flag",
             "het_BP_lm_pvalue", "heteroscedasticity_flag"],
        )
        _add_data_table(doc, headers, rows, widths)
        _add_footnote_text(
            doc,
            "rGE model: genotype ~ exposure + covariates (OLS, HC3 SE); a significant \u03b2 suggests "
            "genotype and exposure are not independent in this sample and does not by itself invalidate "
            "an interaction finding, but complicates its causal interpretation. Heteroscedasticity: "
            "Breusch-Pagan test on the residuals of phenotype ~ SNP + exposure + covariates (no "
            "interaction term); flagged loci already use heteroscedasticity-robust SEs in Table 2/S4 "
            "by default.",
        )

    if not robustness_df.empty:
        _add_table_caption(doc, "Supplementary Table S4", "Robustness of prioritized interaction loci to phenotype transformation and outlier removal.")
        headers, rows, widths = _paper_table(robustness_df, ["SNP", "CHR", "POS", "exposure", "variant", "beta_I", "SE", "pval", "N", "MAF"])
        _add_data_table(doc, headers, rows, widths)
        _add_footnote_text(
            doc,
            "Each prioritized locus (Table 3) was re-tested on the original phenotype, its "
            "log-transform, a rank-based inverse-normal transform, and after removing outliers "
            "(|z| > 3). A locus whose \u03b2_interaction direction and significance are stable across "
            "these variants is less likely to reflect an artifact of the phenotype distribution or of "
            "a small number of influential observations.",
        )

    # Supplementary figures: per-locus boxplot + scatter, stessi loci del report.md (top 10)
    top_interactions = interaction_df.sort_values("pval").head(10).reset_index(drop=True) if not interaction_df.empty else interaction_df
    supp_fig_n = 0
    for i, row in top_interactions.iterrows():
        snp_id, exp_raw = row["SNP"], row["exposure"]
        box_path = os.path.join(fig_dir, f"locus{i + 1}_{snp_id}_boxplot.png")
        scatter_path = os.path.join(fig_dir, f"locus{i + 1}_{snp_id}_scatter.png")
        if _add_image_pair(doc, box_path, scatter_path, width_in=4.3):
            supp_fig_n += 1
            _add_figure_caption(
                doc, f"Supplementary Figure S{supp_fig_n}",
                f"{snp_id} \u00d7 {exp_raw}. Left: phenotype distribution by genotype (dosage 0/1/2). "
                f"Right: phenotype vs. standardized exposure, stratified and fitted separately by genotype.",
            )
    if supp_fig_n == 0:
        _add_paragraph(doc, "No supplementary per-locus figures were available for this cohort.", italic=True)

    out_path = os.path.join(cohort_dir, "report.docx")
    doc.save(out_path)
    log.info("Scritto %s (Results + Supplementary Material, in inglese)", out_path)
    return out_path
